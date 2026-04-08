import os
from pathlib import Path
from docx import Document
from datetime import datetime

UPLOAD_DIR = Path("uploads")
TEMPLATE_PATH = Path("template/NET- Standard NDA -  Mutual Disclsoure.docx")


def _replace_in_paragraph(para, replacements: dict):
    """Merge all runs into first run then apply replacements (handles split runs)."""
    # Rebuild full text by concatenating all runs
    full_text = "".join(r.text for r in para.runs)
    new_text = full_text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != full_text and para.runs:
        # Put everything into the first run, preserve its formatting
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""


def _replace_text_in_docx(doc, replacements: dict):
    """Replace placeholder text across all paragraphs and table cells."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)


def generate_nda_pdf(contract_id: str, company_name: str, description: str, expiry_date: str) -> str:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    # 1. Load the actual NDA DOCX template
    doc = Document(TEMPLATE_PATH)

    # 2. Build replacement map.
    # The template uses [●] (U+25CF), [______], and variations.
    today = datetime.now().strftime("%-d of %B, %Y")
    short_name = company_name.strip() if company_name.strip() else "____"

    replacements = {
        # Company name placeholders
        "[●]": short_name,
        "[●],": f"{short_name},",
        "[______]": short_name,
        # Date placeholders in first paragraph (underscore sequences)
        "_____ of ____,": today.rsplit(" of ", 1)[0] + " of",
        "____ of _____,": today.rsplit(" of ", 1)[0] + " of",
    }

    _replace_text_in_docx(doc, replacements)

    # 3. Render to PDF using reportlab (reading both paragraphs AND tables)
    pdf_path = UPLOAD_DIR / f"{contract_id}.pdf"
    _render_docx_as_pdf(doc, str(pdf_path), company_name)

    return str(pdf_path)


def generate_nda_docx(contract_id: str, company_name: str, description: str, expiry_date: str) -> str:
    """Generate NDA DOCX from the official template while preserving layout."""
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(TEMPLATE_PATH)

    today = datetime.now().strftime("%-d of %B, %Y")
    short_name = company_name.strip() if company_name.strip() else "____"
    purpose = (description or "").strip()

    replacements = {
        "[●]": short_name,
        "[●],": f"{short_name},",
        "[______]": short_name,
        "_____ of ____,": today.rsplit(" of ", 1)[0] + " of",
        "____ of _____,": today.rsplit(" of ", 1)[0] + " of",
        "[PURPOSE]": purpose,
        "[Purpose]": purpose,
        "[DESCRIPTION]": purpose,
        "[Description]": purpose,
        "[EXPIRY_DATE]": expiry_date or "",
    }

    _replace_text_in_docx(doc, replacements)
    docx_path = UPLOAD_DIR / f"{contract_id}_nda.docx"
    doc.save(str(docx_path))
    return str(docx_path)


def _render_docx_as_pdf(doc, pdf_path: str, company_name: str):
    """Faithfully render DOCX paragraphs AND tables into a PDF using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, HRFlowable, KeepTogether
    )
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT

    PAGE_W, PAGE_H = A4
    LM = RM = 2.5 * cm
    TM = BM = 2.0 * cm

    doc_pdf = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        rightMargin=RM, leftMargin=LM,
        topMargin=TM, bottomMargin=BM
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'NDATitle', parent=styles['Normal'],
        fontSize=10, fontName='Helvetica-Bold',
        alignment=TA_CENTER, spaceAfter=8, spaceBefore=4
    )
    body_style = ParagraphStyle(
        'NDABody', parent=styles['Normal'],
        fontSize=9, fontName='Helvetica',
        alignment=TA_JUSTIFY, spaceAfter=5, leading=14
    )
    bold_body_style = ParagraphStyle(
        'NDABoldBody', parent=styles['Normal'],
        fontSize=9, fontName='Helvetica-Bold',
        alignment=TA_JUSTIFY, spaceAfter=5, leading=14
    )
    heading_style = ParagraphStyle(
        'NDAHeading', parent=styles['Normal'],
        fontSize=9.5, fontName='Helvetica-Bold',
        spaceAfter=6, spaceBefore=10, leading=14
    )
    indent_style = ParagraphStyle(
        'NDAIndent', parent=styles['Normal'],
        fontSize=9, fontName='Helvetica',
        alignment=TA_JUSTIFY, spaceAfter=4, leading=14,
        leftIndent=18
    )
    sig_label_style = ParagraphStyle(
        'NDASignLabel', parent=styles['Normal'],
        fontSize=9, fontName='Helvetica-Bold',
        spaceAfter=2, spaceBefore=2
    )
    sig_name_style = ParagraphStyle(
        'NDASignName', parent=styles['Normal'],
        fontSize=9, fontName='Helvetica',
        spaceAfter=2
    )

    def _para_to_flowable(para):
        text = para.text.strip()
        if not text:
            return Spacer(1, 5)

        style_name = para.style.name if para.style else ""
        is_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())

        # Heading detection
        if "Heading 1" in style_name or (text.isupper() and len(text) < 80):
            return Paragraph(text.replace("&", "&amp;"), heading_style)
        if "Heading 2" in style_name:
            return Paragraph(text.replace("&", "&amp;"), bold_body_style)

        # Numbered clauses (e.g. "1.1  Payment")
        if len(text) > 2 and text[:2].replace(".", "").isdigit():
            return Paragraph(text.replace("&", "&amp;"), body_style)

        # Sub-items with letter prefix
        if len(text) > 2 and text[0] == "(" and text[2] == ")":
            return Paragraph(text.replace("&", "&amp;"), indent_style)

        # All bold paragraph
        if is_bold:
            return Paragraph(text.replace("&", "&amp;"), bold_body_style)

        return Paragraph(text.replace("&", "&amp;"), body_style)

    story = []

    # Track which paragraph index each table starts AFTER to interleave tables
    # Build ordered list of content: (type, obj) where type = 'para' | 'table'
    # We use the python-docx element_order approach
    from docx.oxml.ns import qn

    body_el = doc.element.body
    for child in body_el.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Find the paragraph
            for para in doc.paragraphs:
                if para._element is child:
                    story.append(_para_to_flowable(para))
                    break

        elif tag == "tbl":
            # Find the matching table and render it
            for tbl in doc.tables:
                if tbl._element is child:
                    _render_table(tbl, story, sig_label_style, sig_name_style, body_style)
                    break

    doc_pdf.build(story)


def _render_table(tbl, story, sig_label_style, sig_name_style, body_style):
    """Render a DOCX table as a reportlab Table flowable."""
    from reportlab.platypus import Table, TableStyle, Spacer, Paragraph, HRFlowable
    from reportlab.lib import colors
    from reportlab.lib.units import cm

    story.append(Spacer(1, 12))

    rows_data = []
    for row in tbl.rows:
        row_cells = []
        for cell in row.cells:
            lines = []
            for para in cell.paragraphs:
                t = para.text.strip()
                if t:
                    is_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
                    if is_bold or "For and on behalf" in t:
                        lines.append(Paragraph(t.replace("&", "&amp;"), sig_label_style))
                    elif set(t) == {"_"} or t.startswith("___"):
                        lines.append(HRFlowable(width="80%", thickness=0.5, color=colors.black, spaceBefore=16, spaceAfter=4))
                    else:
                        lines.append(Paragraph(t.replace("&", "&amp;"), sig_name_style))
                else:
                    lines.append(Paragraph("", sig_name_style))
            row_cells.append(lines)
        rows_data.append(row_cells)

    if not rows_data:
        return

    # Use equal column widths
    num_cols = max(len(r) for r in rows_data) if rows_data else 1
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    avail_width = A4[0] - 5 * cm  # LM + RM
    col_width = avail_width / num_cols

    rl_table = Table(rows_data, colWidths=[col_width] * num_cols)
    rl_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(rl_table)
