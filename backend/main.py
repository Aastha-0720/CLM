<<<<<<< Updated upstream
import asyncio
from fastapi import FastAPI, UploadFile, File, Form, Body, HTTPException, Query, BackgroundTasks, Header
import base64
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
from io import StringIO
from typing import List, Optional, Dict, Any, Union
=======
import logging
import tempfile
import zipfile
from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Query, Form, Header, Depends
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
>>>>>>> Stashed changes
from bson import ObjectId
from datetime import datetime, timedelta
import pathlib
import shutil
import re
<<<<<<< Updated upstream
import textwrap

from database import db
from models import Contract, CAS, DepartmentReviews, Document, ReviewComment, PyObjectId, AuditLog, ChangeRequest, ContractVersion, DOARule, DOARuleHistory
import services
from services import AIService
=======
from io import BytesIO
import xml.etree.ElementTree as ET

from database import db
from models import Contract, CAS, DepartmentReviews, Document, ReviewComment, PyObjectId, WhitelistEntry, DigInkDocument, PdfFile, PdfAnnotation
import services
from services.pdf_service import generate_nda_pdf, generate_nda_docx
import digink_service
from routes.digink_routes import router as digink_router
>>>>>>> Stashed changes
from openai import OpenAI
import os

UPLOAD_DIR = pathlib.Path(os.getenv("UPLOAD_DIR", "/app/uploads"))
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "50"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

<<<<<<< Updated upstream
=======
# Allowed upload formats
ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx'}
ALLOWED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/msword'
}

>>>>>>> Stashed changes
deepseek_client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY", ""),
    base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
)
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger

scheduler = AsyncIOScheduler()

import aiosmtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
SMTP_FROM_NAME = os.getenv("SMTP_FROM_NAME", "Apeiro CLM")

async def send_email_notification(to_email: str, subject: str, body_html: str):
    if not SMTP_HOST or not SMTP_USER:
        print("SMTP not configured, skipping email")
        return
    try:
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"] = f"{SMTP_FROM_NAME} <{SMTP_USER}>"
        msg["To"] = to_email
        msg.attach(MIMEText(body_html, "html"))
        await aiosmtplib.send(
            msg,
            hostname=SMTP_HOST,
            port=SMTP_PORT,
            username=SMTP_USER,
            password=SMTP_PASSWORD,
            start_tls=True,
        )
        print(f"Email sent to {to_email}")
    except Exception as e:
        print(f"Email send failed: {e}")

ai_service = AIService(deepseek_client, DEEPSEEK_MODEL)


def _clean_extracted_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _extract_plain_text_fallback(text: str) -> Dict[str, Any]:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    if not normalized:
        return {
            "counterparty": "",
            "contractValue": "",
            "effectiveDate": "",
            "expiryDate": "",
            "duration": "",
            "keyDates": "",
            "clauses": [],
            "warnings": ["No extractable text found in the uploaded file."]
        }

    counterparty = ""
    between_match = re.search(
        r"\bbetween\s+([A-Z][A-Za-z0-9&,. \-]{2,80}?)\s+and\s+([A-Z][A-Za-z0-9&,. \-]{2,80}?)(?:,|\.|\s{2,}|effective|dated|\n)",
        normalized,
        re.IGNORECASE,
    )
    if between_match:
        counterparty = between_match.group(2).strip(" ,.")

    value_match = re.search(r"((?:USD|EUR|GBP|INR|Rs\.?|₹|\$)\s?[0-9][0-9,]*(?:\.\d{1,2})?)", normalized, re.IGNORECASE)
    contract_value = value_match.group(1).strip() if value_match else ""

    def _extract_date(label_patterns):
        for pattern in label_patterns:
            match = re.search(pattern, normalized, re.IGNORECASE)
            if match:
                return match.group(1).strip(" .,:;")
        return ""

    effective_date = _extract_date([
        r"(?:effective\s+date|commencement\s+date)\s*[:\-]?\s*([A-Za-z0-9,\/\- ]{6,40})",
        r"effective\s+as\s+of\s+([A-Za-z0-9,\/\- ]{6,40})",
    ])
    expiry_date = _extract_date([
        r"(?:expiry|expiration|termination|end)\s+date\s*[:\-]?\s*([A-Za-z0-9,\/\- ]{6,40})",
        r"expires\s+on\s+([A-Za-z0-9,\/\- ]{6,40})",
    ])

    clause_matches = re.findall(
        r"((?:\d+(?:\.\d+)*)\s*[.)-]?\s*[A-Z][A-Za-z /&-]{2,60})(.*?)(?=(?:\d+(?:\.\d+)*)\s*[.)-]?\s*[A-Z][A-Za-z /&-]{2,60}|$)",
        text or "",
        re.DOTALL,
    )
    clauses = []
    for idx, (heading, body) in enumerate(clause_matches[:12]):
        heading = re.sub(r"\s+", " ", heading).strip()
        body = re.sub(r"\s+", " ", body).strip()
        if not body:
            continue
        clauses.append({
            "id": f"fallback-cl-{idx}",
            "title": heading,
            "type": heading,
            "content": body[:1200],
            "text": body[:1200],
            "department": "Legal",
        })

    return {
        "counterparty": counterparty,
        "contractValue": contract_value,
        "effectiveDate": effective_date,
        "expiryDate": expiry_date,
        "duration": f"{effective_date} to {expiry_date}".strip(" to"),
        "keyDates": f"Effective: {effective_date} | Expiry: {expiry_date}".strip(),
        "clauses": clauses,
        "warnings": []
    }

def get_current_user(authorization: Optional[str] = Header(None)) -> str:
    if not authorization or not isinstance(authorization, str) or not authorization.startswith("Bearer "):
        return "System"
    try:
        # Format: btoa(`${user.email}:${Date.now()}`)
        token = authorization.split(" ")[1]
        decoded = base64.b64decode(token).decode("utf-8")
        email = decoded.split(":")[0]
        return email
    except:
        return "Unknown"

AUDIT_EVENT_TYPES = {
    "CREATED": "CREATED",
    "REVIEW_APPROVED": "REVIEW_APPROVED",
    "WORKFLOW_COMPLETED": "WORKFLOW_COMPLETED",
    "CAS_GENERATED": "CAS_GENERATED",
    "DOA_STEP_APPROVED": "DOA_STEP_APPROVED",
    "DOA_COMPLETED": "DOA_COMPLETED",
    "CONTRACT_COMPLETED": "CONTRACT_COMPLETED",
    "ROUTED": "ROUTED",
    "REVIEW_REJECTED": "REVIEW_REJECTED",
    "CHANGES_REQUESTED": "CHANGES_REQUESTED",
    "RESUBMITTED_FOR_REVIEW": "RESUBMITTED_FOR_REVIEW",
    "UPDATED": "UPDATED",
    "REDLINE_APPLIED": "REDLINE_APPLIED",
}


def _audit_phase_for_event(event_type: str) -> str:
    if event_type in ["CREATED", "ROUTED"]:
        return "Intake"
    if event_type in ["REVIEW_APPROVED", "WORKFLOW_COMPLETED", "REVIEW_REJECTED", "CHANGES_REQUESTED", "RESUBMITTED_FOR_REVIEW", "REDLINE_APPLIED"]:
        return "Review"
    if event_type == "CAS_GENERATED":
        return "CAS"
    if event_type in ["DOA_STEP_APPROVED", "DOA_COMPLETED"]:
        return "DOA"
    if event_type == "CONTRACT_COMPLETED":
        return "Completion"
    return "General"


def _infer_audit_event(
    action: str,
    user_name: str,
    role: str = "System",
    department: Optional[str] = None,
    details: Optional[str] = None,
    notes: Optional[str] = None,
    event_type: Optional[str] = None,
    message: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    normalized_action = (action or "").strip()
    normalized_metadata = dict(metadata or {})

    if event_type:
        normalized_event_type = event_type
    elif normalized_action == "Created":
        normalized_event_type = AUDIT_EVENT_TYPES["CREATED"]
    elif normalized_action == "Approved":
        normalized_event_type = AUDIT_EVENT_TYPES["REVIEW_APPROVED"]
    elif normalized_action == "Workflow Approved":
        normalized_event_type = AUDIT_EVENT_TYPES["WORKFLOW_COMPLETED"]
    elif normalized_action == "CAS Generated":
        normalized_event_type = AUDIT_EVENT_TYPES["CAS_GENERATED"]
    elif normalized_action == "DOA Approved":
        normalized_event_type = AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"]
    elif normalized_action == "Completed":
        normalized_event_type = AUDIT_EVENT_TYPES["CONTRACT_COMPLETED"]
    elif normalized_action == "Rejected":
        normalized_event_type = AUDIT_EVENT_TYPES["REVIEW_REJECTED"]
    elif normalized_action == "Changes Requested":
        normalized_event_type = AUDIT_EVENT_TYPES["CHANGES_REQUESTED"]
    elif normalized_action == "Resubmitted for Review":
        normalized_event_type = AUDIT_EVENT_TYPES["RESUBMITTED_FOR_REVIEW"]
    elif normalized_action == "Routed" or normalized_action == "Moved to Next Stage" or normalized_action == "Submitted for Review":
        normalized_event_type = AUDIT_EVENT_TYPES["ROUTED"]
    elif normalized_action == "Redline Applied":
        normalized_event_type = AUDIT_EVENT_TYPES["REDLINE_APPLIED"]
    else:
        normalized_event_type = AUDIT_EVENT_TYPES["UPDATED"]

    if normalized_event_type == AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"]:
        normalized_metadata["step"] = normalized_metadata.get("step") or role or department or "Unknown"
    if normalized_event_type == AUDIT_EVENT_TYPES["REVIEW_APPROVED"]:
        normalized_metadata["department"] = normalized_metadata.get("department") or department or role

    if message:
        normalized_message = message
    elif normalized_event_type == AUDIT_EVENT_TYPES["CREATED"]:
        normalized_message = details or f"Contract created by {user_name}."
    elif normalized_event_type == AUDIT_EVENT_TYPES["REVIEW_APPROVED"]:
        review_department = department or role or "Review"
        normalized_message = f"{review_department} review approved."
    elif normalized_event_type == AUDIT_EVENT_TYPES["WORKFLOW_COMPLETED"]:
        normalized_message = "All required reviews completed."
    elif normalized_event_type == AUDIT_EVENT_TYPES["CAS_GENERATED"]:
        normalized_message = "CAS generated."
    elif normalized_event_type == AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"]:
        doa_step = normalized_metadata.get("step") or role or department or "DOA"
        normalized_message = f"{doa_step} DOA step approved."
    elif normalized_event_type == AUDIT_EVENT_TYPES["DOA_COMPLETED"]:
        normalized_message = "All DOA approval steps completed."
    elif normalized_event_type == AUDIT_EVENT_TYPES["CONTRACT_COMPLETED"]:
        normalized_message = "Contract completed."
    elif normalized_event_type == AUDIT_EVENT_TYPES["REVIEW_REJECTED"]:
        normalized_message = details or "Contract rejected."
    elif normalized_event_type == AUDIT_EVENT_TYPES["CHANGES_REQUESTED"]:
        normalized_message = details or "Changes requested during review."
    elif normalized_event_type == AUDIT_EVENT_TYPES["RESUBMITTED_FOR_REVIEW"]:
        normalized_message = details or "Contract resubmitted for review."
    elif normalized_event_type == AUDIT_EVENT_TYPES["ROUTED"]:
        normalized_message = details or "Workflow advanced."
    elif normalized_event_type == AUDIT_EVENT_TYPES["REDLINE_APPLIED"]:
        normalized_message = details or "Redline applied."
    else:
        normalized_message = details or normalized_action or "Contract activity updated."

    return {
        "action": normalized_action or normalized_event_type.replace("_", " ").title(),
        "eventType": normalized_event_type,
        "message": normalized_message,
        "actor": user_name,
        "phase": _audit_phase_for_event(normalized_event_type),
        "metadata": normalized_metadata,
        "department": department,
        "role": role,
        "details": details,
        "notes": notes,
    }


def _normalize_audit_log_document(log: Dict[str, Any]) -> Dict[str, Any]:
    normalized = dict(log)
    if normalized.get("_id") is not None:
        normalized["id"] = str(normalized["_id"])
        del normalized["_id"]

    normalized_payload = _infer_audit_event(
        action=normalized.get("action", ""),
        user_name=normalized.get("actor") or normalized.get("userName") or "System",
        role=normalized.get("role", "System"),
        department=normalized.get("department"),
        details=normalized.get("details"),
        notes=normalized.get("notes"),
        event_type=normalized.get("eventType"),
        message=normalized.get("message"),
        metadata=normalized.get("metadata") or {},
    )
    normalized.update(normalized_payload)
    normalized["userName"] = normalized.get("actor") or normalized.get("userName") or "System"
    normalized["timestamp"] = normalized.get("timestamp") or datetime.utcnow().isoformat()
    return normalized


def _curate_audit_logs(logs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normalized_logs = [_normalize_audit_log_document(log) for log in logs]
    normalized_logs.sort(key=lambda item: item.get("timestamp", ""))

    deduped: List[Dict[str, Any]] = []
    for log in normalized_logs:
        if deduped:
            prev = deduped[-1]
            if (
                prev.get("eventType") == log.get("eventType")
                and prev.get("actor") == log.get("actor")
                and prev.get("message") == log.get("message")
                and prev.get("department") == log.get("department")
                and prev.get("timestamp") == log.get("timestamp")
            ):
                continue
        deduped.append(log)

    if any(log.get("eventType") == AUDIT_EVENT_TYPES["CAS_GENERATED"] for log in deduped):
        deduped = [log for log in deduped if log.get("eventType") != AUDIT_EVENT_TYPES["WORKFLOW_COMPLETED"]]

    doa_steps = [log for log in deduped if log.get("eventType") == AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"]]
    if len(doa_steps) > 1 and any(log.get("eventType") == AUDIT_EVENT_TYPES["CONTRACT_COMPLETED"] for log in deduped):
        latest_step = doa_steps[-1]
        grouped_event = {
            "id": f"doa-summary-{latest_step.get('timestamp')}",
            "contractId": latest_step.get("contractId"),
            "userName": "System",
            "actor": "System",
            "role": "System",
            "action": "DOA Completed",
            "eventType": AUDIT_EVENT_TYPES["DOA_COMPLETED"],
            "message": "All DOA approval steps completed.",
            "phase": "DOA",
            "department": "DOA",
            "details": "All DOA approval steps completed.",
            "notes": None,
            "metadata": {
                "steps": [step.get("metadata", {}).get("step") or step.get("department") for step in doa_steps]
            },
            "timestamp": latest_step.get("timestamp"),
        }
        deduped = [log for log in deduped if log.get("eventType") != AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"]]
        deduped.append(grouped_event)

    deduped.sort(key=lambda item: item.get("timestamp", ""), reverse=True)
    return deduped


async def log_audit(
    contract_id: str,
    action: str,
    user_name: str,
    role: str = "System",
    department: Optional[str] = None,
    details: Optional[str] = None,
    notes: Optional[str] = None,
    event_type: Optional[str] = None,
    message: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None,
):
    try:
        normalized_payload = _infer_audit_event(
            action=action,
            user_name=user_name,
            role=role,
            department=department,
            details=details,
            notes=notes,
            event_type=event_type,
            message=message,
            metadata=metadata,
        )
        latest = await db.audit_logs.find_one(
            {"contractId": str(contract_id)},
            sort=[("timestamp", -1)]
        )
        if latest:
            latest_normalized = _normalize_audit_log_document(latest)
            if (
                latest_normalized.get("eventType") == normalized_payload["eventType"]
                and latest_normalized.get("actor") == normalized_payload["actor"]
                and latest_normalized.get("department") == normalized_payload["department"]
                and latest_normalized.get("message") == normalized_payload["message"]
                and latest_normalized.get("notes") == normalized_payload["notes"]
            ):
                return

        if latest and latest.get("action") == action and latest.get("userName") == user_name and latest.get("role") == role and latest.get("department") == department and latest.get("details") == details and latest.get("notes") == notes:
            return

        log = AuditLog(
            contractId=str(contract_id),
            userName=user_name,
            actor=normalized_payload["actor"],
            role=role,
            action=normalized_payload["action"],
            eventType=normalized_payload["eventType"],
            message=normalized_payload["message"],
            phase=normalized_payload["phase"],
            department=department,
            details=details,
            notes=notes,
            metadata=normalized_payload["metadata"],
        )
        await db.audit_logs.insert_one(log.model_dump(by_alias=True, exclude_none=True))
    except Exception as e:
        print(f"Error creating audit log: {e}")

app = FastAPI(title="CLM Admin Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DEMO_USERS = [
<<<<<<< Updated upstream
    {"name": "Admin User", "email": "admin@apeiro.com", "role": "Admin", "password": "Admin@2026", "status": "Active"},
    {"name": "Legal Counsel", "email": "legal@apeiro.com", "role": "Legal", "password": "Legal@2026", "status": "Active"},
    {"name": "Finance Controller", "email": "finance@apeiro.com", "role": "Finance", "password": "Finance@2026", "status": "Active"},
    {"name": "Compliance Officer", "email": "compliance@apeiro.com", "role": "Compliance", "password": "Comply@2026", "status": "Active"},
    {"name": "Procurement Lead", "email": "procurement@apeiro.com", "role": "Procurement", "password": "Procure@2026", "status": "Active"},
    {"name": "Sales Manager", "email": "sales@apeiro.com", "role": "Sales", "password": "Sales@2026", "status": "Active"},
    {"name": "Operations Manager", "email": "manager@apeiro.com", "role": "Manager", "password": "Manager@2026", "status": "Active"},
    {"name": "Chief Executive Officer", "email": "ceo@apeiro.com", "role": "CEO", "password": "CEO@2026", "status": "Active"},
=======
    {"name": ADMIN_NAME, "email": ADMIN_EMAIL, "role": "Admin", "password": "Admin@2026", "status": "Active"},
    {"name": "Legal Counsel", "email": f"legal@{WHITELISTED_DOMAIN}", "role": "Legal", "password": "Legal@2026", "status": "Active"},
    {"name": "Finance Controller", "email": f"finance@{WHITELISTED_DOMAIN}", "role": "Finance", "password": "Finance@2026", "status": "Active"},
    {"name": "Compliance Officer", "email": f"compliance@{WHITELISTED_DOMAIN}", "role": "Compliance", "password": "Comply@2026", "status": "Active"},
    {"name": "Procurement Lead", "email": f"procurement@{WHITELISTED_DOMAIN}", "role": "Procurement", "password": "Procure@2026", "status": "Active"},
    {"name": "Sales Manager", "email": f"sales@{WHITELISTED_DOMAIN}", "role": "Sales", "password": "Sales@2026", "status": "Active"},
    {"name": "Operations Manager", "email": f"manager@{WHITELISTED_DOMAIN}", "role": "Manager", "password": "Manager@2026", "status": "Active"},
    {"name": "Chief Executive Officer", "email": f"ceo@{WHITELISTED_DOMAIN}", "role": "CEO", "password": "CEO@2026", "status": "Active"},
    {"name": "Chief Legal Officer", "email": f"clo@{WHITELISTED_DOMAIN}", "role": "CLO", "password": "CLO@2026", "status": "Active"},
    {"name": USER_NAME, "email": USER_EMAIL, "role": "User", "password": "User@2026", "status": "Active"},
    {"name": SUPERADMIN_NAME, "email": SUPERADMIN_EMAIL, "role": "Superadmin", "password": "Super@2026", "status": "Active"},
>>>>>>> Stashed changes
]

async def check_expiring_contracts():
    """Runs daily. Finds contracts expiring in 30, 14, or 7 days and creates notifications."""
    from datetime import date, timedelta
    today = date.today()
    warning_windows = [30, 14, 7]  # days before expiry
    
    cursor = db.contracts.find({"expiry_date": {"$exists": True, "$ne": None, "$ne": ""}})
    contracts = await cursor.to_list(length=500)
    
    for contract in contracts:
        expiry_raw = contract.get("expiry_date", "")
        if not expiry_raw:
            continue
        try:
            # Handle potential datetime objects or ISO strings
            expiry_str = str(expiry_raw).split('T')[0].split(' ')[0]
            expiry_date = date.fromisoformat(expiry_str)
        except Exception:
            continue
        
        days_left = (expiry_date - today).days
        
        if days_left in warning_windows:
            contract_id = str(contract["_id"])
            title = contract.get("title", "Unknown Contract")
            
            # Avoid duplicate notifications for same contract+days_left combo
            existing = await db.notifications.find_one({
                "contract_id": contract_id,
                "lifecycle_alert": True,
                "days_left": days_left
            })
            if existing:
                continue
            
            message = f"⚠️ '{title}' expires in {days_left} days ({expiry_date.isoformat()}). Action required."
            
            await db.notifications.insert_one({
                "for_role": "Admin",
                "message": message,
                "contract_id": contract_id,
                "contract_title": title,
                "action": "Go to Dashboard",
                "lifecycle_alert": True,
                "days_left": days_left,
                "read": False,
                "createdAt": datetime.utcnow().isoformat()
            })
            
            # Also notify the contract owner role if available
            owner_role = contract.get("department", "")
            if owner_role:
                await db.notifications.insert_one({
                    "for_role": owner_role,
                    "message": message,
                    "contract_id": contract_id,
                    "contract_title": title,
                    "action": "Go to Dashboard",
                    "lifecycle_alert": True,
                    "days_left": days_left,
                    "read": False,
                    "createdAt": datetime.utcnow().isoformat()
                })
            
            # Send email to contract owner if SMTP configured
            owner_name = contract.get("contract_owner") or contract.get("submittedBy", "")
            if owner_name:
                user_doc = await db.users.find_one({"name": owner_name})
                if user_doc and user_doc.get("email"):
                    await send_email_notification(
                        to_email=user_doc["email"],
                        subject=f"[Apeiro CLM] Contract Expiry Alert: {title}",
                        body_html=f"<p>This is an automated reminder.</p><p><strong>{title}</strong> is expiring in <strong>{days_left} days</strong> on {expiry_date.isoformat()}.</p><p>Please log in to the CLM system to take action.</p>"
                    )
            
            print(f"Lifecycle alert created for '{title}' — {days_left} days left")

@app.on_event("startup")
async def startup_event():
    print("Backend startup: Checking user seeding...")
    try:
        count = await db.users.count_documents({})
        if count == 0:
            print(f"Seeding {len(DEMO_USERS)} demo users into 'clm_platform.users'...")
            await db.users.insert_many(DEMO_USERS)
            print("Seeding completed successfully.")
        else:
            print(f"Database already has {count} users. Skipping seeding.")
    except Exception as e:
        print(f"ERROR during user seeding: {e}")
        
    scheduler.add_job(check_expiring_contracts, CronTrigger(hour=8, minute=0), id="lifecycle_check", replace_existing=True)
    scheduler.start()
    print("Lifecycle scheduler started — runs daily at 8:00 AM")

@app.post("/api/admin/seed-users")
async def manual_seed_users():
    """Manual endpoint to force seed demo users if they are missing."""
    count = await db.users.count_documents({})
    if count == 0:
        await db.users.insert_many(DEMO_USERS)
        return {"message": f"Seeded {len(DEMO_USERS)} users successfully."}
    return {"message": f"Users already exist ({count} users). No action taken."}

@app.on_event("shutdown")
async def shutdown_event():
    scheduler.shutdown()

@app.get("/api/dashboard/stats")
async def get_dashboard_stats():
    total = await db.contracts.count_documents({})
    under_review = await db.contracts.count_documents({"stage": "Under Review"})
    pending_approval = await db.contracts.count_documents({"stage": "DOA Approval"})
    approved = await db.contracts.count_documents({"status": "Approved"})
    
    current_month = datetime.utcnow().strftime("%Y-%m")
    approved_this_month = await db.contracts.count_documents({
        "status": "Approved",
        "createdAt": {"$regex": f"^{current_month}"}
    })
    
    return {
        "totalContracts": total,
        "underReview": under_review,
        "pendingApproval": pending_approval,
        "approved": approved,
        "approvedThisMonth": approved_this_month
    }

<<<<<<< Updated upstream
@app.post("/api/upload")
async def upload_contracts(file: UploadFile = File(...)):
    if not file.filename.endswith('.csv'):
        raise HTTPException(status_code=400, detail="Only CSV files are supported")
    
    contents = await file.read()
    try:
        df = pd.read_csv(StringIO(contents.decode('utf-8')))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error parsing CSV: {str(e)}")
    
    inserted_count = 0
    for _, row in df.iterrows():
        company_name = str(row.get('company', '')).strip()
        if not company_name or company_name == 'nan':
            company_name = 'Unknown Company'
        
        title = str(row.get('title', f"{company_name} Agreement")).strip()
        if not title or title == 'nan':
            title = f"{company_name} Agreement"
        
        try:
            value = float(str(row.get('value', 0)).replace(',', '').replace('$', ''))
        except:
            value = 0.0
        
        department = str(row.get('department', 'Legal')).strip()
        if department == 'nan':
            department = 'Legal'
        
        stage = str(row.get('stage', 'Under Review')).strip()
        if stage == 'nan':
            stage = 'Under Review'
        
        status = str(row.get('status', 'Pending')).strip()
        if status == 'nan':
            status = 'Pending'
        
        submitted_by = str(row.get('submittedBy', 'Admin')).strip()
        if submitted_by == 'nan':
            submitted_by = 'Admin'
        
        risk_level = str(row.get('riskLevel', 'Medium')).strip()
        if risk_level == 'nan':
            risk_level = 'Medium'
        
        contract = Contract(
            title=title,
            company=company_name,
            value=value,
            department=department,
            stage="Under Review",
            status="Pending"
        )
        
        await db.contracts.insert_one(contract.model_dump(by_alias=True, exclude_none=True))
        inserted_count += 1
        
    return {"message": f"Successfully processed and created {inserted_count} contracts."}

=======
>>>>>>> Stashed changes
@app.get("/api/contracts")
async def get_contracts(stage: str = None, department: str = None):
    query = {}
    if stage:
        query["stage"] = stage
    if department:
        query["$or"] = [
            {"department": department},
            {"review_mode": "parallel", "review_stages": department},
            {"review_mode": "parallel", "required_reviewers": department},
            {"review_mode": "parallel", "workflow": department},
        ]
        
    cursor = db.contracts.find(query)
    contracts = await cursor.to_list(length=1000)
    
    for c in contracts:
        c["id"] = str(c["_id"])
        del c["_id"]
        normalized = _normalize_contract_document(c)
        c.clear()
        c.update(normalized)
        
    return contracts

def _matches_keywords(text: str, keywords: List[str]) -> bool:
    return any(keyword in text for keyword in keywords)


def _build_rule_based_routing(
    title: str,
    value: float,
    category: str,
    risk: str,
    clauses: list,
    business_unit: str = "",
) -> Dict[str, Any]:
    department_reasons: Dict[str, List[str]] = {dept: [] for dept in REVIEW_DEPARTMENTS}
    applied_rules: List[str] = []

    normalized_category = (category or "").strip().lower()
    normalized_risk = (risk or "").strip().lower()
    normalized_business_unit = (business_unit or "").strip().lower()
    full_text = f"{title} {category} {business_unit}".lower()

    clause_texts = []
    for clause in clauses or []:
        clause_text = " ".join([
            str(clause.get("title", "")),
            str(clause.get("type", "")),
            str(clause.get("content", "")),
            str(clause.get("text", "")),
            str(clause.get("label", "")),
        ]).lower()
        clause_texts.append(clause_text)
        full_text += f" {clause_text}"

    legal_keywords = ["termination", "terminate", "liability", "indemn", "damages", "dispute", "governing law", "intellectual", "warranty", "breach"]
    finance_keywords = ["payment", "invoice", "pricing", "price", "fee", "fees", "financial", "milestone", "tax", "discount", "commercial"]
    compliance_keywords = ["compliance", "regulatory", "gdpr", "privacy", "data protection", "audit", "security", "sox", "hipaa", "sanctions"]
    procurement_keywords = ["vendor", "supplier", "procurement", "purchase", "sourcing", "delivery", "sla", "service level", "goods", "material", "po "]

    if any(_matches_keywords(text, legal_keywords) for text in clause_texts):
        department_reasons["Legal"].append("Clause analysis found legal-risk provisions such as termination or liability.")
        applied_rules.append("clauses:legal")
    if any(_matches_keywords(text, finance_keywords) for text in clause_texts):
        department_reasons["Finance"].append("Clause analysis found commercial or payment-related terms.")
        applied_rules.append("clauses:finance")
    if any(_matches_keywords(text, compliance_keywords) for text in clause_texts):
        department_reasons["Compliance"].append("Clause analysis found compliance, privacy, audit, or regulatory language.")
        applied_rules.append("clauses:compliance")
    if any(_matches_keywords(text, procurement_keywords) for text in clause_texts):
        department_reasons["Procurement"].append("Clause analysis found vendor, supplier, delivery, or SLA obligations.")
        applied_rules.append("clauses:procurement")

    if _matches_keywords(normalized_category, ["vendor", "procurement", "purchase", "supplier", "hardware", "sourcing"]):
        department_reasons["Procurement"].append("Contract category indicates vendor or procurement involvement.")
        applied_rules.append("category:procurement")
    if _matches_keywords(normalized_category, ["services", "msa", "nda", "agreement", "license", "licensing", "legal"]):
        department_reasons["Legal"].append("Contract category requires legal review.")
        applied_rules.append("category:legal")

    if value >= 100000:
        department_reasons["Finance"].append("Contract value exceeds the finance review threshold.")
        applied_rules.append("value:finance-100k")
    elif value >= 50000:
        department_reasons["Finance"].append("Contract value meets the finance review threshold.")
        applied_rules.append("value:finance-50k")

    if normalized_risk in ["high", "critical"]:
        department_reasons["Compliance"].append("High-risk contract requires compliance review.")
        applied_rules.append("risk:compliance")
        department_reasons["Legal"].append("High-risk contract requires legal review.")
        applied_rules.append("risk:legal")
    elif normalized_risk == "medium":
        department_reasons["Legal"].append("Medium-risk contract requires legal review.")
        applied_rules.append("risk:legal-medium")

    if _matches_keywords(full_text, ["payment_terms", "payment terms", "invoice", "fee", "pricing"]):
        department_reasons["Finance"].append("Routing detected payment terms in extracted contract content.")
        applied_rules.append("content:payment-terms")
    if _matches_keywords(full_text, ["termination", "liability", "indemnity", "indemnification"]):
        department_reasons["Legal"].append("Routing detected termination or liability terms in extracted contract content.")
        applied_rules.append("content:legal-core")
    if _matches_keywords(full_text, ["privacy", "gdpr", "data protection", "compliance", "audit"]):
        department_reasons["Compliance"].append("Routing detected privacy or regulatory terms in extracted contract content.")
        applied_rules.append("content:compliance-core")
    if _matches_keywords(full_text, ["vendor", "supplier", "purchase", "procurement", "delivery", "sla"]):
        department_reasons["Procurement"].append("Routing detected supplier or procurement terms in extracted contract content.")
        applied_rules.append("content:procurement-core")

    if _matches_keywords(normalized_business_unit, ["procurement", "supply chain", "operations"]):
        department_reasons["Procurement"].append(f"Business unit '{business_unit}' suggests procurement involvement.")
        applied_rules.append("business_unit:procurement")
    if _matches_keywords(normalized_business_unit, ["finance", "accounts", "treasury"]):
        department_reasons["Finance"].append(f"Business unit '{business_unit}' suggests finance review.")
        applied_rules.append("business_unit:finance")

    if not any(department_reasons.values()):
        department_reasons["Legal"].append("Default legal review applied because no specific routing signal was detected.")
        applied_rules.append("default:legal")

    workflow = [dept for dept in REVIEW_DEPARTMENTS if department_reasons[dept]]
    workflow = _dedupe_workflow(workflow)
    routing_decisions = [{"department": dept, "reason": "; ".join(reasons)} for dept, reasons in department_reasons.items() if reasons]
    routing_reasons = [decision["reason"] for decision in routing_decisions]

    return {
        "workflow": workflow,
        "routing_decisions": routing_decisions,
        "routing_reasons": routing_reasons,
        "applied_rules": list(dict.fromkeys(applied_rules)),
    }


<<<<<<< Updated upstream
async def calculate_review_stages(title: str, value: float, category: str, risk: str, clauses: list) -> list[str]:
    routing = _build_rule_based_routing(title, value, category, risk, clauses)
    workflow = list(routing["workflow"])

    try:
        from services import ai_service
        import asyncio
        ai_route = await asyncio.wait_for(
            ai_service.route_review(title, value, category, "Global", risk),
            timeout=12.0
        )
        if ai_route and "parallelReviewers" in ai_route:
            primary = ai_route.get("primaryReviewer", {}).get("department")
            if primary and primary in REVIEW_DEPARTMENTS:
                workflow.insert(0, str(primary))
            for dept in ai_route["parallelReviewers"]:
                dept_str = str(dept)
                if dept_str in REVIEW_DEPARTMENTS:
                    workflow.append(dept_str)
    except Exception as e:
        print(f"AI Routing failed, continuing with rule-based: {e}")

    final_reviewers = _dedupe_workflow(workflow)
    if not final_reviewers:
        final_reviewers = ["Legal"]
    return final_reviewers


async def resolve_review_plan(
    title: str,
    value: float,
    category: str,
    risk: str,
    clauses: list,
    business_unit: str = "",
    requested_mode: Optional[str] = None,
) -> Dict[str, Any]:
    base_routing = _build_rule_based_routing(title, value, category, risk, clauses, business_unit)
    workflow = list(base_routing["workflow"])
    normalized_category = (category or "").strip().lower()
    normalized_risk = (risk or "").strip().lower()
    reasons: List[str] = list(base_routing["routing_reasons"])
    applied_rules: List[str] = list(base_routing["applied_rules"])
    routing_decisions: List[Dict[str, str]] = list(base_routing["routing_decisions"])

    review_mode = "sequential"
    settings_doc = await db.settings.find_one({"key": "review_routing_rules"})
    settings_value = settings_doc.get("value", {}) if settings_doc else {}

    try:
        category_parallel = {
            str(item).strip().lower()
            for item in settings_value.get("parallel_categories", [])
            if str(item).strip()
        }
        risk_parallel = {
            str(item).strip().lower()
            for item in settings_value.get("parallel_risks", [])
            if str(item).strip()
        }
        value_parallel_threshold = _parse_contract_value(settings_value.get("parallel_value_threshold"))

        if normalized_category in category_parallel:
            review_mode = "parallel"
            reasons.append("Parallel review enabled by category routing rule.")
            applied_rules.append("settings:parallel-category")
        if normalized_risk in risk_parallel:
            review_mode = "parallel"
            reasons.append("Parallel review enabled by risk routing rule.")
            applied_rules.append("settings:parallel-risk")
        if value_parallel_threshold and value >= value_parallel_threshold:
            review_mode = "parallel"
            reasons.append("Parallel review enabled by value threshold rule.")
            applied_rules.append("settings:parallel-value-threshold")
    except Exception as e:
        print(f"Review routing settings parsing failed: {e}")

    if requested_mode in ["sequential", "parallel"]:
        review_mode = requested_mode
        applied_rules.append("request:review_mode")
        reasons.append(f"Review mode explicitly set to {requested_mode}.")
    elif normalized_risk in ["high", "critical"] and len(workflow) > 1:
        review_mode = "parallel"
        reasons.append("Parallel review enabled automatically for multi-department high-risk review.")
        applied_rules.append("risk:auto-parallel")

    if not reasons:
        reasons.append("Default business rules applied to determine required reviewers.")
        applied_rules.append("default:routing")

    workflow = _dedupe_workflow(workflow)
    if not workflow:
        workflow = ["Legal"]
        routing_decisions = [{"department": "Legal", "reason": "Default legal review applied because no routing rule matched."}]
        reasons = ["Default legal review applied because no routing rule matched."]
        applied_rules.append("default:legal")

    return {
        "workflow": workflow,
        "review_mode": review_mode,
        "reasons": list(dict.fromkeys(reasons)),
        "applied_rules": list(dict.fromkeys(applied_rules)),
        "routing_decisions": routing_decisions,
=======
@app.get("/api/user/approvals")
async def get_user_approvals(current_user: dict = Depends(get_current_user)):
    user_email = current_user.get("email")
    user_role = current_user.get("role")
    
    # Only standard users or initiators typically track their own approvals here
    # Optimization: Filter by submittedBy in MongoDB
    query = {
        "submittedBy": user_email,
        "stage": {"$in": ["CAS", "CAS Generated", "DOA Approval"]}
>>>>>>> Stashed changes
    }


def _parse_contract_value(raw_value: Any) -> float:
    try:
        if raw_value is None:
            return 0.0
        if isinstance(raw_value, str):
            raw_value = re.sub(r"[^0-9.]", "", raw_value)
        return float(raw_value) if raw_value not in ["", None] else 0.0
    except Exception:
        return 0.0


def _first_present(data: Dict[str, Any], *keys: str, default: Any = None) -> Any:
    for key in keys:
        if key in data and data.get(key) is not None:
            return data.get(key)
    return default


def _normalize_string(value: Any, default: str = "") -> str:
    if value is None:
        return default
    return str(value).strip()


def _normalize_risk_classification(value: Any) -> str:
    normalized = _normalize_string(value, "Medium")
    allowed = {"low": "Low", "medium": "Medium", "high": "High", "critical": "Critical"}
    return allowed.get(normalized.lower(), "Medium")


def _normalize_contract_input(data: Dict[str, Any]) -> Dict[str, Any]:
    category = _first_present(data, "category", default="General") or "General"
    risk = _normalize_risk_classification(_first_present(data, "risk_classification", "riskLevel", default="Medium"))
    business_unit = _first_present(data, "business_unit", "businessUnit", default="")
    contract_owner = _first_present(data, "contract_owner", "contractOwner", default="")
    sales_opportunity_id = _first_present(
        data,
        "sales_opportunity_id",
        "salesOpportunityId",
        "opportunity_id",
        "opportunityId",
        default="",
    )

    return {
        "title": _normalize_string(_first_present(data, "title", default="Untitled Contract"), "Untitled Contract") or "Untitled Contract",
        "company": _normalize_string(_first_present(data, "company", "counterparty", default="Unknown"), "Unknown") or "Unknown",
        "value": _parse_contract_value(_first_present(data, "value", default=0)),
        "duration": _normalize_string(_first_present(data, "duration", default="")),
        "category": _normalize_string(category, "General") or "General",
        "risk_classification": risk,
        "business_unit": _normalize_string(business_unit),
        "contract_owner": _normalize_string(contract_owner),
        "sales_opportunity_id": _normalize_string(sales_opportunity_id),
        "submittedBy": _normalize_string(_first_present(data, "submittedBy", default="Admin"), "Admin") or "Admin",
        "created_by": _normalize_string(_first_present(data, "created_by", "submittedBy", default="Admin"), "Admin") or "Admin",
        "clauses": _first_present(data, "clauses", default=[]) or [],
        "draftText": _first_present(data, "draftText", "contract_content", default="") or "",
        "expiry_date": _normalize_string(_first_present(data, "expiry_date", "expiryDate", default=None)) or None,
        "counterparty_email": _normalize_string(_first_present(data, "counterparty_email", "counterpartyEmail", default=None)) or None,
        "signer_name": _normalize_string(_first_present(data, "signer_name", "signerName", default=None)) or None,
    }


def _build_send_stage(department: str) -> str:
    return "Under Review" if department == "Legal" else f"{department} Review"


REVIEW_DEPARTMENTS = ["Procurement", "Legal", "Finance", "Compliance"]
DEFAULT_REVIEW_DAYS = int(os.getenv("DEFAULT_REVIEW_DAYS", "3"))


def _due_at_iso(days: int = DEFAULT_REVIEW_DAYS) -> str:
    return (datetime.utcnow() + timedelta(days=days)).isoformat()


def _dedupe_workflow(stages: List[str]) -> List[str]:
    seen = set()
    ordered = []
    for stage in stages:
        if stage in REVIEW_DEPARTMENTS and stage not in seen:
            ordered.append(stage)
            seen.add(stage)
    return ordered


def _normalize_contract_document(contract: Dict[str, Any]) -> Dict[str, Any]:
    normalized = dict(contract)
    workflow = normalized.get("workflow") or normalized.get("review_stages") or normalized.get("required_reviewers") or []
    workflow = _dedupe_workflow(workflow)
    if not workflow:
        fallback_department = normalized.get("department")
        workflow = [fallback_department] if fallback_department in REVIEW_DEPARTMENTS else ["Legal"]

    current_stage_index = normalized.get("current_stage_index")
    if current_stage_index is None:
        department = normalized.get("department")
        try:
            current_stage_index = workflow.index(department)
        except ValueError:
            current_stage_index = 0

    normalized["workflow"] = workflow
    normalized["review_stages"] = workflow
    normalized["required_reviewers"] = workflow
    normalized["current_stage_index"] = max(0, min(int(current_stage_index), max(len(workflow) - 1, 0)))
    normalized["current_department"] = normalized.get("department")
    normalized["contract_name"] = normalized.get("contract_name") or normalized.get("title")
    normalized["created_by"] = normalized.get("created_by") or normalized.get("submittedBy") or "Admin"
    normalized["routing_reasons"] = normalized.get("routing_reasons") or []
    normalized["routing_rules_applied"] = normalized.get("routing_rules_applied") or []
    normalized["routing_decisions"] = normalized.get("routing_decisions") or []
    normalized["change_requests"] = normalized.get("change_requests") or []
    normalized["active_change_request_ids"] = normalized.get("active_change_request_ids") or []
    normalized["current_version"] = int(normalized.get("current_version") or 1)
    normalized["versions"] = normalized.get("versions") or []
    normalized["redline_history"] = normalized.get("redline_history") or []
    normalized["internal_notes"] = normalized.get("internal_notes") or normalized.get("internalNotes") or ""
    normalized["internalNotes"] = normalized["internal_notes"]
    normalized["routingNotes"] = normalized["internal_notes"]
    normalized["workflow_events"] = normalized.get("workflow_events") or normalized.get("workflowEvents") or []
    normalized["workflowEvents"] = normalized["workflow_events"]
    normalized["business_unit"] = normalized.get("business_unit") or normalized.get("businessUnit") or ""
    normalized["businessUnit"] = normalized["business_unit"]
    normalized["contract_owner"] = normalized.get("contract_owner") or normalized.get("contractOwner") or ""
    normalized["contractOwner"] = normalized["contract_owner"]
    normalized["risk_classification"] = normalized.get("risk_classification") or normalized.get("riskLevel") or "Medium"
    normalized["riskLevel"] = normalized["risk_classification"]
    normalized["category"] = normalized.get("category") or "General"
    normalized["duration"] = normalized.get("duration") or ""
    normalized["sales_opportunity_id"] = (
        normalized.get("sales_opportunity_id")
        or normalized.get("salesOpportunityId")
        or normalized.get("opportunity_id")
        or normalized.get("opportunityId")
        or ""
    )
    normalized["salesOpportunityId"] = normalized["sales_opportunity_id"]
    normalized["opportunity_id"] = normalized["sales_opportunity_id"]
    normalized["opportunityId"] = normalized["sales_opportunity_id"]

    due_at = normalized.get("dueAt")
    status = normalized.get("status", "Pending")
    if due_at and status in ["Pending", "Under Review"]:
        try:
            if datetime.fromisoformat(str(due_at)) < datetime.utcnow():
                normalized["status"] = "Overdue"
        except Exception:
            pass

    if normalized.get("status") == "Completed" and not normalized.get("completedAt"):
        normalized["completedAt"] = normalized.get("updatedAt") or normalized.get("createdAt")

    return normalized


def _contract_snapshot(contract: Dict[str, Any]) -> Dict[str, Any]:
    snapshot = dict(contract)
    snapshot.pop("_id", None)
    return snapshot


def _build_workflow_event(action: str, performed_by: str, department: str = "", notes: str = "", metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    return {
        "action": action,
        "performedBy": performed_by,
        "department": department,
        "notes": notes or "",
        "timestamp": datetime.utcnow().isoformat(),
        "metadata": metadata or {},
    }


async def _latest_contract_file_url(contract_id: str) -> Optional[str]:
    try:
        latest_document = await db.documents.find_one({"contractId": contract_id}, sort=[("version", -1), ("uploadedAt", -1)])
        if latest_document and latest_document.get("_id"):
            return f"/api/documents/{str(latest_document['_id'])}/download"
    except Exception as e:
        print(f"Failed to resolve latest contract file URL: {e}")
    return None


def _change_request_summary(change_request: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": str(change_request.get("_id") or change_request.get("id") or ""),
        "action": change_request.get("action", "request_changes"),
        "department": change_request.get("department", ""),
        "reviewer": change_request.get("reviewer") or change_request.get("requestedBy", ""),
        "comment": change_request.get("comment") or change_request.get("description", ""),
        "clause_reference": change_request.get("clause_reference") or change_request.get("clauseTitle") or change_request.get("fieldName") or "",
        "timestamp": change_request.get("createdAt") or datetime.utcnow().isoformat(),
        "status": change_request.get("status", "Open"),
        "linkedVersion": change_request.get("linkedVersion"),
    }


def _build_version_summary(
    version: int,
    file_url: Optional[str],
    uploaded_by: str,
    changes: str,
    change_request_ids: Optional[List[str]] = None,
    timestamp: Optional[str] = None,
) -> Dict[str, Any]:
    return {
        "version": version,
        "file_url": file_url,
        "uploaded_by": uploaded_by,
        "timestamp": timestamp or datetime.utcnow().isoformat(),
        "changes": changes,
        "change_request_ids": change_request_ids or [],
    }


async def _store_contract_version(
    contract_id: str,
    previous_contract: Dict[str, Any],
    next_snapshot: Dict[str, Any],
    created_by: str,
    change_type: str,
    summary: str,
    file_url: Optional[str] = None,
    change_request_ids: Optional[List[str]] = None,
) -> int:
    previous_version = int(previous_contract.get("current_version") or 1)
    next_version = previous_version + 1
    version_doc = ContractVersion(
        contractId=contract_id,
        version=next_version,
        changeType=change_type,
        summary=summary,
        file_url=file_url,
        changeRequestIds=change_request_ids or [],
        createdBy=created_by,
        snapshot=_contract_snapshot(next_snapshot),
        previousSnapshot=_contract_snapshot(previous_contract),
    )
    await db.contract_versions.insert_one(version_doc.model_dump(by_alias=True, exclude_none=True))
    return next_version


def _can_generate_cas(contract: Dict[str, Any]) -> bool:
    normalized = _normalize_contract_document(contract)
    workflow = normalized.get("workflow") or []
    if not workflow:
        return False
    if normalized.get("current_stage_index") != len(workflow) - 1:
        return False
    reviews = normalized.get("reviews") or {}
    return all(reviews.get(dept, {}).get("status") == "Approved" for dept in workflow)


def _normalize_cas_document(cas: Dict[str, Any]) -> Dict[str, Any]:
    normalized = dict(cas)
    if normalized.get("_id") is not None:
        normalized["id"] = str(normalized["_id"])
        del normalized["_id"]
    normalized["businessUnit"] = normalized.get("businessUnit") or normalized.get("business_unit") or ""
    normalized["business_unit"] = normalized["businessUnit"]
    normalized["agreementType"] = normalized.get("agreementType") or normalized.get("agreement_type") or ""
    normalized["agreement_type"] = normalized["agreementType"]
    normalized["keyNotes"] = normalized.get("keyNotes") or normalized.get("key_notes") or ""
    normalized["key_notes"] = normalized["keyNotes"]
    normalized["reviewDepartments"] = normalized.get("reviewDepartments") or normalized.get("review_departments") or []
    normalized["review_departments"] = normalized["reviewDepartments"]
    normalized["routingReasons"] = normalized.get("routingReasons") or normalized.get("routing_reasons") or []
    normalized["routing_reasons"] = normalized["routingReasons"]
    normalized["approvalChain"] = normalized.get("approvalChain") or []
    normalized["cost_center"] = normalized.get("cost_center") or ""
    normalized["project_name"] = normalized.get("project_name") or ""
    normalized["effective_date"] = normalized.get("effective_date") or ""
    normalized["execution_date"] = normalized.get("execution_date") or ""
    normalized["doaLevel"] = normalized.get("doaLevel") or normalized.get("doa_level") or ""
    normalized["doaRuleId"] = normalized.get("doaRuleId") or normalized.get("doa_rule_id") or ""
    normalized["doaRuleName"] = normalized.get("doaRuleName") or normalized.get("doa_rule_name") or ""
    normalized["doaRoutingSource"] = normalized.get("doaRoutingSource") or normalized.get("doa_routing_source") or ""
    return normalized


async def _hydrate_cas_document(cas: Dict[str, Any]) -> Dict[str, Any]:
    normalized = _normalize_cas_document(cas)
    source_contract: Dict[str, Any] = {}
    contract_id = normalized.get("contractId")
    if contract_id and ObjectId.is_valid(contract_id):
        contract_doc = await db.contracts.find_one({"_id": ObjectId(contract_id)})
        if contract_doc:
            source_contract = _normalize_contract_document(contract_doc)

    source_business_unit = (
        source_contract.get("businessUnit")
        or source_contract.get("business_unit")
        or source_contract.get("department")
        or source_contract.get("category")
        or "N/A"
    )
    source_agreement_type = (
        source_contract.get("category")
        or "N/A"
    )
    source_department = (
        normalized.get("department")
        or ", ".join(source_contract.get("workflow") or [])
        or source_contract.get("department")
        or source_business_unit
        or "N/A"
    )

    normalized["businessUnit"] = normalized.get("businessUnit") or source_business_unit
    normalized["business_unit"] = normalized["businessUnit"]
    normalized["agreementType"] = normalized.get("agreementType") or source_agreement_type
    normalized["agreement_type"] = normalized["agreementType"]
    normalized["department"] = source_department
    normalized["cost_center"] = normalized.get("cost_center") or source_contract.get("cost_center") or normalized["businessUnit"] or "N/A"
    normalized["project_name"] = normalized.get("project_name") or source_contract.get("project_name") or source_contract.get("title") or "N/A"
    normalized["effective_date"] = normalized.get("effective_date") or source_contract.get("effective_date") or source_contract.get("effectiveDate") or source_contract.get("createdAt") or ""
    normalized["execution_date"] = normalized.get("execution_date") or source_contract.get("execution_date") or source_contract.get("executionDate") or source_contract.get("updatedAt") or normalized.get("createdAt") or ""
    normalized["reviewDepartments"] = normalized.get("reviewDepartments") or source_contract.get("workflow") or ([source_contract.get("department")] if source_contract.get("department") else [])
    normalized["review_departments"] = normalized["reviewDepartments"]

    approval_chain = []
    for idx, step in enumerate(normalized.get("approvalChain") or []):
        step_role = step.get("role") or ""
        mapped_role = "Evaluator" if step_role == "Endorser" else step_role
        step_name = step.get("name") or ""
        if mapped_role == "Evaluator" and step_name in ["", "Dept Head"]:
            step_name = normalized["businessUnit"]
        elif mapped_role == "Reviewer" and step_name in ["", "Legal Counsel"]:
            step_name = normalized["department"]
        elif mapped_role == "Approver" and not step_name:
            step_name = normalized.get("doaApprover") or "Approver"
        approval_chain.append({
            **step,
            "role": mapped_role,
            "name": step_name or f"Step {idx + 1}",
        })
    normalized["approvalChain"] = approval_chain
    return normalized


async def _generate_dynamic_cas_notes(contract: Dict[str, Any]) -> str:
    direct_notes = (
        contract.get("aiSummary")
        or contract.get("executiveSummary")
        or contract.get("internal_notes")
        or contract.get("internalNotes")
        or contract.get("draftText")
        or ""
    )
    if direct_notes:
        return str(direct_notes).strip()

    reviews = contract.get("reviews") or {}
    try:
        return await ai_service.generate_cas_notes(
            title=contract.get("title", ""),
            company=contract.get("company", ""),
            value=contract.get("value", 0),
            contract_type=contract.get("category", ""),
            business_unit=contract.get("business_unit") or contract.get("businessUnit") or "",
            department=", ".join(contract.get("workflow") or [contract.get("department", "")]).strip(", "),
            risk_summary=contract.get("risk_classification", "Medium"),
            reviews=reviews,
            key_issues="; ".join(
                filter(
                    None,
                    [str((reviews.get(dept) or {}).get("comments", "")).strip() for dept in REVIEW_DEPARTMENTS]
                )
            ) or "All routed reviews completed.",
            initiator_notes=contract.get("internal_notes") or contract.get("submittedBy", ""),
        )
    except Exception as e:
        print(f"Failed to generate AI CAS notes: {e}")
        return f"CAS generated for {contract.get('title', 'contract')}. All required reviews completed and the document has been moved into DOA approval."


async def _build_cas_for_contract(contract: Dict[str, Any], thresholds: Dict[str, Any]) -> CAS:
    normalized_contract = _normalize_contract_document(contract)
    workflow = normalized_contract.get("workflow") or []
    routed_departments = workflow or ([normalized_contract.get("department")] if normalized_contract.get("department") else [])
    department_label = ", ".join(routed_departments)
    business_unit = (
        contract.get("businessUnit")
        or contract.get("business_unit")
        or contract.get("department")
        or contract.get("category")
        or "N/A"
    )
    agreement_type = contract.get("category") or "N/A"
    key_notes = await _generate_dynamic_cas_notes(normalized_contract)
    doa_config = await _resolve_doa_configuration(contract, thresholds)

    return services.generate_cas_document(
        contract_id=str(normalized_contract.get("id") or normalized_contract.get("_id")),
        contract_title=normalized_contract.get("title", ""),
        value=normalized_contract.get("value", 0),
        initiator=normalized_contract.get("submittedBy", "Admin"),
        thresholds=thresholds,
        agreement_type=agreement_type,
        business_unit=business_unit,
        department=department_label or normalized_contract.get("department") or business_unit,
        cost_center=normalized_contract.get("cost_center") or business_unit or "N/A",
        project_name=normalized_contract.get("project_name") or normalized_contract.get("title", "") or "N/A",
        effective_date=normalized_contract.get("effective_date") or normalized_contract.get("effectiveDate") or normalized_contract.get("createdAt") or "",
        execution_date=normalized_contract.get("execution_date") or normalized_contract.get("executionDate") or normalized_contract.get("updatedAt") or normalized_contract.get("createdAt") or "",
        key_notes=key_notes,
        review_departments=routed_departments,
        routing_reasons=normalized_contract.get("routing_reasons") or [],
        doa_approver=doa_config["approver"],
        approval_chain_roles=doa_config["approval_chain_roles"],
        doa_level=doa_config["level"],
        doa_rule=doa_config["rule"],
    )


def _build_simple_pdf_bytes(title: str, lines: List[str]) -> bytes:
    safe_lines = []
    for raw_line in lines:
        chunks = textwrap.wrap(str(raw_line), width=92) or [""]
        safe_lines.extend(chunks)

    content_lines = ["BT", "/F1 11 Tf", "50 790 Td", "14 TL"]
    title_line = title.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content_lines.append(f"({title_line}) Tj")
    for line in safe_lines:
        escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        content_lines.append("T*")
        content_lines.append(f"({escaped}) Tj")
    content_lines.append("ET")
    content = "\n".join(content_lines).encode("latin-1", errors="replace")

    objects = []
    objects.append(b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append(b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
    objects.append(b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n")
    objects.append(b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
    objects.append(f"5 0 obj << /Length {len(content)} >> stream\n".encode("latin-1") + content + b"\nendstream endobj\n")

    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf += obj
    xref_pos = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    pdf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode("latin-1")
    pdf += f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode("latin-1")
    return pdf


def _normalize_rule_value(value: Any) -> str:
    return str(value or "").strip().lower()


def _build_value_based_doa_chain(value: float) -> Dict[str, Any]:
    if value <= 10000:
        return {
            "approvers": ["Manager"],
            "final_approver": "Manager",
            "level": "L1",
        }
    if value <= 50000:
        return {
            "approvers": ["Director"],
            "final_approver": "Director",
            "level": "L2",
        }
    return {
        "approvers": ["VP"],
        "final_approver": "VP",
        "level": "L3",
    }


def _rule_matches_contract(rule: Dict[str, Any], contract: Dict[str, Any]) -> bool:
    conditions = rule.get("conditions") or {}
    contract_value = float(contract.get("value") or 0)
    opportunity_value = float(contract.get("opportunityValue") or contract.get("opportunity_value") or 0)
    contract_risk = _normalize_rule_value(contract.get("riskLevel") or contract.get("risk_classification"))
    contract_bu = _normalize_rule_value(contract.get("businessUnit") or contract.get("business_unit"))
    contract_category = _normalize_rule_value(contract.get("category"))

    min_value = conditions.get("min_value")
    max_value = conditions.get("max_value")
    min_opportunity_value = conditions.get("min_opportunity_value")
    max_opportunity_value = conditions.get("max_opportunity_value")

    if min_value is not None and contract_value < float(min_value):
        return False
    if max_value is not None and float(max_value) > 0 and contract_value > float(max_value):
        return False
    if min_opportunity_value is not None and opportunity_value < float(min_opportunity_value):
        return False
    if max_opportunity_value is not None and float(max_opportunity_value) > 0 and opportunity_value > float(max_opportunity_value):
        return False
    if conditions.get("risk") and _normalize_rule_value(conditions.get("risk")) != contract_risk:
        return False
    if conditions.get("businessUnit") and _normalize_rule_value(conditions.get("businessUnit")) != contract_bu:
        return False
    if conditions.get("category") and _normalize_rule_value(conditions.get("category")) != contract_category:
        return False
    return True


async def _resolve_doa_rule(contract: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    cursor = db.doa_rules.find({"is_active": True}).sort([("priority", 1), ("updated_at", -1)])
    rules = await cursor.to_list(length=500)
    for rule in rules:
        if _rule_matches_contract(rule, contract):
            rule["id"] = str(rule["_id"])
            del rule["_id"]
            return rule
    return None


async def _resolve_doa_configuration(contract: Dict[str, Any], thresholds: Dict[str, Any]) -> Dict[str, Any]:
    normalized_contract = _normalize_contract_document(contract)
    normalized_contract["opportunityValue"] = (
        contract.get("opportunityValue")
        or contract.get("opportunity_value")
        or contract.get("value")
        or 0
    )
    matched_rule = await _resolve_doa_rule(normalized_contract)
    if matched_rule:
        approvers = [str(item).strip() for item in matched_rule.get("approvers") or [] if str(item).strip()]
        final_approver = approvers[-1] if approvers else "Approver"
        level = f"R{matched_rule.get('priority', 1)}"
        return {
            "approver": final_approver,
            "approval_chain_roles": ["Initiator", *approvers],
            "level": level,
            "rule": matched_rule,
            "source": "rule-engine",
        }

    fallback = _build_value_based_doa_chain(float(normalized_contract.get("value") or 0))
    final_approver = fallback["final_approver"]
    return {
        "approver": final_approver,
        "approval_chain_roles": ["Initiator", "Evaluator", "Reviewer", final_approver],
        "level": fallback["level"],
        "rule": None,
        "source": "value-threshold",
    }

<<<<<<< Updated upstream
@app.post("/api/contracts/create")
async def create_single_contract(data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    try:
        normalized_input = _normalize_contract_input(data)
        title = normalized_input["title"]
        company = normalized_input["company"]
        value = normalized_input["value"]
        category = normalized_input["category"]
        risk = normalized_input["risk_classification"]
        clauses = normalized_input["clauses"]
        
        review_plan = await resolve_review_plan(
            title=title,
            value=value,
            category=category,
            risk=risk,
            clauses=clauses,
            business_unit=normalized_input["business_unit"],
            requested_mode=str(data.get("review_mode", "")).lower() or None,
=======

DOCUMENT_CATEGORIES = {
    "Self-Service (AI Instant)": {"NDA", "Simple Service Agreement", "Data Processing Agreement (DPA)", "Memorandum of Understanding (MOU)"},
    "Business Contracts (Legal Review)": {"Master Service Agreement (MSA)", "Statement of Work (SOW)", "Service Level Agreement (SLA)"},
    "Vendor / Procurement": {"Vendor Agreement", "Supplier Contract"},
    "Compliance / Legal": {"Data Privacy Agreement", "Compliance Agreement", "Regulatory Document"},
    "Legal Assistance": {"Contract Review Request", "Legal Advice", "Draft Custom Contract", "Clause Review"},
}


def get_document_category(template_type: str) -> str:
    if not template_type:
        return "Legal Assistance"
    for category, doc_types in DOCUMENT_CATEGORIES.items():
        if template_type in doc_types:
            return category
    return "Legal Assistance"


def infer_document_type(text: str, filename: str = "") -> str:
    haystack = f"{filename} {text}".lower()
    rules = [
        ("Master Service Agreement (MSA)", ["master service agreement", "msa"]),
        ("Statement of Work (SOW)", ["statement of work", "sow"]),
        ("Service Level Agreement (SLA)", ["service level agreement", "sla"]),
        ("Data Processing Agreement (DPA)", ["data processing agreement", "dpa"]),
        ("Data Privacy Agreement", ["data privacy", "privacy agreement", "privacy addendum"]),
        ("Compliance Agreement", ["compliance agreement", "compliance terms"]),
        ("Regulatory Document", ["regulatory", "statutory compliance", "governing authority"]),
        ("Vendor Agreement", ["vendor agreement", "vendor terms"]),
        ("Supplier Contract", ["supplier", "procurement contract"]),
        ("Simple Service Agreement", ["service agreement", "professional services"]),
        ("Memorandum of Understanding (MOU)", ["memorandum of understanding", "mou"]),
        ("NDA", ["non-disclosure", "non disclosure", "nda", "confidentiality agreement"]),
    ]
    for doc_type, keywords in rules:
        if any(k in haystack for k in keywords):
            return doc_type
    return "Contract Review Request"


def infer_clause_categories(text: str) -> List[str]:
    haystack = text.lower()
    category_keywords = {
        "Payment": ["payment", "invoice", "fee", "pricing", "charges", "net 30", "net 45"],
        "Legal": ["liability", "indemnity", "warranty", "termination", "dispute", "governing law"],
        "Compliance": ["compliance", "regulatory", "gdpr", "privacy", "security", "audit"],
    }
    detected = []
    for category, keywords in category_keywords.items():
        if any(k in haystack for k in keywords):
            detected.append(category)
    return detected or ["Legal"]


def extract_metadata_from_text(text: str) -> dict:
    compact_text = " ".join(text.split())
    party_candidates = []
    party_patterns = [
        r"between\s+(.{2,80}?)\s+and\s+(.{2,80}?)(?:\.|,|\s)",
        r"this\s+agreement\s+is\s+made\s+between\s+(.{2,80}?)\s+and\s+(.{2,80}?)(?:\.|,|\s)",
    ]
    for pattern in party_patterns:
        match = re.search(pattern, compact_text, flags=re.IGNORECASE)
        if match:
            party_candidates = [match.group(1).strip(), match.group(2).strip()]
            break

    date_matches = re.findall(
        r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|"
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+\d{1,2},?\s+\d{2,4})\b",
        compact_text,
        flags=re.IGNORECASE
    )

    return {
        "party_names": party_candidates,
        "dates": list(dict.fromkeys(date_matches))[:10]
    }


def extract_text_from_upload(filename: str, contents: bytes) -> str:
    ext = pathlib.Path(filename or "").suffix.lower()
    if ext == ".pdf":
        import PyPDF2
        reader = PyPDF2.PdfReader(BytesIO(contents))
        return "\n".join([(page.extract_text() or "") for page in reader.pages]).strip()
    if ext == ".docx":
        import docx as _docx
        doc_file = _docx.Document(BytesIO(contents))
        return "\n".join([p.text for p in doc_file.paragraphs]).strip()
    if ext == ".doc":
        # Best-effort extraction for legacy DOC in this demo setup.
        return contents.decode("latin-1", errors="ignore").strip()
    return contents.decode("utf-8", errors="ignore").strip()


def build_generated_text(template_type: str, company_name: str, description: str) -> str:
    return (
        f"{template_type}\n\n"
        f"This document is generated for Apeiro Digital and {company_name}.\n\n"
        f"Business Context:\n{description or 'Requested by business user for legal processing.'}\n\n"
        "Key obligations, scope, payment, legal protections, and compliance controls are included "
        "based on the selected template type."
    )


def build_clause_shell(clause_categories: List[str], description: str) -> List[dict]:
    clause_templates = {
        "Payment": "Fees, invoicing cadence, and payment timelines will follow mutually agreed commercial terms.",
        "Legal": "Liability, indemnity, term, termination, and governing law will be defined by legal standards.",
        "Compliance": "Data handling, audit rights, and applicable regulatory obligations must be met by both parties.",
    }
    clauses = []
    for idx, category in enumerate(clause_categories, start=1):
        clauses.append({
            "id": f"c_{idx}",
            "type": category,
            "text": clause_templates.get(category, "Standard legal clause to be reviewed by Legal."),
            "department": "Legal",
            "status": "Pending",
            "riskLevel": "Medium" if category in {"Legal", "Compliance"} else "Low"
        })
    if description:
        clauses.append({
            "id": f"c_{len(clauses) + 1}",
            "type": "Business Context",
            "text": description,
            "department": "Legal",
            "status": "Pending",
            "riskLevel": "Low"
        })
    return clauses


def _html_to_plain_text_for_docx(raw: str) -> str:
    text = raw or ""
    # Preserve paragraph and line boundaries before stripping tags.
    text = re.sub(r'</(p|div|h1|h2|h3|li|blockquote)>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<li[^>]*>', '• ', text, flags=re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_paragraph_texts_from_html(raw: str) -> List[str]:
    html = raw or ""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        nodes = soup.select(".docx p")
        if not nodes:
            nodes = soup.find_all(["p", "li", "h1", "h2", "h3", "h4", "h5", "h6", "div"])
        paragraph_texts = []
        for node in nodes:
            text = node.get_text("\n", strip=False).replace("\r", "")
            paragraph_texts.append(text.replace("\u00a0", " "))
        return paragraph_texts
    except Exception:
        # Best-effort fallback if parser fails
        plain = _html_to_plain_text_for_docx(html)
        return plain.split("\n")


def _set_wt_text(wt: ET.Element, text: str) -> None:
    xml_space_attr = "{http://www.w3.org/XML/1998/namespace}space"
    wt.text = text
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        wt.set(xml_space_attr, "preserve")
    elif xml_space_attr in wt.attrib:
        del wt.attrib[xml_space_attr]


def _patch_docx_ooxml_paragraphs(source_path: pathlib.Path, output_path: pathlib.Path, paragraph_texts: List[str]) -> bool:
    if not source_path.exists():
        return False
    if not paragraph_texts:
        return False

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = ns["w"]

    with tempfile.TemporaryDirectory(prefix="docx_patch_") as td:
        temp_dir = pathlib.Path(td)
        with zipfile.ZipFile(source_path, "r") as zf:
            zf.extractall(temp_dir)

        document_xml = temp_dir / "word" / "document.xml"
        if not document_xml.exists():
            return False

        tree = ET.parse(document_xml)
        root = tree.getroot()
        paragraphs = root.findall(".//w:body//w:p", ns)
        para_idx = 0

        for p in paragraphs:
            if para_idx >= len(paragraph_texts):
                break
            new_text = paragraph_texts[para_idx] or ""
            para_idx += 1

            text_nodes = p.findall(".//w:t", ns)
            if text_nodes:
                _set_wt_text(text_nodes[0], new_text)
                for wt in text_nodes[1:]:
                    _set_wt_text(wt, "")
                continue

            # Paragraph with no text runs: create one while preserving paragraph props.
            run = ET.Element(f"{{{w_ns}}}r")
            wt = ET.Element(f"{{{w_ns}}}t")
            _set_wt_text(wt, new_text)
            run.append(wt)
            p.append(run)

        # Append extra paragraphs if user added more block lines than existing OOXML paragraphs.
        body = root.find(".//w:body", ns)
        if body is not None and para_idx < len(paragraph_texts):
            sect_pr = body.find("w:sectPr", ns)
            insert_at = list(body).index(sect_pr) if sect_pr is not None else len(list(body))
            for extra in paragraph_texts[para_idx:]:
                p = ET.Element(f"{{{w_ns}}}p")
                run = ET.Element(f"{{{w_ns}}}r")
                wt = ET.Element(f"{{{w_ns}}}t")
                _set_wt_text(wt, extra or "")
                run.append(wt)
                p.append(run)
                body.insert(insert_at, p)
                insert_at += 1

        tree.write(document_xml, encoding="UTF-8", xml_declaration=True)

        with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as out_zip:
            for path in temp_dir.rglob("*"):
                if path.is_file():
                    out_zip.write(path, arcname=str(path.relative_to(temp_dir)))

    return True


async def _create_docx_version(
    contract_id: str,
    content: str,
    user_label: str,
    category_hint: str = "General",
    base_docx_path: str = None,
    rendered_html: Optional[str] = None,
    paragraph_texts: Optional[List[str]] = None
):
    from docx import Document as WordDocument
    from docx.shared import Pt

    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract_dir = UPLOAD_DIR / contract_id
    contract_dir.mkdir(parents=True, exist_ok=True)
    next_docx_version = int(contract.get("docx_version", 0)) + 1
    new_docx_path = contract_dir / f"edited_v{next_docx_version}.docx"

    plain_text = _html_to_plain_text_for_docx(content)
    lines = plain_text.split('\n')

    # Prefer updating an existing DOCX skeleton to preserve alignment/layout.
    source_path = base_docx_path or contract.get("latest_docx_path")
    docx = None
    if source_path:
        try:
            source_file = pathlib.Path(source_path)
            if source_file.exists():
                docx = WordDocument(str(source_file))
        except Exception:
            docx = None

    # Phase 2: strict OOXML paragraph patching to preserve legal numbering, indentation and styles.
    if source_path and paragraph_texts:
        try:
            source_file = pathlib.Path(source_path)
            if _patch_docx_ooxml_paragraphs(source_file, new_docx_path, paragraph_texts):
                await db.contracts.update_one(
                    {"_id": ObjectId(contract_id)},
                    {"$set": {
                        "latest_docx_path": str(new_docx_path),
                        "docx_version": next_docx_version,
                    }}
                )

                last_doc = await db.documents.find_one({"contractId": contract_id}, sort=[("version", -1)])
                version = (last_doc["version"] + 1) if last_doc else 1
                new_doc_record = Document(
                    contractId=contract_id,
                    fileName=new_docx_path.name,
                    fileType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    fileSize=new_docx_path.stat().st_size,
                    storagePath=str(new_docx_path),
                    version=version,
                    uploadedBy=user_label,
                    category=last_doc["category"] if last_doc else category_hint,
                    renderedHtml=rendered_html,
                    renderMode="html" if rendered_html else "docx-preview"
                )
                inserted = await db.documents.insert_one(new_doc_record.model_dump(by_alias=True, exclude_none=True))
                return {
                    "document_id": str(inserted.inserted_id),
                    "path": str(new_docx_path),
                    "version": next_docx_version,
                }
        except Exception as exc:
            logger.warning(f"OOXML patch path failed for {contract_id}: {exc}. Falling back to python-docx update.")

    if docx is None:
        docx = WordDocument()
        normal_style = docx.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)

    editable_paragraphs = list(docx.paragraphs)
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                editable_paragraphs.extend(cell.paragraphs)

    if editable_paragraphs:
        # Update paragraph text while keeping each paragraph's style/alignment.
        for idx, para in enumerate(editable_paragraphs):
            if idx >= len(lines):
                # Preserve untouched template paragraphs to avoid layout collapse.
                continue
            new_line = lines[idx]
            if para.runs:
                para.runs[0].text = new_line
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_line)
        # If content has more lines than available paragraphs, append new paragraphs.
        if len(lines) > len(editable_paragraphs):
            for extra in lines[len(editable_paragraphs):]:
                docx.add_paragraph(extra)
    else:
        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                docx.add_paragraph("")
                continue
            if line.startswith('### '):
                p = docx.add_paragraph(line[4:].strip())
                p.style = docx.styles['Heading 3']
            elif line.startswith('## '):
                p = docx.add_paragraph(line[3:].strip())
                p.style = docx.styles['Heading 2']
            elif line.startswith('# '):
                p = docx.add_paragraph(line[2:].strip())
                p.style = docx.styles['Heading 1']
            elif line.isupper() and len(line) < 120:
                p = docx.add_paragraph(line)
                if p.runs:
                    p.runs[0].bold = True
            else:
                docx.add_paragraph(line)

    docx.save(str(new_docx_path))

    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "latest_docx_path": str(new_docx_path),
            "docx_version": next_docx_version,
        }}
    )

    last_doc = await db.documents.find_one({"contractId": contract_id}, sort=[("version", -1)])
    version = (last_doc["version"] + 1) if last_doc else 1
    new_doc_record = Document(
        contractId=contract_id,
        fileName=new_docx_path.name,
        fileType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        fileSize=new_docx_path.stat().st_size,
        storagePath=str(new_docx_path),
        version=version,
        uploadedBy=user_label,
        category=last_doc["category"] if last_doc else category_hint,
        renderedHtml=rendered_html,
        renderMode="html" if rendered_html else "docx-preview"
    )
    inserted = await db.documents.insert_one(new_doc_record.model_dump(by_alias=True, exclude_none=True))
    return {
        "document_id": str(inserted.inserted_id),
        "path": str(new_docx_path),
        "version": next_docx_version,
    }

@app.post("/api/self-service/request")
async def create_self_service_request(data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    user_email = current_user.get("email", data.get("submittedBy"))
    template_type = data.get("template_type", "NDA")
    document_category = data.get("document_category") or get_document_category(template_type)
    request_mode = data.get("request_mode", "generated")
    is_self_service_category = document_category == "Self-Service (AI Instant)"
    
    contract = Contract(
        title=data.get("title", f"{data.get('company', 'Unknown')} - {template_type}"),
        company=data.get("company", "Unknown"),
        department=data.get("department", "Sales"),
        value=float(data.get("value", 0)) if data.get("value") else 0.0,
        stage="Under Review",
        status="Pending",
        submittedBy=user_email,
        required_reviewers=["Legal"]
    )
    
    contract_dump = contract.model_dump(by_alias=True, exclude_none=True)
    contract_dump["source"] = "self-service"
    contract_dump["template_type"] = template_type
    contract_dump["document_category"] = document_category
    contract_dump["request_mode"] = request_mode
    contract_dump["route"] = "ai-instant" if is_self_service_category else "legal-counsel"
    
    company_name = data.get("company", "Unknown")
    description = data.get("description", "")
    contract_dump["description"] = description

    clause_categories = infer_clause_categories(f"{template_type} {description}")
    contract_dump["extractedText"] = build_generated_text(template_type, company_name, description)
    contract_dump["clauses"] = build_clause_shell(clause_categories, description)
    contract_dump["ai_classification_result"] = {
        "document_type": template_type,
        "document_category": document_category,
        "clause_categories": clause_categories,
        "routing_decision": "AI Template Generation" if is_self_service_category else "Legal Counsel Review",
        "confidence": "high"
    }
    
    # Store expiry date as empty string if not provided
    contract_dump["expiry_date"] = data.get("expiry_date", "")
    
    result = await db.contracts.insert_one(contract_dump)
    contract_id = str(result.inserted_id)

    if is_self_service_category and template_type == "NDA":
        # Generate NDA DOCX from official template and store as primary editable document.
        nda_docx_path = generate_nda_docx(contract_id, company_name, description, contract_dump["expiry_date"])
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": {
                "latest_docx_path": nda_docx_path,
                "docx_version": 1
            }}
        )
        nda_docx_file = pathlib.Path(nda_docx_path)
        new_doc_record = Document(
            contractId=contract_id,
            fileName=nda_docx_file.name,
            fileType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            fileSize=nda_docx_file.stat().st_size if nda_docx_file.exists() else 0,
            storagePath=nda_docx_path,
            version=1,
            uploadedBy=user_email,
            category="Original"
        )
        await db.documents.insert_one(new_doc_record.model_dump(by_alias=True, exclude_none=True))

        # Keep NDA PDF generation as fallback preview.
        pdf_path = generate_nda_pdf(contract_id, company_name, description, contract_dump["expiry_date"])
        pdf_file = PdfFile(
            contractId=contract_id,
            template_type=contract_dump["template_type"],
            file_path=pdf_path,
            generated_by=user_email,
            version=1
        )
        await db.pdf_files.insert_one(pdf_file.model_dump(by_alias=True, exclude_none=True))
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": {
                "pdf_path": pdf_path,
                "pdf_generated": True,
                "pdf_version": 1
            }}
>>>>>>> Stashed changes
        )
        final_reviewers = review_plan["workflow"]
        review_mode = review_plan["review_mode"]
            
        status_val = data.get("status") or "Under Review"
        
        if status_val == "Draft":
            stage_val = "Draft"
        else:
            if review_mode == "sequential":
                stage_val = _build_send_stage(final_reviewers[0])
            else:
                stage_val = "Under Review"

        contract = Contract(
            title=title,
            contract_name=title,
            company=company,
            value=value,
            department=data.get("department") or normalized_input["business_unit"] or (final_reviewers[0] if final_reviewers else "Legal"),
            stage=stage_val,
            status=status_val,
            submittedBy=normalized_input["submittedBy"],
            created_by=normalized_input["created_by"],
            updatedAt=datetime.utcnow().isoformat(),
            dueAt=_due_at_iso() if status_val != "Draft" else None,
            counterparty_email=normalized_input["counterparty_email"],
            signer_name=normalized_input["signer_name"],
            duration=normalized_input["duration"],
            category=category,
            risk_classification=risk,
            business_unit=normalized_input["business_unit"],
            contract_owner=normalized_input["contract_owner"] or normalized_input["submittedBy"],
            sales_opportunity_id=normalized_input["sales_opportunity_id"] or None,
            opportunity_id=normalized_input["sales_opportunity_id"] or None,
            expiry_date=normalized_input["expiry_date"],
            clauses=clauses,
            draftText=normalized_input["draftText"],
            workflow=final_reviewers,
            current_stage_index=0,
            required_reviewers=final_reviewers,
            review_stages=final_reviewers,
            routing_reasons=review_plan["reasons"],
            routing_rules_applied=review_plan["applied_rules"],
            routing_decisions=review_plan["routing_decisions"],
            review_mode=review_mode
        )
        
        # Initialize first reviewer status to Pending if not draft
        if stage_val != "Draft":
            if review_mode == "sequential":
                first_reviewer = final_reviewers[0]
                getattr(contract.reviews, first_reviewer).status = "Pending"
                setattr(contract, f"{first_reviewer.lower()}_status", "Pending")
            else:
                for reviewer in final_reviewers:
                    getattr(contract.reviews, reviewer).status = "Pending"
                    setattr(contract, f"{reviewer.lower()}_status", "Pending")
        
        result = await asyncio.wait_for(
            db.contracts.insert_one(contract.model_dump(by_alias=True, exclude_none=True)),
            timeout=5.0
        )
        
        await log_audit(str(result.inserted_id), "Created", user_name, role="System", details=f"Contract '{title}' created by {user_name}")
        
        if stage_val != "Draft":
            await log_audit(str(result.inserted_id), "Routed", "System", role="System", details=f"Workflow initialized. Stages: {', '.join(final_reviewers)}")
            for reviewer in final_reviewers:
                await db.notifications.insert_one({
                    "for_role": reviewer,
                    "message": f"New contract '{title}' requires your review.",
                    "contract_id": str(result.inserted_id),
                    "contract_title": title,
                    "action": "Review Required",
                    "read": False,
                    "createdAt": datetime.utcnow().isoformat()
                })
                
        return {
            "message": "Contract created successfully",
            "id": str(result.inserted_id),
            "reviewers": final_reviewers,
            "mode": review_mode
        }
    except asyncio.TimeoutError:
        return JSONResponse(status_code=504, content={"message": "Request timed out. Please try again."})
    except Exception as e:
        print(f"Contract creation failed: {e}")
        return JSONResponse(status_code=400, content={"message": f"Failed to create contract: {str(e)}"})

@app.put("/api/contracts/{contract_id}")
async def update_contract(contract_id: str, data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    try:
        # Validate ObjectId
        try:
            from bson import ObjectId
            if not ObjectId.is_valid(contract_id):
                return JSONResponse(status_code=400, content={"message": "Invalid contract ID format"})
        except ImportError:
            pass

        existing = await db.contracts.find_one({"_id": PyObjectId(contract_id)})
        if not existing:
            return JSONResponse(status_code=404, content={"message": "Contract not found"})
        
        # Prepare update data
        update_fields = {}
        normalized_input = _normalize_contract_input(data)
        relevant_keys = [
            "title", "company", "value", "department", "status", "stage", 
            "duration", "category", "risk_classification", "business_unit",
            "contract_owner", "sales_opportunity_id", "opportunity_id", "expiry_date", "clauses", "draftText"
        ]

        alias_updates = {
            "title": normalized_input["title"],
            "company": normalized_input["company"],
            "value": normalized_input["value"],
            "duration": normalized_input["duration"],
            "category": normalized_input["category"],
            "risk_classification": normalized_input["risk_classification"],
            "business_unit": normalized_input["business_unit"],
            "contract_owner": normalized_input["contract_owner"],
            "sales_opportunity_id": normalized_input["sales_opportunity_id"],
            "opportunity_id": normalized_input["sales_opportunity_id"],
            "expiry_date": normalized_input["expiry_date"],
            "clauses": normalized_input["clauses"],
            "draftText": normalized_input["draftText"],
        }

        for key in relevant_keys:
            if key in data:
                update_fields[key] = alias_updates.get(key, data[key])

        if any(key in data for key in ["businessUnit", "business_unit"]):
            update_fields["business_unit"] = normalized_input["business_unit"]
        if any(key in data for key in ["contractOwner", "contract_owner"]):
            update_fields["contract_owner"] = normalized_input["contract_owner"]
        if any(key in data for key in ["riskLevel", "risk_classification"]):
            update_fields["risk_classification"] = normalized_input["risk_classification"]
        if any(key in data for key in ["salesOpportunityId", "sales_opportunity_id", "opportunityId", "opportunity_id"]):
            update_fields["sales_opportunity_id"] = normalized_input["sales_opportunity_id"]
            update_fields["opportunity_id"] = normalized_input["sales_opportunity_id"]
        if "expiryDate" in data and "expiry_date" not in update_fields:
            update_fields["expiry_date"] = normalized_input["expiry_date"]
        
        # SAFEGUARD: Workflow reset on move to Draft
        if data.get("status") == "Draft" and existing.get("status") != "Draft":
            update_fields["stage"] = "Draft"
            update_fields["review_stages"] = []
            update_fields["reviews"] = {}
            for r in ["legal", "finance", "compliance", "procurement"]:
                update_fields[f"{r}_status"] = "Not Started"
                
        # Initialize workflow if submitting for Review
        elif data.get("status") == "Under Review" and existing.get("status") != "Under Review":
            requested_mode = data.get("review_mode") or existing.get("review_mode", "sequential")
            
            # Recalculate review stages dynamically
            c_title = data.get("title") or existing.get("title", "Untitled")
            c_val = normalized_input["value"] if "value" in data else existing.get("value", 0.0)
                    
            c_cat = normalized_input["category"] if "category" in data else existing.get("category", "General")
            c_risk = normalized_input["risk_classification"] if any(k in data for k in ["risk_classification", "riskLevel"]) else existing.get("risk_classification", "Medium")
            c_clauses = normalized_input["clauses"] if "clauses" in data else existing.get("clauses", [])
            
            review_plan = await resolve_review_plan(
                title=c_title,
                value=c_val,
                category=c_cat,
                risk=c_risk,
                clauses=c_clauses,
                business_unit=data.get("business_unit") or data.get("businessUnit") or existing.get("business_unit", ""),
                requested_mode=requested_mode,
            )
            review_stages = review_plan["workflow"]
            review_mode = review_plan["review_mode"]
            update_fields["workflow"] = review_stages
            update_fields["review_stages"] = review_stages
            update_fields["required_reviewers"] = review_stages
            update_fields["review_mode"] = review_mode
            update_fields["routing_reasons"] = review_plan["reasons"]
            update_fields["routing_rules_applied"] = review_plan["applied_rules"]
            update_fields["routing_decisions"] = review_plan["routing_decisions"]
            first_reviewer = review_stages[0] if review_stages else "Legal"
            update_fields["current_stage_index"] = 0
            update_fields["department"] = first_reviewer
            update_fields["updatedAt"] = datetime.utcnow().isoformat()
            update_fields["dueAt"] = _due_at_iso()
            update_fields["rejection_reason"] = None
            update_fields["rejected_by"] = None
            update_fields["rejectedAt"] = None
            
            update_fields["status"] = "Under Review"
            if review_mode == "sequential":
                update_fields["stage"] = _build_send_stage(first_reviewer)
                for r in review_stages:
                    update_fields[f"reviews.{r}.status"] = "Pending" if r == first_reviewer else "Not Started"
                    update_fields[f"{r.lower()}_status"] = "Pending" if r == first_reviewer else "Not Started"
            else:
                update_fields["stage"] = "Under Review"
                for r in review_stages:
                    update_fields[f"reviews.{r}.status"] = "Pending"
                    update_fields[f"{r.lower()}_status"] = "Pending"
                    
            # Notification logic handled in audit section
            data["_new_review_stages"] = review_stages
            data["_first_reviewer"] = first_reviewer
            data["_c_title"] = c_title
            data["_review_mode"] = review_mode

        should_store_version = any(key in update_fields for key in ["clauses", "draftText", "redline_history"])
        next_snapshot = None
        if should_store_version:
            next_snapshot = dict(existing)
            for key, value in update_fields.items():
                if "." not in key:
                    next_snapshot[key] = value
            update_fields["current_version"] = int(existing.get("current_version") or 1) + 1
            next_snapshot["current_version"] = update_fields["current_version"]
            existing_versions = existing.get("versions") or []
            update_fields["versions"] = [
                *existing_versions,
                _build_version_summary(
                    version=update_fields["current_version"],
                    file_url=await _latest_contract_file_url(contract_id),
                    uploaded_by=user_name,
                    changes="Contract content updated with a new version snapshot.",
                ),
            ]
            next_snapshot["versions"] = update_fields["versions"]

        import asyncio
        if update_fields:
            await asyncio.wait_for(
                db.contracts.update_one(
                    {"_id": PyObjectId(contract_id)},
                    {"$set": update_fields}
                ),
                timeout=5.0
            )

        if should_store_version and next_snapshot is not None:
            await _store_contract_version(
                contract_id=contract_id,
                previous_contract=existing,
                next_snapshot=next_snapshot,
                created_by=user_name,
                change_type="redline" if "clauses" in update_fields else "content-update",
                summary="Contract content updated with a new version snapshot.",
            )
        
        # Audit log for submission/update
        if data.get("status") == "Under Review" and existing.get("status") != "Under Review":
            await log_audit(contract_id, "Submitted for Review", user_name, role="System")
            
            r_stages = data.get("_new_review_stages", [])
            await log_audit(contract_id, "Routed", "System", role="System", details=f"Workflow restarted. Stages: {', '.join(r_stages)}")
            
            c_title = data.get("_c_title", "Untitled")
            r_mode = data.get("_review_mode", "sequential")
            
            from datetime import datetime
            if r_mode == "sequential":
                first_req = data.get("_first_reviewer")
                if first_req:
                    await db.notifications.insert_one({
                        "for_role": first_req,
                        "message": f"Contract '{c_title}' has been submitted for review. You are the first reviewer.",
                        "contract_id": contract_id,
                        "contract_title": c_title,
                        "action": f"Go to {first_req} Review",
                        "department": "System",
                        "read": False,
                        "createdAt": datetime.utcnow().isoformat()
                    })
            else:
                r_stages = data.get("_new_review_stages", [])
                for r in r_stages:
                    await db.notifications.insert_one({
                        "for_role": r,
                        "message": f"Contract '{c_title}' has been submitted for parallel review.",
                        "contract_id": contract_id,
                        "contract_title": c_title,
                        "action": f"Go to {r} Review",
                        "department": "System",
                        "read": False,
                        "createdAt": datetime.utcnow().isoformat()
                    })
        elif data.get("status") == "Draft" and existing.get("status") != "Draft":
            await log_audit(contract_id, "Workflow Reset to Draft", user_name, role="System")
        else:
            await log_audit(contract_id, "Contract Updated", user_name, role="System")

        return {"message": "Contract updated successfully", "id": contract_id}
    except asyncio.TimeoutError:
        return JSONResponse(status_code=504, content={"message": "Update timed out. Please try again."})
    except Exception as e:
        print(f"Contract update failed: {e}")
        return JSONResponse(status_code=400, content={"message": f"Failed to update contract: {str(e)}"})    


@app.post("/api/contracts/send")
async def send_contract(data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    print("Incoming request:", data)

    try:
        contract_payload = data.get("contract_data") if isinstance(data.get("contract_data"), dict) else data
        normalized_input = _normalize_contract_input(contract_payload)
        requested_department = (data.get("department") or "").strip()
        routing_notes = _normalize_string(data.get("notes"))
        if not requested_department:
            raise HTTPException(status_code=400, detail="Missing department")

        contract_id = data.get("contract_id")
        value = _parse_contract_value(data.get("value", normalized_input["value"]))
        submitted_by = normalized_input["submittedBy"] or user_name or "Admin"
        title = normalized_input["title"] or "New Contract"
        company = normalized_input["company"] or "Unknown"
        clauses = normalized_input["clauses"]
        category = normalized_input["category"]
        risk_level = normalized_input["risk_classification"]

        review_plan = await resolve_review_plan(
            title=title,
            value=value,
            category=category,
            risk=risk_level,
            clauses=clauses,
            business_unit=normalized_input["business_unit"],
            requested_mode=str(contract_payload.get("review_mode", "")).lower() or None,
        )
        workflow = review_plan["workflow"]
        routing_decisions = list(review_plan["routing_decisions"])
        if requested_department in workflow:
            workflow = [requested_department] + [dept for dept in workflow if dept != requested_department]
        else:
            workflow = [requested_department] + workflow
            routing_decisions.insert(0, {
                "department": requested_department,
                "reason": "Manually selected by Sales override."
            })
        workflow = _dedupe_workflow(workflow)

        department = workflow[0]
        stage = _build_send_stage(department)
        review_mode = review_plan["review_mode"]
        if review_mode == "parallel":
            stage = "Under Review"

        review_update = {
            "department": department,
            "value": value,
            "status": "Under Review",
            "stage": stage,
            "contract_name": title,
            "submittedBy": submitted_by,
            "created_by": normalized_input["created_by"] or submitted_by,
            "updatedAt": datetime.utcnow().isoformat(),
            "dueAt": _due_at_iso(),
            "completedAt": None,
            "rejectedAt": None,
            "rejection_reason": None,
            "rejected_by": None,
            "duration": normalized_input["duration"],
            "category": category,
            "risk_classification": risk_level,
            "business_unit": normalized_input["business_unit"],
            "contract_owner": normalized_input["contract_owner"] or submitted_by,
            "sales_opportunity_id": normalized_input["sales_opportunity_id"],
            "opportunity_id": normalized_input["sales_opportunity_id"],
            "expiry_date": normalized_input["expiry_date"],
            "internal_notes": routing_notes,
            "workflow": workflow,
            "required_reviewers": workflow,
            "review_stages": workflow,
            "review_mode": review_mode,
            "routing_reasons": review_plan["reasons"],
            "routing_rules_applied": review_plan["applied_rules"],
            "routing_decisions": routing_decisions,
            "current_stage_index": 0,
            f"reviews.{department}.status": "Pending",
            f"{department.lower()}_status": "Pending",
        }
        for dept_name in REVIEW_DEPARTMENTS:
            if dept_name not in workflow:
                review_update[f"{dept_name.lower()}_status"] = "Not Started"
                review_update[f"reviews.{dept_name}.status"] = "Not Started"
            elif review_mode == "parallel":
                review_update[f"{dept_name.lower()}_status"] = "Pending"
                review_update[f"reviews.{dept_name}.status"] = "Pending"
            elif dept_name != department:
                review_update[f"{dept_name.lower()}_status"] = "Not Started"
                review_update[f"reviews.{dept_name}.status"] = "Not Started"

        if contract_id:
            if not ObjectId.is_valid(contract_id):
                raise HTTPException(status_code=400, detail="Invalid contract_id")

            existing = await db.contracts.find_one({"_id": ObjectId(contract_id)})
            if not existing:
                raise HTTPException(status_code=404, detail="Contract not found")

            if data.get("notes") is not None:
                review_update["draftText"] = data.get("notes")
            if routing_notes:
                existing_events = existing.get("workflow_events") or []
                review_update["workflow_events"] = [
                    *existing_events,
                    _build_workflow_event(
                        action=f"sent_to_{department.lower()}",
                        performed_by=user_name,
                        department=department,
                        notes=routing_notes,
                        metadata={"review_mode": review_mode, "workflow": workflow},
                    )
                ]

            await db.contracts.update_one(
                {"_id": ObjectId(contract_id)},
                {"$set": review_update}
            )

            await log_audit(
                contract_id,
                "Sent to Department",
                user_name,
                role="System",
                department=department,
                details=f"Contract routed to {department}",
                notes=routing_notes or None
            )
        else:
            contract = Contract(
                title=title,
                contract_name=title,
                company=company,
                value=value,
                department=department,
                stage=stage,
                status="Under Review",
                submittedBy=submitted_by,
                created_by=normalized_input["created_by"] or submitted_by,
                updatedAt=datetime.utcnow().isoformat(),
                dueAt=_due_at_iso(),
                counterparty_email=normalized_input["counterparty_email"],
                signer_name=normalized_input["signer_name"],
                duration=normalized_input["duration"],
                category=category,
                risk_classification=risk_level,
                business_unit=normalized_input["business_unit"],
                contract_owner=normalized_input["contract_owner"] or submitted_by,
                sales_opportunity_id=normalized_input["sales_opportunity_id"] or None,
                opportunity_id=normalized_input["sales_opportunity_id"] or None,
                expiry_date=normalized_input["expiry_date"],
                clauses=clauses,
                draftText=data.get("notes") or normalized_input["draftText"],
                internal_notes=routing_notes,
                workflow=workflow,
                current_stage_index=0,
                required_reviewers=workflow,
                review_stages=workflow,
                routing_reasons=review_plan["reasons"],
                routing_rules_applied=review_plan["applied_rules"],
                routing_decisions=routing_decisions,
                workflow_events=[
                    _build_workflow_event(
                        action=f"sent_to_{department.lower()}",
                        performed_by=user_name,
                        department=department,
                        notes=routing_notes,
                        metadata={"review_mode": review_mode, "workflow": workflow},
                    )
                ] if routing_notes else [],
                review_mode=review_mode
            )
            if review_mode == "parallel":
                for reviewer in workflow:
                    getattr(contract.reviews, reviewer).status = "Pending"
                    setattr(contract, f"{reviewer.lower()}_status", "Pending")
            else:
                getattr(contract.reviews, department).status = "Pending"
                setattr(contract, f"{department.lower()}_status", "Pending")

            result = await db.contracts.insert_one(contract.model_dump(by_alias=True, exclude_none=True))
            contract_id = str(result.inserted_id)

            await log_audit(
                contract_id,
                "Created",
                user_name,
                role="System",
                details=f"Contract '{title}' created and sent to {department}",
                notes=routing_notes or None
            )

        await db.notifications.insert_one({
            "for_role": department,
            "message": f"Contract '{title}' requires your review.",
            "contract_id": str(contract_id),
            "contract_title": title,
            "action": f"Go to {department} Review",
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })

        return {
                "message": "Contract sent successfully",
            "id": str(contract_id),
            "department": department,
            "status": "Under Review",
            "stage": stage,
            "workflow": workflow,
            "review_mode": review_mode
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Send contract failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

<<<<<<< Updated upstream
=======
    # Notify Legal department
    await add_notification(
        for_role="Legal",
        message=f"New {request_mode.title()} request '{contract.title}' for {contract.company} requires review.",
        type="assignment",
        contract_id=contract_id,
        contract_title=contract.title,
        action="Go to Legal Review"
    )
    
    await log_action("Create Self-Service Contract", user_email, current_user.get("role", "User"),
        f"{request_mode.title()} request '{contract.title}' created under category '{document_category}'.",
        contract_id)
    
    return {
        "message": "Self-service request submitted successfully",
        "id": contract_id
    }
>>>>>>> Stashed changes


@app.post("/api/self-service/upload-request")
async def create_upload_request(
    file: UploadFile = File(...),
    title: str = Form(""),
    company: str = Form(""),
    department: str = Form("Sales"),
    value: str = Form("0"),
    description: str = Form(""),
    submittedBy: str = Form(""),
    current_user: dict = Depends(get_current_user)
):
    user_email = current_user.get("email", submittedBy) or submittedBy or "Unknown"
    file_name = file.filename or "uploaded_document"
    file_ext = pathlib.Path(file_name).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{file_ext}'. Only PDF, DOC, and DOCX files are allowed."
        )

    contents = await file.read()
    size_mb = len(contents) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(status_code=413, detail=f"File too large. Max {MAX_FILE_SIZE_MB}MB allowed.")

    extracted_text = extract_text_from_upload(file_name, contents)
    inferred_document_type = infer_document_type(extracted_text, file_name)
    inferred_category = get_document_category(inferred_document_type)
    clause_categories = infer_clause_categories(extracted_text)
    extracted_metadata = extract_metadata_from_text(extracted_text)
    inferred_company = company or (extracted_metadata.get("party_names", ["Unknown", "Unknown"])[-1] if extracted_metadata.get("party_names") else "Unknown")

    contract = Contract(
        title=title or f"{inferred_company} - {inferred_document_type}",
        company=inferred_company or "Unknown",
        department=department or "Sales",
        value=float(value) if value else 0.0,
        stage="Under Review",
        status="Pending",
        submittedBy=user_email,
        required_reviewers=["Legal"]
    )
    contract_dump = contract.model_dump(by_alias=True, exclude_none=True)
    contract_dump["source"] = "self-service"
    contract_dump["request_mode"] = "uploaded"
    contract_dump["template_type"] = inferred_document_type
    contract_dump["document_category"] = inferred_category
    contract_dump["route"] = "legal-counsel"
    contract_dump["description"] = description or "User uploaded an external document for legal review."
    contract_dump["extractedText"] = extracted_text[:20000] if extracted_text else ""
    contract_dump["clauses"] = build_clause_shell(clause_categories, contract_dump["description"])
    contract_dump["ai_classification_result"] = {
        "document_type": inferred_document_type,
        "document_category": inferred_category,
        "clause_categories": clause_categories,
        "metadata": extracted_metadata,
        "routing_decision": "Legal Counsel Review",
        "confidence": "medium"
    }

    insert_result = await db.contracts.insert_one(contract_dump)
    contract_id = str(insert_result.inserted_id)

    # Save uploaded file as the first contract document version.
    contract_dir = UPLOAD_DIR / contract_id
    contract_dir.mkdir(parents=True, exist_ok=True)
    storage_name = f"v1_{file_name}"
    file_path = contract_dir / storage_name
    with open(file_path, "wb") as f:
        f.write(contents)

    mime_type = file.content_type
    if not mime_type or mime_type == "application/octet-stream":
        if file_ext == ".pdf":
            mime_type = "application/pdf"
        elif file_ext == ".docx":
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_ext == ".doc":
            mime_type = "application/msword"
        else:
            mime_type = "application/octet-stream"

    doc = Document(
        contractId=contract_id,
        fileName=file_name,
        fileType=mime_type,
        fileSize=len(contents),
        storagePath=str(file_path),
        version=1,
        uploadedBy=user_email,
        category="Original",
        tags=["self-service", "uploaded", inferred_document_type]
    )
    await db.documents.insert_one(doc.model_dump(by_alias=True, exclude_none=True))

    await add_notification(
        for_role="Legal",
        message=f"New uploaded request '{contract.title}' for {contract.company} requires review.",
        type="assignment",
        contract_id=contract_id,
        contract_title=contract.title,
        action="Go to Legal Review"
    )
    await log_action(
        "Create Uploaded Contract Request",
        user_email,
        current_user.get("role", "User"),
        f"Uploaded contract '{contract.title}' classified as '{inferred_document_type}'.",
        contract_id
    )

    return {
        "message": "Uploaded request submitted successfully",
        "id": contract_id,
        "request_mode": "uploaded",
        "document_category": inferred_category,
        "ai_classification_result": contract_dump["ai_classification_result"]
    }

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PDF & Annotations API
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.get("/api/contracts/{contract_id}/pdf")
async def get_contract_pdf(contract_id: str):
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract or not contract.get("pdf_path"):
        raise HTTPException(status_code=404, detail="PDF not found")
    path = Path(contract["pdf_path"])
    if not path.exists():
        raise HTTPException(status_code=404, detail="PDF file missing")
    return FileResponse(path, media_type="application/pdf", filename=f"{contract.get('title')}.pdf")

@app.get("/api/contracts/{contract_id}/pdf-annotations")
async def get_pdf_annotations(contract_id: str, doc_type: str = Query("pdf"), document_id: Optional[str] = Query(None)):
    if doc_type == "pdf":
        query = {
            "contractId": contract_id,
            "$or": [{"docType": "pdf"}, {"docType": {"$exists": False}}]
        }
    else:
        query = {"contractId": contract_id, "docType": doc_type}
    if doc_type == "docx" and document_id:
        query["targetDocumentId"] = document_id
    cursor = db.pdf_annotations.find(query).sort("createdAt", 1)
    annotations = await cursor.to_list(length=1000)
    for a in annotations:
        a["_id"] = str(a["_id"])
    return annotations

@app.post("/api/contracts/{contract_id}/pdf-annotations")
async def add_pdf_annotation(contract_id: str, data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    doc_type = data.get("doc_type", "pdf")
    ann = PdfAnnotation(
        contractId=contract_id,
        page=data.get("page", 1),
        x=data.get("x", 0.0),
        y=data.get("y", 0.0),
        width=data.get("width", 0.0),
        height=data.get("height", 0.0),
        text=data.get("text", ""),
        author=data.get("author", current_user.get("email", "Unknown")),
        role=data.get("role", current_user.get("role", "User")),
        color=data.get("color", "#f59e0b"),
        docType=doc_type,
        targetDocumentId=data.get("document_id")
    )
    result = await db.pdf_annotations.insert_one(ann.model_dump(by_alias=True, exclude_none=True))
    new_doc = await db.pdf_annotations.find_one({"_id": result.inserted_id})
    if new_doc:
        new_doc["_id"] = str(new_doc["_id"])
    return new_doc

@app.delete("/api/contracts/{contract_id}/pdf-annotations/{annotation_id}")
async def delete_pdf_annotation(contract_id: str, annotation_id: str, current_user: dict = Depends(get_current_user)):
    if not ObjectId.is_valid(annotation_id):
        raise HTTPException(status_code=400, detail="Invalid annotation ID")
    result = await db.pdf_annotations.delete_one({
        "_id": ObjectId(annotation_id),
        "contractId": contract_id
    })
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Annotation not found")
    return {"message": "Annotation deleted"}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Self-Service Actions API
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.post("/api/contracts/{contract_id}/send-to-clo")
async def send_to_clo(contract_id: str, data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    comment = data.get("review_comment", "")
    reviewer_name = data.get("reviewer_name", "Legal Counsel")
    
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "stage": "CLO Review",
            "reviews.Legal.status": "Sent to CLO",
            "reviews.Legal.comments": comment,
            "reviews.Legal.reviewedBy": reviewer_name,
            "reviews.Legal.reviewedAt": datetime.utcnow().isoformat()
        }}
    )
    
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    
    await add_notification(
        for_role="CLO",
        message=f"Contract '{contract.get('title')}' sent by Legal for CLO approval.",
        type="assignment",
        contract_id=contract_id,
        contract_title=contract.get("title"),
        action="Go to CLO Review"
    )
    
    await log_action("Escalate to CLO", reviewer_name, "Legal", f"Sent to CLO with comment: {comment[:50]}", contract_id)
    return {"message": "Sent to CLO"}

@app.post("/api/contracts/{contract_id}/return-to-user")
async def return_to_user(contract_id: str, data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    comment = data.get("review_comment", "")
    reviewer_name = data.get("reviewer_name", "Legal Counsel")
    
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "stage": "Draft",
            "status": "Changes Requested",
            "reviews.Legal.status": "Changes Requested",
            "reviews.Legal.comments": comment,
            "reviews.Legal.reviewedBy": reviewer_name,
            "reviews.Legal.reviewedAt": datetime.utcnow().isoformat()
        }}
    )
    
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    
    await add_notification(
        for_user=contract.get("submittedBy"),
        message=f"Legal requested changes for '{contract.get('title')}': {comment[:50]}",
        type="alert",
        contract_id=contract_id,
        contract_title=contract.get("title"),
        action="View Contract"
    )
    
    await log_action("Return to User", reviewer_name, "Legal", f"Returned to user with comment: {comment[:50]}", contract_id)
    return {"message": "Returned to user"}

@app.post("/api/contracts/{contract_id}/clo-decision")
async def clo_decision(contract_id: str, data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    decision = data.get("decision")
    comment = data.get("comment", "")
    clo_name = data.get("clo_name", "Chief Legal Officer")
    
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
        
    set_fields = {
        "reviews.CLO.status": decision.capitalize(),
        "reviews.CLO.comments": comment,
        "reviews.CLO.reviewedBy": clo_name,
        "reviews.CLO.reviewedAt": datetime.utcnow().isoformat()
    }
    
    if decision == "approve" or decision == "send-to-cas":
        set_fields["stage"] = "CAS Generated"
        set_fields["status"] = "Approved"
        await add_notification(
            for_role="CLO",
            message=f"CLO approved '{contract.get('title')}'. CAS record created and ready for review.",
            type="assignment",
            contract_id=contract_id,
            contract_title=contract.get("title"),
            action="Go to CAS"
        )
        # Also notify the contract submitter
        await add_notification(
            for_user=contract.get("submittedBy"),
            message=f"Your contract '{contract.get('title')}' has been approved by the CLO and entered CAS.",
            type="approval",
            contract_id=contract_id,
            contract_title=contract.get("title"),
            action="View Contract"
        )
        log_msg = f"CLO Approved with comment: {comment[:50]}"
        # Create CAS document so it shows up in CLO CAS panel
        cas_doc = services.generate_cas_document(
            contract_id=str(contract_id),
            contract_title=contract.get("title", ""),
            value=contract.get("value", 0),
            initiator=contract.get("submittedBy", "Admin")
        )
        await db.cas.insert_one(cas_doc.model_dump(by_alias=True, exclude_none=True))
    elif decision == "revise":
        set_fields["stage"] = "Under Review"
        set_fields["reviews.Legal.status"] = "Changes Requested"
        await add_notification(
            for_role="Legal",
            message=f"CLO requested revisions for '{contract.get('title')}': {comment[:80]}",
            type="alert",
            contract_id=contract_id,
            contract_title=contract.get("title"),
            action="Go to Review"
        )
        log_msg = f"CLO Requested Revision: {comment[:50]}"
    else:
        set_fields["stage"] = "Rejected"
        set_fields["status"] = "Rejected"
        await add_notification(
            for_user=contract.get("submittedBy"),
            message=f"Contract '{contract.get('title')}' was rejected by CLO.",
            type="alert",
            contract_id=contract_id,
            contract_title=contract.get("title"),
            action="View Details"
        )
        log_msg = f"CLO Rejected: {comment[:50]}"
        
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": set_fields}
    )
    
    await log_action("CLO Decision", clo_name, "CLO", log_msg, contract_id)
    return {"message": f"CLO decision '{decision}' processed"}

@app.post("/api/contracts/{contract_id}/return-to-legal")
async def return_to_legal(contract_id: str, data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    comment = data.get("review_comment", "")
    reviewer_name = current_user.get("name", "Chief Legal Officer")
    
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "stage": "Under Review",
            "status": "Changes Requested",
            "reviews.CLO.status": "Returned to Legal",
            "reviews.CLO.comments": comment,
            "reviews.CLO.reviewedBy": reviewer_name,
            "reviews.CLO.reviewedAt": datetime.utcnow().isoformat()
        }}
    )
    
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    
    await add_notification(
        for_role="Legal",
        message=f"CLO returned '{contract.get('title')}' for revision: {comment[:50]}",
        type="alert",
        contract_id=contract_id,
        contract_title=contract.get("title"),
        action="Go to Legal Review"
    )
    
    await log_action("Return to Legal", reviewer_name, "CLO", f"Returned to Legal with comment: {comment[:50]}", contract_id)
    return {"message": "Returned to Legal"}

@app.post("/api/contracts/{contract_id}/resubmit")
async def resubmit_contract(contract_id: str, data: dict = Body({}), current_user: dict = Depends(get_current_user)):
    user_email = current_user.get("email", "UnknownUser")
    
    # ─── Clear All PDF Annotations ─── 
    # When user resubmits, we delete ALL annotations so Legal/CLO get a clean document
    await db.pdf_annotations.delete_many({"contractId": contract_id})
    
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "stage": "Under Review",
            "status": "Pending",
        }}
    )
    
    # Clear the Legal review status so it re-appears in Legal queue
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "reviews.Legal.status": "Pending",
        }}
    )
    
    # Increment pdf_version to track re-review cycle
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$inc": {"pdf_version": 1}}
    )

    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    
    await add_notification(
        for_role="Legal",
        message=f"User resubmitted '{contract.get('title')}' with changes. Document is now clean for re-review.",
        type="assignment",
        contract_id=contract_id,
        contract_title=contract.get("title"),
        action="Go to Legal Review"
    )
    
    await log_action("Resubmit Contract", user_email, current_user.get("role", "User"), 
                     f"Resubmitted contract (annotations cleared, pdf_version incremented)", contract_id)
    return {"message": "Contract resubmitted successfully", "annotations_cleared": True}


@app.delete("/api/contracts/{contract_id}/pdf-annotations/clear-all")
async def clear_all_annotations(contract_id: str, current_user: dict = Depends(get_current_user)):
    """Clear all PDF annotations for a contract (called on resubmit)."""
    result = await db.pdf_annotations.delete_many({"contractId": contract_id})
    await log_action("Clear Annotations", current_user.get("email", "Unknown"), 
                     current_user.get("role", "User"), 
                     f"Cleared {result.deleted_count} annotations", contract_id)
    return {"message": f"Cleared {result.deleted_count} annotations"}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Comments & Reviews API
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@app.post("/api/contracts/{contract_id}/comments")
async def add_review_comment(contract_id: str, data: dict = Body(...)):
    comment = ReviewComment(
        contractId=contract_id,
        department=data.get("department", ""),
        clauseId=data.get("clauseId"),
        comment=data.get("comment", ""),
        commentedBy=data.get("commentedBy", "Admin"),
        parentId=data.get("parentId")
    )
    result = await db.review_comments.insert_one(
        comment.model_dump(by_alias=True, exclude_none=True)
    )
    new_doc = await db.review_comments.find_one({"_id": result.inserted_id})
    if new_doc:
        new_doc["_id"] = str(new_doc["_id"])
    return new_doc

@app.get("/api/contracts/{contract_id}/comments")
async def get_review_comments(contract_id: str, department: Optional[str] = None):
    query = {"contractId": contract_id}
    if department:
        query["department"] = department
        
    cursor = db.review_comments.find(query).sort("createdAt", 1)
    comments = await cursor.to_list(length=1000)
    for c in comments:
        c["_id"] = str(c["_id"])
    return comments

@app.delete("/api/comments/{comment_id}")
async def delete_review_comment(comment_id: str):
    try:
        result = await db.review_comments.delete_one({"_id": PyObjectId(comment_id)})
        if result.deleted_count == 1:
            return {"message": "Comment deleted successfully"}
        raise HTTPException(status_code=404, detail="Comment not found")
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/contracts/{contract_id}/save-review-comment")
async def save_review_decision(contract_id: str, data: dict = Body(...)):
    # 1. Save decision as a comment
    comment = ReviewComment(
        contractId=contract_id,
        department=data.get("department", ""),
        comment=data.get("comment", ""),
        commentedBy=data.get("reviewedBy", "Admin"),
        status=data.get("status", "Review")
    )
    await db.review_comments.insert_one(
        comment.model_dump(by_alias=True, exclude_none=True)
    )
    
    # 2. Update contract review data inline for backwards compatibility
    dept = data.get("department", "")
    status = data.get("status", "Pending")
    
    update_data = {
        f"reviews.{dept}.status": status,
        f"reviews.{dept}.comments": data.get("comment", ""),
        f"reviews.{dept}.reviewedBy": data.get("reviewedBy", "Admin"),
        f"reviews.{dept}.reviewedAt": datetime.utcnow().isoformat()
    }
    
    await db.contracts.update_one(
        {"_id": PyObjectId(contract_id)},
        {"$set": update_data}
    )
    return {"message": "Review decision saved successfully"}

@app.post("/api/contracts/{contract_id}/change-requests")
async def create_change_request(contract_id: str, data: dict = Body(...)):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
    contract = await db.contracts.find_one({"_id": PyObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    cr = ChangeRequest(
        contractId=contract_id,
        action=data.get("action", "request_changes"),
        department=data.get("department", ""),
        requestedBy=data.get("requestedBy", "Admin"),
        reviewer=data.get("reviewer") or data.get("requestedBy", "Admin"),
        title=data.get("title"),
        requestType=data.get("requestType", "Clause Change"),
        clauseId=data.get("clauseId"),
        clauseTitle=data.get("clauseTitle"),
        clause_reference=data.get("clause_reference") or data.get("clauseTitle") or data.get("fieldName"),
        fieldName=data.get("fieldName"),
        currentValue=data.get("currentValue"),
        proposedValue=data.get("proposedValue"),
        priority=data.get("priority", "Medium"),
        description=data.get("description", ""),
        comment=data.get("comment") or data.get("description", ""),
        status=data.get("status", "Open"),
        contractVersion=int(contract.get("current_version") or 1)
    )
    result = await db.change_requests.insert_one(
        cr.model_dump(by_alias=True, exclude_none=True)
    )
    await db.contracts.update_one(
        {"_id": PyObjectId(contract_id)},
        {
            "$addToSet": {"active_change_request_ids": str(result.inserted_id)},
            "$push": {
                "change_requests": {
                    "id": str(result.inserted_id),
                    "action": cr.action,
                    "department": cr.department,
                    "reviewer": cr.reviewer or cr.requestedBy,
                    "comment": cr.comment or cr.description,
                    "clause_reference": cr.clause_reference or "",
                    "timestamp": cr.createdAt,
                    "status": cr.status,
                    "linkedVersion": cr.linkedVersion,
                }
            },
            "$set": {"updatedAt": datetime.utcnow().isoformat()}
        }
    )
    new_doc = await db.change_requests.find_one({"_id": result.inserted_id})
    if new_doc:
        new_doc["_id"] = str(new_doc["_id"])
    return new_doc

@app.get("/api/contracts/{contract_id}/change-requests")
async def get_change_requests(contract_id: str, department: Optional[str] = None):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
    query = {"contractId": contract_id}
    if department:
        query["department"] = department
    cursor = db.change_requests.find(query).sort("createdAt", 1)
    results = await cursor.to_list(length=1000)
    for r in results:
        r["_id"] = str(r["_id"])
    return results

@app.put("/api/change-requests/{cr_id}")
async def update_change_request(cr_id: str, data: dict = Body(...)):
    if not ObjectId.is_valid(cr_id):
        raise HTTPException(status_code=400, detail="Invalid change request ID")
    update_fields = {}
    if "status" in data:
        update_fields["status"] = data["status"]
    if "resolution" in data:
        update_fields["resolution"] = data["resolution"]
    if "resolvedBy" in data:
        update_fields["resolvedBy"] = data["resolvedBy"]
    update_fields["updatedAt"] = datetime.utcnow().isoformat()
    if data.get("status") in ["Resolved", "Rejected"]:
        update_fields["resolvedAt"] = datetime.utcnow().isoformat()
    if not update_fields:
        raise HTTPException(status_code=400, detail="No valid fields to update")
    await db.change_requests.update_one(
        {"_id": PyObjectId(cr_id)},
        {"$set": update_fields}
    )
    updated = await db.change_requests.find_one({"_id": PyObjectId(cr_id)})
    if updated and updated.get("contractId") and data.get("status") in ["Resolved", "Rejected"]:
        await db.contracts.update_one(
            {"_id": PyObjectId(updated["contractId"])},
            {"$pull": {"active_change_request_ids": str(updated["_id"])}, "$set": {"updatedAt": datetime.utcnow().isoformat()}}
        )
    if updated and updated.get("contractId"):
        contract = await db.contracts.find_one({"_id": PyObjectId(updated["contractId"])})
        if contract:
            updated_summary = _change_request_summary(updated)
            change_requests = contract.get("change_requests") or []
            replaced = False
            for idx, item in enumerate(change_requests):
                if str(item.get("id")) == str(updated["_id"]):
                    change_requests[idx] = updated_summary
                    replaced = True
                    break
            if not replaced:
                change_requests.append(updated_summary)
            await db.contracts.update_one(
                {"_id": PyObjectId(updated["contractId"])},
                {"$set": {"change_requests": change_requests, "updatedAt": datetime.utcnow().isoformat()}}
            )
    if updated:
        updated["_id"] = str(updated["_id"])
    return updated

@app.post("/api/contracts/{contract_id}/escalate")
async def escalate_contract(contract_id: str, data: dict = Body(...)):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
        
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
        
    reason = data.get("reason", "")
    department = data.get("department", "Unknown")
    escalatedBy = data.get("escalatedBy", "Admin")
    
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "escalated": True,
            "escalatedBy": escalatedBy,
            "escalatedAt": datetime.utcnow().isoformat(),
            "escalationReason": reason
        }}
    )
    
    title = contract.get("title", "Unknown Contract")
    
    # Create notification for Manager and CEO roles
    for target_role in ["Manager", "CEO"]:
        await db.notifications.insert_one({
            "for_role": target_role,
            "message": f"Contract '{title}' escalated by {department}. Reason: {reason}",
            "contract_id": contract_id,
            "type": "escalation",
            "department": department,
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })
        
    return {"message": "Contract escalated successfully"}

@app.delete("/api/contracts/{contract_id}")
async def delete_contract(contract_id: str):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(400, "Invalid contract ID")
    result = await db.contracts.delete_one({"_id": ObjectId(contract_id)})
    if result.deleted_count == 0:
        raise HTTPException(404, "Contract not found")
    return {"message": "Contract deleted successfully"}


@app.get("/api/contracts/{contract_id}/versions")
async def get_contract_versions(contract_id: str):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    cursor = db.contract_versions.find({"contractId": contract_id}).sort("version", -1)
    versions = await cursor.to_list(length=200)
    for version in versions:
        version["id"] = str(version["_id"])
        del version["_id"]
    return versions


@app.post("/api/contracts/{contract_id}/redlines")
async def apply_contract_redline(contract_id: str, data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    clauses = contract.get("clauses") or []
    clause_id = data.get("clauseId")
    clause_title = data.get("clauseTitle")
    redlined_text = data.get("redlinedText")
    if not redlined_text:
        raise HTTPException(status_code=400, detail="redlinedText is required")

    target_clause = None
    updated_clauses = []
    for clause in clauses:
        matches_id = clause_id and str(clause.get("id")) == str(clause_id)
        matches_title = clause_title and str(clause.get("title")) == str(clause_title)
        if matches_id or matches_title:
            target_clause = clause
            clause_redlines = clause.get("redline_history") or []
            clause_redlines.append({
                "id": str(ObjectId()),
                "clauseId": clause.get("id"),
                "clauseTitle": clause.get("title"),
                "department": data.get("department", ""),
                "originalText": data.get("originalText") or clause.get("content") or clause.get("text") or "",
                "redlinedText": redlined_text,
                "justification": data.get("justification"),
                "issues": data.get("issues", []),
                "createdBy": data.get("createdBy") or user_name,
                "createdAt": datetime.utcnow().isoformat(),
            })
            updated_clause = {
                **clause,
                "content": redlined_text,
                "text": redlined_text,
                "status": "Reviewed",
                "redline_history": clause_redlines,
            }
            updated_clauses.append(updated_clause)
        else:
            updated_clauses.append(clause)

    if target_clause is None:
        raise HTTPException(status_code=404, detail="Clause not found for redline application")

    redline_entry = {
        "id": str(ObjectId()),
        "clauseId": target_clause.get("id"),
        "clauseTitle": target_clause.get("title"),
        "department": data.get("department", ""),
        "originalText": data.get("originalText") or target_clause.get("content") or target_clause.get("text") or "",
        "redlinedText": redlined_text,
        "justification": data.get("justification"),
        "issues": data.get("issues", []),
        "createdBy": data.get("createdBy") or user_name,
        "createdAt": datetime.utcnow().isoformat(),
    }

    next_snapshot = dict(contract)
    next_snapshot["clauses"] = updated_clauses
    next_snapshot["redline_history"] = [*(contract.get("redline_history") or []), redline_entry]
    next_snapshot["current_version"] = int(contract.get("current_version") or 1) + 1
    next_snapshot["versions"] = [
        *(contract.get("versions") or []),
        _build_version_summary(
            version=next_snapshot["current_version"],
            file_url=await _latest_contract_file_url(contract_id),
            uploaded_by=data.get("createdBy") or user_name,
            changes=f"Updated clause '{target_clause.get('title', '')}' via redline.",
        ),
    ]
    next_snapshot["updatedAt"] = datetime.utcnow().isoformat()

    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {
            "$set": {
                "clauses": updated_clauses,
                "redline_history": next_snapshot["redline_history"],
                "current_version": next_snapshot["current_version"],
                "versions": next_snapshot["versions"],
                "updatedAt": next_snapshot["updatedAt"],
            }
        }
    )
    await _store_contract_version(
        contract_id=contract_id,
        previous_contract=contract,
        next_snapshot=next_snapshot,
        created_by=user_name,
        change_type="redline",
        summary=f"Redline applied to clause '{target_clause.get('title', '')}'.",
    )

    await log_audit(contract_id, "Redline Applied", user_name, role=data.get("department", "System"), department=data.get("department"), details=f"Redline applied to clause '{target_clause.get('title', '')}'.")
    refreshed = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    refreshed["id"] = str(refreshed["_id"])
    del refreshed["_id"]
    return _normalize_contract_document(refreshed)


@app.post("/api/contracts/{contract_id}/resubmit-review")
async def resubmit_contract_for_review(contract_id: str, data: dict = Body({}), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = _normalize_contract_document(contract)
    requested_mode = data.get("review_mode") or contract.get("review_mode", "sequential")
    review_plan = await resolve_review_plan(
        title=contract.get("title", "Untitled"),
        value=contract.get("value", 0),
        category=contract.get("category", "General"),
        risk=contract.get("risk_classification", "Medium"),
        clauses=contract.get("clauses", []),
        business_unit=contract.get("business_unit", ""),
        requested_mode=requested_mode,
    )
    workflow = review_plan["workflow"]
    review_mode = review_plan["review_mode"]
    first_reviewer = workflow[0] if workflow else "Legal"

    update_fields: Dict[str, Any] = {
        "workflow": workflow,
        "review_stages": workflow,
        "required_reviewers": workflow,
        "review_mode": review_mode,
        "routing_reasons": review_plan["reasons"],
        "routing_rules_applied": review_plan["applied_rules"],
        "routing_decisions": review_plan["routing_decisions"],
        "status": "Under Review",
        "stage": "Under Review" if review_mode == "parallel" else _build_send_stage(first_reviewer),
        "department": first_reviewer,
        "current_stage_index": 0,
        "dueAt": _due_at_iso(),
        "rejection_reason": None,
        "rejected_by": None,
        "rejectedAt": None,
        "last_review_resubmitted_at": datetime.utcnow().isoformat(),
        "updatedAt": datetime.utcnow().isoformat(),
    }

    active_crs = await db.change_requests.find({
        "contractId": contract_id,
        "status": {"$in": ["Open", "Pending", "Resubmitted"]}
    }).to_list(length=500)
    active_cr_ids = [str(item["_id"]) for item in active_crs]
    update_fields["active_change_request_ids"] = active_cr_ids
    current_version = int(contract.get("current_version") or 1)
    next_version = current_version + 1
    change_summary = data.get("changes") or data.get("change_summary") or "Contract resubmitted after requested changes."
    latest_file_url = await _latest_contract_file_url(contract_id)

    for dept in REVIEW_DEPARTMENTS:
        dept_change_request_ids = [str(item["_id"]) for item in active_crs if item.get("department") == dept]
        if dept in workflow:
            is_pending = review_mode == "parallel" or dept == first_reviewer
            update_fields[f"reviews.{dept}.status"] = "Pending" if is_pending else "Not Started"
            update_fields[f"reviews.{dept}.comments"] = ""
            update_fields[f"reviews.{dept}.reviewedBy"] = ""
            update_fields[f"reviews.{dept}.reviewedAt"] = None
            update_fields[f"reviews.{dept}.decisionType"] = None
            update_fields[f"reviews.{dept}.requiresChanges"] = False
            update_fields[f"reviews.{dept}.changeRequestIds"] = dept_change_request_ids
            update_fields[f"{dept.lower()}_status"] = "Pending" if is_pending else "Not Started"
        else:
            update_fields[f"reviews.{dept}.status"] = "Not Started"
            update_fields[f"reviews.{dept}.comments"] = ""
            update_fields[f"reviews.{dept}.reviewedBy"] = ""
            update_fields[f"reviews.{dept}.reviewedAt"] = None
            update_fields[f"reviews.{dept}.decisionType"] = None
            update_fields[f"reviews.{dept}.requiresChanges"] = False
            update_fields[f"reviews.{dept}.changeRequestIds"] = []
            update_fields[f"{dept.lower()}_status"] = "Not Started"

    existing_versions = contract.get("versions") or []
    version_entry = _build_version_summary(
        version=next_version,
        file_url=latest_file_url,
        uploaded_by=user_name,
        changes=change_summary,
        change_request_ids=active_cr_ids,
    )
    update_fields["current_version"] = next_version
    update_fields["versions"] = [*existing_versions, version_entry]
    update_fields["workflow_events"] = [
        *(contract.get("workflow_events") or []),
        _build_workflow_event(
            action="resubmitted_for_review",
            performed_by=user_name,
            department="System",
            notes=change_summary,
            metadata={"review_mode": review_mode, "workflow": workflow, "change_request_ids": active_cr_ids, "version": next_version}
        )
    ]

    next_snapshot = dict(contract)
    for key, value in update_fields.items():
        if "." not in key:
            next_snapshot[key] = value

    await db.contracts.update_one({"_id": ObjectId(contract_id)}, {"$set": update_fields})
    await db.change_requests.update_many(
        {"contractId": contract_id, "status": {"$in": ["Open", "Pending"]}},
        {"$set": {"status": "Resubmitted", "updatedAt": datetime.utcnow().isoformat(), "linkedVersion": next_version}}
    )

    await _store_contract_version(
        contract_id=contract_id,
        previous_contract=contract,
        next_snapshot=next_snapshot,
        created_by=user_name,
        change_type="resubmission",
        summary=change_summary,
        file_url=latest_file_url,
        change_request_ids=active_cr_ids,
    )

    existing_change_requests = contract.get("change_requests") or []
    refreshed_change_requests = []
    active_change_request_map = {}
    for item in active_crs:
        item["linkedVersion"] = next_version
        active_change_request_map[str(item["_id"])] = _change_request_summary(item)

    for item in existing_change_requests:
        item_id = str(item.get("id") or "")
        if item_id and item_id in active_change_request_map:
            refreshed_change_requests.append(active_change_request_map.pop(item_id))
        else:
            refreshed_change_requests.append(item)

    for item in active_change_request_map.values():
        refreshed_change_requests.append(item)
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {"change_requests": refreshed_change_requests if refreshed_change_requests else existing_change_requests}}
    )

    await log_audit(contract_id, "Resubmitted for Review", user_name, role="System", details=f"Contract resubmitted in {review_mode} mode. Workflow: {', '.join(workflow)}")
    for dept in workflow:
        await db.notifications.insert_one({
            "for_role": dept,
            "message": f"Contract '{contract.get('title', '')}' has been resubmitted for review.",
            "contract_id": contract_id,
            "contract_title": contract.get("title", ""),
            "action": f"Go to {dept} Review",
            "department": "System",
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })

    return {
        "message": "Contract resubmitted for review successfully",
        "workflow": workflow,
        "review_mode": review_mode,
        "active_change_request_ids": active_cr_ids
    }

@app.post("/api/contracts/{contract_id}/review")
async def submit_review(contract_id: str, department: str = Body(...), status: str = Body(...), comments: str = Body(""), reviewer: str = Body("Admin")):
    from datetime import datetime
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
        
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
        
    contract = _normalize_contract_document(contract)
    workflow = contract.get("workflow", ["Legal"])
    review_mode = contract.get("review_mode", "sequential")
    if department not in workflow:
        raise HTTPException(status_code=403, detail=f"Department '{department}' is not an authorized reviewer for this contract.")
    
    # SAFEGUARD: Prevent Duplicate Approval
    existing_status = contract.get("reviews", {}).get(department, {}).get("status")
    if existing_status == "Approved":
        raise HTTPException(status_code=400, detail=f"Department '{department}' has already approved this contract.")

    current_stage_index = int(contract.get("current_stage_index", 0))
    current_department = workflow[current_stage_index] if workflow else department
    if contract.get("status") not in ["Under Review", "Overdue", "Pending"]:
        raise HTTPException(status_code=400, detail=f"Contract is not in a reviewable state ({contract.get('status')}).")
    if department != current_department and review_mode != "parallel":
        raise HTTPException(status_code=403, detail=f"Contract is currently assigned to {current_department}, not {department}.")

    update_data = {
        f"reviews.{department}.status": status,
        f"reviews.{department}.comments": comments,
        f"reviews.{department}.reviewedBy": reviewer,
        f"reviews.{department}.reviewedAt": datetime.utcnow().isoformat(),
        f"{department.lower()}_status": status,
        "updatedAt": datetime.utcnow().isoformat()
    }
    
    if status == "Rejected":
        rejection_update = {
            **update_data,
            "status": "Rejected",
            "stage": "Rejected",
            "department": "Sales",
            "rejection_reason": comments or "Rejected without comment",
            "rejected_by": reviewer,
            "rejectedAt": datetime.utcnow().isoformat(),
            "dueAt": None,
            "current_stage_index": 0
        }
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": rejection_update}
        )
        
        await log_audit(contract_id, "Rejected", reviewer, role=department, department=department, details=f"Rejected by {department}. Comments: {comments}")
        
        c_title = contract.get("title", "")
        owner_dept = contract.get("created_by") or contract.get("submittedBy") or "Admin"
        await db.notifications.insert_one({
            "for_role": owner_dept,
            "message": f"Contract '{c_title}' rejected by {department}.",
            "contract_id": contract_id,
            "contract_title": c_title,
            "action": "Review Rejection",
            "department": department,
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })
        
        if owner_dept != "Admin":
            await db.notifications.insert_one({
                "for_role": "Admin",
                "message": f"Contract '{c_title}' rejected by {department}.",
                "contract_id": contract_id,
                "contract_title": c_title,
                "action": "View Draft",
                "department": department,
                "read": False,
                "createdAt": datetime.utcnow().isoformat()
            })
        
        # Email notification for Rejection
        try:
            initiator_name = contract.get("submittedBy", "")
            if initiator_name:
                user_doc = await db.users.find_one({"name": initiator_name})
                if user_doc and user_doc.get("email"):
                    await send_email_notification(
                        to_email=user_doc["email"],
                        subject=f"[Apeiro CLM] Contract Rejected: {contract.get('title', '')}",
                        body_html=f"""
                        <div style="font-family: sans-serif; max-width: 600px; padding: 20px; border: 1px solid #eee; border-radius: 10px;">
                            <h2 style="color: #ef4444;">Contract Rejected</h2>
                            <p>Hi {initiator_name},</p>
                            <p>Your contract <strong>{contract.get('title', '')}</strong> has been rejected by the <strong>{department}</strong> department.</p>
                            <p><strong>Reason/Comments:</strong> {comments or 'No specific reason provided.'}</p>
                            <hr style="border: none; border-top: 1px solid #eee; margin: 20px 0;" />
                            <p style="font-size: 12px; color: #666;">This is an automated message from Apeiro CLM.</p>
                        </div>
                        """
                    )
        except Exception as e:
            print(f"Rejection email failed: {e}")
            
        return {"message": f"Contract rejected by {department}"}

<<<<<<< Updated upstream
    if status == "Changes Requested":
        open_change_requests = await db.change_requests.find({
            "contractId": contract_id,
            "department": department,
            "status": {"$in": ["Open", "Pending", "Resubmitted"]}
        }).to_list(length=200)

        if not open_change_requests:
            auto_change_request = ChangeRequest(
                contractId=contract_id,
                action="request_changes",
                department=department,
                requestedBy=reviewer,
                reviewer=reviewer,
                requestType="General Review Change",
                description=comments or f"{department} requested contract changes.",
                comment=comments or f"{department} requested contract changes.",
                status="Open",
                contractVersion=int(contract.get("current_version") or 1),
            )
            auto_result = await db.change_requests.insert_one(
                auto_change_request.model_dump(by_alias=True, exclude_none=True)
            )
            auto_doc = await db.change_requests.find_one({"_id": auto_result.inserted_id})
            if auto_doc:
                open_change_requests = [auto_doc]
                await db.contracts.update_one(
                    {"_id": ObjectId(contract_id)},
                    {
                        "$addToSet": {"active_change_request_ids": str(auto_result.inserted_id)},
                        "$push": {"change_requests": _change_request_summary(auto_doc)},
                        "$set": {"updatedAt": datetime.utcnow().isoformat()}
                    }
                )

        change_request_ids = [str(item["_id"]) for item in open_change_requests]
        change_update = {
            **update_data,
            f"reviews.{department}.decisionType": "Changes Requested",
            f"reviews.{department}.requiresChanges": True,
            f"reviews.{department}.changeRequestIds": change_request_ids,
            "status": "Changes Requested",
            "stage": "Changes Requested",
            "department": contract.get("created_by") or contract.get("submittedBy") or "Sales",
            "dueAt": None,
            "active_change_request_ids": list({*(contract.get("active_change_request_ids") or []), *change_request_ids}),
            "workflow_events": [
                *(contract.get("workflow_events") or []),
                _build_workflow_event(
                    action="request_changes",
                    performed_by=reviewer,
                    department=department,
                    notes=comments,
                    metadata={"change_request_ids": change_request_ids}
                )
            ],
=======
    # ─── Notification insertion for next department ───
    notification_map = {
        "Legal": {
            "for_role": "Finance",
            "message": "Legal has approved '{title}'. You can now review it.",
            "action": "Go to Finance Review"
        },
        "Finance": {
            "for_role": "Compliance",
            "message": "Finance has approved '{title}'. You can now review it.",
            "action": "Go to Compliance Review"
        },
        "Compliance": {
            "for_role": "Procurement",
            "message": "Compliance has approved '{title}'. You can now review it.",
            "action": "Go to Procurement Review"
        },
        "Procurement": {
            "for_role": "CLO",
            "message": "All reviews complete for {title}. CAS generated!",
            "action": "Go to CAS"
>>>>>>> Stashed changes
        }
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": change_update}
        )

        await log_audit(
            contract_id,
            "Changes Requested",
            reviewer,
            role=department,
            department=department,
            details=f"Changes requested by {department}. Comments: {comments}",
            notes=comments
        )

        c_title = contract.get("title", "")
        contract_owner = contract.get("contract_owner") or contract.get("submittedBy") or "Admin"
        await db.notifications.insert_one({
            "for_role": contract_owner,
            "message": f"Contract '{c_title}' requires changes from {department}.",
            "contract_id": contract_id,
            "contract_title": c_title,
            "action": "Update Contract",
            "department": department,
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })

        return {
            "message": f"Changes requested by {department}",
            "status": "Changes Requested",
            "active_change_request_ids": change_request_ids
        }

    next_stage = contract.get("stage")
    next_department = department
    next_index = current_stage_index
    next_status = "Under Review"

    if review_mode == "parallel" and status == "Approved":
        updated_reviews = contract.get("reviews", {}).copy()
        updated_reviews[department] = {
            **updated_reviews.get(department, {}),
            "status": status,
            "comments": comments,
            "reviewedBy": reviewer,
            "reviewedAt": datetime.utcnow().isoformat(),
        }
        pending_parallel_reviewers = [
            dept for dept in workflow
            if updated_reviews.get(dept, {}).get("status") != "Approved"
        ]

        if pending_parallel_reviewers:
            update_data["status"] = "Under Review"
            update_data["stage"] = "Under Review"
            update_data["department"] = contract.get("department") or workflow[0]
            update_data["current_stage_index"] = 0

            await db.contracts.update_one(
                {"_id": ObjectId(contract_id)},
                {"$set": {**update_data, "stage": "Under Review"}}
            )

            await log_audit(
                contract_id,
                "Approved",
                reviewer,
                role=department,
                department=department,
                details=f"Approved by {department}. Comments: {comments}",
                event_type=AUDIT_EVENT_TYPES["REVIEW_APPROVED"],
                message=f"{department} review approved.",
                metadata={"department": department, "comments": comments},
            )
            await log_audit(
                contract_id,
                "Parallel Review Pending",
                "System",
                role="System",
                details=f"Waiting on parallel approvals from: {', '.join(pending_parallel_reviewers)}."
            )

            return {
                "message": "Review submitted successfully",
                "pending_parallel_reviewers": pending_parallel_reviewers,
                "review_mode": "parallel",
                "all_parallel_reviews_complete": False
            }

        next_status = "Approved"
        next_stage = "CAS Generated"
        update_data["status"] = "Approved"
        update_data["stage"] = next_stage
        update_data["department"] = department
        update_data["current_stage_index"] = max(len(workflow) - 1, 0)
        update_data["dueAt"] = None
        await log_audit(
            contract_id,
            "Workflow Approved",
            reviewer,
            role=department,
            department=department,
            details="All required parallel reviewers approved. Awaiting CAS/DOA.",
            event_type=AUDIT_EVENT_TYPES["WORKFLOW_COMPLETED"],
            message="All required reviews completed.",
            metadata={"review_mode": "parallel", "department": department},
        )
        await log_audit(
            contract_id,
            "CAS Generated",
            "System",
            role="System",
            details="All reviews completed. CAS generated automatically.",
            event_type=AUDIT_EVENT_TYPES["CAS_GENERATED"],
            message="CAS generated after review completion.",
        )
        await db.notifications.insert_one({
            "for_role": "Admin",
            "message": f"All parallel reviews complete for '{contract.get('title', '')}'. CAS generated!",
            "contract_id": contract_id,
            "contract_title": contract.get("title", ""),
            "action": "Go to CAS",
            "department": department,
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })
    elif current_stage_index + 1 < len(workflow):
        next_index = current_stage_index + 1
        next_department = workflow[next_index]
        next_stage = _build_send_stage(next_department)
        update_data["department"] = next_department
        update_data["current_stage_index"] = next_index
        update_data["status"] = "Under Review"
        update_data["stage"] = next_stage
        update_data["dueAt"] = _due_at_iso()
        update_data[f"reviews.{next_department}.status"] = "Pending"
        update_data[f"{next_department.lower()}_status"] = "Pending"

        await log_audit(
            contract_id,
            "Moved to Next Stage",
            "System",
            role="System",
            details=f"Moved from {department} to {next_department}.",
            event_type=AUDIT_EVENT_TYPES["ROUTED"],
            message=f"Workflow advanced from {department} to {next_department}.",
            metadata={"from": department, "to": next_department},
        )
        await db.notifications.insert_one({
            "for_role": next_department,
            "message": f"{department} approved '{contract.get('title', '')}'. It is now assigned to you.",
            "contract_id": contract_id,
            "contract_title": contract.get("title", ""),
            "action": f"Go to {next_department} Review",
            "department": department,
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })
    else:
        next_status = "Approved"
        next_stage = "CAS Generated"
        update_data["status"] = "Approved"
        update_data["stage"] = next_stage
        update_data["department"] = department
        update_data["current_stage_index"] = current_stage_index
        update_data["dueAt"] = None
        await log_audit(
            contract_id,
            "Workflow Approved",
            reviewer,
            role=department,
            department=department,
            details=f"All required departments approved. Awaiting CAS/DOA.",
            event_type=AUDIT_EVENT_TYPES["WORKFLOW_COMPLETED"],
            message="All required reviews completed.",
            metadata={"review_mode": "sequential", "department": department},
        )
        await log_audit(
            contract_id,
            "CAS Generated",
            "System",
            role="System",
            details="All reviews completed. CAS generated automatically.",
            event_type=AUDIT_EVENT_TYPES["CAS_GENERATED"],
            message="CAS generated after review completion.",
        )
        await db.notifications.insert_one({
            "for_role": "Admin",
            "message": f"All reviews complete for '{contract.get('title', '')}'. CAS generated!",
            "contract_id": contract_id,
            "contract_title": contract.get("title", ""),
            "action": "Go to CAS",
            "department": department,
            "read": False,
            "createdAt": datetime.utcnow().isoformat()
        })

    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {**update_data, "stage": next_stage}}
    )
    
    if status == "Approved":
        await log_audit(
            contract_id,
            "Approved",
            reviewer,
            role=department,
            department=department,
            details=f"Approved by {department}. Comments: {comments}",
            event_type=AUDIT_EVENT_TYPES["REVIEW_APPROVED"],
            message=f"{department} review approved.",
            metadata={"department": department, "comments": comments},
        )
    
    if next_status == "Approved":
        updated_contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
        if not _can_generate_cas(updated_contract):
            return {"message": "Review submitted successfully"}

        existing_cas = await db.cas.find_one({"contractId": str(contract_id)})
        if existing_cas:
            await db.contracts.update_one(
                {"_id": ObjectId(contract_id)},
                {"$set": {"stage": "DOA Approval"}}
            )
            return {"message": "Review submitted successfully"}

        settings_doc = await db.settings.find_one({"key": "doa_thresholds"})
        thresholds = settings_doc.get("value") if settings_doc else {"Manager": {"min": 0, "max": 10000}, "CEO": {"min": 10001, "max": 0}}
        
        cas_doc = await _build_cas_for_contract(updated_contract, thresholds)
        await db.cas.insert_one(cas_doc.model_dump(by_alias=True, exclude_none=True))
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": {"stage": "DOA Approval", "doa_stage": cas_doc.doa_stage, "doa_status": cas_doc.doa_status}}
        )
        
    return {"message": "Review submitted successfully"}

@app.post("/api/contracts/{contract_id}/generate-cas")
async def manual_generate_cas(contract_id: str):
    # Endpoint to manually generate CAS if needed
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    if not _can_generate_cas(contract):
        raise HTTPException(status_code=400, detail="CAS is locked until all workflow stages are approved.")

    existing_cas = await db.cas.find_one({"contractId": str(contract_id)})
    if existing_cas:
        raise HTTPException(status_code=400, detail="CAS already exists for this contract.")
        
    # Fetch dynamic thresholds
    settings_doc = await db.settings.find_one({"key": "doa_thresholds"})
    thresholds = settings_doc.get("value") if settings_doc else {"Manager": {"min": 0, "max": 10000}, "CEO": {"min": 10001, "max": 0}}
    
    cas_doc = await _build_cas_for_contract(contract, thresholds)
    result = await db.cas.insert_one(cas_doc.model_dump(by_alias=True, exclude_none=True))
    
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {"stage": "DOA Approval", "status": "Approved", "doa_stage": cas_doc.doa_stage, "doa_status": cas_doc.doa_status}}
    )
    
    await log_audit(contract_id, "CAS Generated", "Admin", role="Admin", details="Manual CAS generation triggered by Admin.")
    await log_audit(contract_id, "Routed", "Admin", role="Admin", details="Workflow advanced to DOA Approval.")
    
    return {
        "message": "CAS generated and contract moved to DOA Approval",
        "cas_id": str(result.inserted_id),
        "cas": _normalize_cas_document(cas_doc.model_dump(by_alias=True, exclude_none=True))
    }

@app.get("/api/cas")
async def get_cas():
    cursor = db.cas.find({})
    cas_records = await cursor.to_list(length=1000)
    return [await _hydrate_cas_document(c) for c in cas_records]


@app.get("/api/cas/{cas_id}")
async def get_cas_by_id(cas_id: str):
    if not ObjectId.is_valid(cas_id):
        raise HTTPException(400, "Invalid CAS ID")
    cas = await db.cas.find_one({"_id": ObjectId(cas_id)})
    if not cas:
        raise HTTPException(404, "CAS not found")
    return await _hydrate_cas_document(cas)


@app.get("/api/cas/{cas_id}/export")
async def export_cas_pdf(cas_id: str):
    if not ObjectId.is_valid(cas_id):
        raise HTTPException(400, "Invalid CAS ID")
    cas = await db.cas.find_one({"_id": ObjectId(cas_id)})
    if not cas:
        raise HTTPException(404, "CAS not found")

    normalized = await _hydrate_cas_document(cas)
    approval_chain = normalized.get("approvalChain") or []
    approval_lines = [
        f"{step.get('role', 'Step')}: {step.get('status', 'Pending')} | Name: {step.get('name', '')} | Approved By: {step.get('approvedBy') or '-'} | Time: {step.get('timestamp') or '-'}"
        for step in approval_chain
    ]

    lines = [
        f"Contract Title: {normalized.get('contractTitle', '')}",
        f"Contract ID: {normalized.get('contractId', '')}",
        f"CAS ID: {normalized.get('id', '')}",
        f"Status: {normalized.get('status', '')}",
        f"Business Unit: {normalized.get('businessUnit', '')}",
        f"Department: {normalized.get('department', '')}",
        f"Agreement Type: {normalized.get('agreementType', '')}",
        f"Value: {normalized.get('value', 0)}",
        f"Cost Center: {normalized.get('cost_center', '') or ''}",
        f"Project Name: {normalized.get('project_name', '') or ''}",
        f"Effective Date: {normalized.get('effective_date', '') or ''}",
        f"Execution Date: {normalized.get('execution_date', '') or ''}",
        f"Review Departments: {', '.join(normalized.get('reviewDepartments') or [])}",
        f"Routing Reasons: {'; '.join(normalized.get('routingReasons') or [])}",
        "",
        "Key Notes:",
        normalized.get("keyNotes", "") or "",
        "",
        "Approval Chain:",
        *approval_lines,
    ]

    pdf_bytes = _build_simple_pdf_bytes("Contract Approval Sheet", lines)
    filename = f"cas-{normalized.get('id', cas_id)}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@app.delete("/api/cas/{cas_id}")
async def delete_cas_record(cas_id: str):
    if not ObjectId.is_valid(cas_id):
        raise HTTPException(400, "Invalid CAS ID")
    result = await db.cas.delete_one({"_id": ObjectId(cas_id)})
    if result.deleted_count == 0:
        raise HTTPException(404, "CAS not found")
    return {"message": "CAS deleted successfully"}

<<<<<<< Updated upstream
=======
@app.get("/api/contracts/{contract_id}/editor-content")
async def get_editor_content(contract_id: str):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    # 1. Prefer editable text extracted from the latest DOCX to preserve paragraph boundaries.
    latest_docx_path = contract.get("latest_docx_path")
    if latest_docx_path:
        try:
            from docx import Document as WordDocument
            source = pathlib.Path(latest_docx_path)
            if source.exists():
                doc_obj = WordDocument(str(source))
                lines = [p.text for p in doc_obj.paragraphs]
                for table in doc_obj.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                lines.append(p.text)
                extracted = "\n".join(lines).strip()
                if len(extracted) > 20:
                    return {"content": extracted}
        except Exception:
            pass

    # 2. Return previously saved draft if it exists and is substantive
    if contract.get("draftText") and len(contract["draftText"].strip()) > 20:
        return {"content": contract["draftText"]}

    # 3. Extract text from the contract's own PDF file (self-service NDA path)
    pdf_path_str = contract.get("pdf_path")
    if pdf_path_str:
        try:
            pdf_path = pathlib.Path(pdf_path_str)
            if pdf_path.exists():
                import PyPDF2
                from io import BytesIO
                with open(pdf_path, "rb") as f:
                    raw = f.read()
                reader = PyPDF2.PdfReader(BytesIO(raw))
                extracted = ""
                for page in reader.pages:
                    extracted += (page.extract_text() or "") + "\n"
                extracted = extracted.strip()
                if len(extracted) > 30:
                    return {"content": extracted}
        except Exception as exc:
            pass  # fall through to next option

    # 4. Use extractedText field stored during contract generation
    extracted_text = (contract.get("extractedText") or "").strip()
    if len(extracted_text) > 30:
        return {"content": extracted_text}

    # 5. Try the uploaded documents collection (manually uploaded contracts)
    doc = await db.documents.find_one({"contractId": contract_id}, sort=[("uploadedAt", -1)])
    if doc:
        file_path = pathlib.Path(doc["storagePath"])
        if file_path.exists():
            try:
                with open(file_path, "rb") as f:
                    contents = f.read()
                text = ""
                if doc["fileName"].endswith('.pdf'):
                    import PyPDF2
                    from io import BytesIO
                    pdf_reader = PyPDF2.PdfReader(BytesIO(contents))
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ""
                elif doc["fileName"].endswith('.docx'):
                    import docx as _docx
                    from io import BytesIO
                    doc_file = _docx.Document(BytesIO(contents))
                    text = "\n".join([p.text for p in doc_file.paragraphs])
                else:
                    text = contents.decode('utf-8', errors='ignore')
                if text.strip() and len(text.strip()) > 30:
                    return {"content": text.strip()}
            except Exception:
                pass

    # 6. Fallback: build readable text from contract metadata + clauses
    title = contract.get("title", "Contract")
    company = contract.get("company", "Counterparty")
    description = contract.get("description", "")
    clauses = contract.get("clauses", [])
    lines = [
        title.upper(),
        "",
        f"Between: Apeiro Digital Ltd (\"Company\")",
        f"And:     {company}",
        "",
    ]
    if description:
        lines += [f"Purpose: {description}", ""]
    if clauses:
        lines += ["CLAUSES", ""]
        for c in clauses:
            lines.append(f"{c.get('type', 'Clause')}")
            lines.append(c.get('text', ''))
            lines.append("")
    else:
        lines += [
            "TERMS AND CONDITIONS",
            "",
            "Both parties agree to hold all confidential information in strict confidence.",
            "",
            "This agreement shall remain in effect for one (1) year from the Effective Date.",
        ]
    return {"content": "\n".join(lines)}

@app.post("/api/contracts/{contract_id}/save-editor-content")
async def save_editor_content(contract_id: str, data: dict = Body(...), current_user: dict = Depends(get_current_user)):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    content = data.get("content", "")
    paragraph_texts = data.get("paragraph_texts")
    base_document_id = data.get("base_document_id")
    normalized_content = (content or "").strip()
    looks_like_html = bool(normalized_content and re.search(r"<[a-zA-Z][^>]*>", normalized_content))
    plain_text_fallback = _html_to_plain_text_for_docx(content) if looks_like_html else (content or "")
    if looks_like_html and (not isinstance(paragraph_texts, list) or not paragraph_texts):
        paragraph_texts = _extract_paragraph_texts_from_html(content)
    if isinstance(paragraph_texts, list):
        paragraph_texts = [str(t or "") for t in paragraph_texts]
    else:
        paragraph_texts = None

    base_docx_path = None
    if base_document_id and ObjectId.is_valid(base_document_id):
        base_doc = await db.documents.find_one({"_id": ObjectId(base_document_id)})
        if base_doc and str(base_doc.get("fileName", "")).lower().endswith(".docx"):
            base_docx_path = base_doc.get("storagePath")
    user_label = current_user.get("email") or data.get("user", "Admin")

    # ── 1. Persist the edited text as draftText ───────────────────────────────
    await db.contracts.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "draftText": plain_text_fallback,
            "draftHtml": content if looks_like_html else None
        }}
    )

    try:
        docx_info = await _create_docx_version(
            contract_id,
            plain_text_fallback,
            user_label,
            base_docx_path=base_docx_path,
            rendered_html=content if looks_like_html else None,
            paragraph_texts=paragraph_texts
        )
        docx_generated = True

    except Exception as e:
        logger.error(f"DOCX generation failed for {contract_id}: {e}")
        docx_generated = False
        docx_info = None

    await log_action(
        "Edit Contract", user_label, current_user.get("role", "User"),
        f"Saved edited content{' + generated DOCX' if docx_generated else ' (DOCX generation failed)'}",
        contract_id
    )

    if docx_generated:
        return {
            "message": "Content saved and DOCX generated successfully",
            "docx_generated": True,
            "document_id": docx_info["document_id"] if docx_info else None
        }
    else:
        return {
            "message": "Content saved to database. DOCX generation failed.",
            "docx_generated": False
        }


@app.post("/api/contracts/{contract_id}/ensure-docx")
async def ensure_docx_for_review(contract_id: str, current_user: dict = Depends(get_current_user)):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    existing_docx = await db.documents.find_one(
        {
            "contractId": contract_id,
            "$or": [
                {"fileType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
                {"fileName": {"$regex": r"\.docx$", "$options": "i"}}
            ]
        },
        sort=[("version", -1)]
    )
    if existing_docx:
        return {"document_id": str(existing_docx["_id"]), "created": False}

    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    user_label = current_user.get("email") or "System"
    try:
        if (contract.get("template_type") or "").upper() == "NDA":
            nda_docx_path = generate_nda_docx(
                contract_id,
                contract.get("company", "Unknown"),
                contract.get("description", ""),
                contract.get("expiry_date", "")
            )
            await db.contracts.update_one(
                {"_id": ObjectId(contract_id)},
                {"$set": {"latest_docx_path": nda_docx_path, "docx_version": int(contract.get("docx_version", 0)) + 1}}
            )
            path_obj = pathlib.Path(nda_docx_path)
            last_doc = await db.documents.find_one({"contractId": contract_id}, sort=[("version", -1)])
            version = (last_doc["version"] + 1) if last_doc else 1
            rec = Document(
                contractId=contract_id,
                fileName=path_obj.name,
                fileType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                fileSize=path_obj.stat().st_size if path_obj.exists() else 0,
                storagePath=nda_docx_path,
                version=version,
                uploadedBy=user_label,
                category="Original"
            )
            inserted = await db.documents.insert_one(rec.model_dump(by_alias=True, exclude_none=True))
            return {"document_id": str(inserted.inserted_id), "created": True}

        seed_content = (
            contract.get("draftText")
            or contract.get("extractedText")
            or contract.get("description")
            or contract.get("title")
            or "Contract content not available."
        )
        docx_info = await _create_docx_version(contract_id, seed_content, user_label, category_hint="Original")
        return {"document_id": docx_info["document_id"], "created": True}
    except Exception as e:
        logger.error(f"Failed to ensure DOCX for contract {contract_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate DOCX for review")

>>>>>>> Stashed changes
@app.post("/api/cas/{cas_id}/approve")
async def approve_cas(cas_id: str, action: str = Body(...)):
    # Action can be Approve or Reject
    cas = await db.cas.find_one({"_id": ObjectId(cas_id)})
    if not cas:
        raise HTTPException(status_code=404, detail="CAS not found")
        
    new_status = "Approved" if action == "Approve" else "Rejected"
    await db.cas.update_one(
        {"_id": ObjectId(cas_id)},
        {"$set": {"status": new_status}}
    )
    
    contract_id = cas.get("contractId")
    if contract_id and ObjectId.is_valid(contract_id):
        contract_status = "Approved" if action == "Approve" else "Rejected"
        contract_stage = "Approved" if action == "Approve" else "Under Review"
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": {
                "status": contract_status,
                "stage": contract_stage
            }}
        )
        
    return {"message": f"CAS {new_status}"}

@app.post("/api/cas/{cas_id}/approve-step")
async def approve_cas_step(cas_id: str, data: dict = Body(...)):
    step_index = data.get("stepIndex")
    approved_by = data.get("approvedBy", "Admin")
    timestamp = datetime.utcnow().isoformat()

    if not ObjectId.is_valid(cas_id):
        raise HTTPException(400, "Invalid CAS ID")

    cas = await db.cas.find_one({"_id": ObjectId(cas_id)})
    if not cas:
        raise HTTPException(404, "CAS not found")
    chain = cas.get("approvalChain", [])
    if step_index is None or step_index < 0 or step_index >= len(chain):
        raise HTTPException(400, "Invalid approval step")

    await db.cas.update_one(
        {"_id": ObjectId(cas_id)},
        {"$set": {
            f"approvalChain.{step_index}.status": "Approved",
            f"approvalChain.{step_index}.approvedBy": approved_by,
            f"approvalChain.{step_index}.timestamp": timestamp
        }}
    )
    
    step_role = chain[step_index].get("role", "Unknown")
    await log_audit(
        cas.get("contractId"),
        "DOA Approved",
        approved_by,
        role=step_role,
        department=step_role,
        details=f"DOA Step '{step_role}' approved by {approved_by}.",
        event_type=AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"],
        message=f"{step_role} DOA step approved.",
        metadata={"step": step_role, "approvedBy": approved_by},
    )

    cas = await db.cas.find_one({"_id": ObjectId(cas_id)})
    chain = cas.get("approvalChain", [])
    if all(s.get("status") == "Approved" for s in chain):
        await db.cas.update_one(
            {"_id": ObjectId(cas_id)},
            {"$set": {"status": "Approved"}}
        )
        if cas.get("contractId") and ObjectId.is_valid(cas.get("contractId")):
            await db.contracts.update_one(
                {"_id": ObjectId(cas.get("contractId"))},
                {"$set": {"status": "Completed", "stage": "Completed", "completedAt": datetime.utcnow().isoformat()}}
            )
            await log_audit(
                cas.get("contractId"),
                "DOA Completed",
                "System",
                role="System",
                details="All DOA approval steps completed.",
                event_type=AUDIT_EVENT_TYPES["DOA_COMPLETED"],
                message="All DOA approval steps completed.",
                metadata={"steps": [step.get("role") for step in chain]},
            )
            await log_audit(
                cas.get("contractId"),
                "Completed",
                "System",
                role="System",
                details="Final DOA approval completed. Contract is now fully completed.",
                event_type=AUDIT_EVENT_TYPES["CONTRACT_COMPLETED"],
                message="Contract completed.",
            )
            
            # Email notification for Final Approval
            try:
                initiator_name = cas.get("initiator", "")
                if initiator_name:
                    user_doc = await db.users.find_one({"name": initiator_name})
                    if user_doc and user_doc.get("email"):
                        await send_email_notification(
                            to_email=user_doc["email"],
                            subject=f"[Apeiro CLM] Final Approval: {cas.get('contractTitle', '')}",
                            body_html=f"""
                            <div style="font-family: sans-serif; max-width: 600px; padding: 20px; border: 1px solid #10b981; border-radius: 10px;">
                                <h2 style="color: #10b981;">Execution Ready!</h2>
                                <p>Hi {initiator_name},</p>
                                <p>Great news! The contract <strong>{cas.get('contractTitle', '')}</strong> has received final DOA approval.</p>
                                <p>The contract status is now <strong>Approved</strong> and the CAS is signed.</p>
                                <hr style="border: none; border-top: 1px solid #eee; margin: 20px 0;" />
                                <p style="font-size: 12px; color: #666;">This is an automated message from Apeiro CLM.</p>
                            </div>
                            """
                        )
            except Exception as e:
                print(f"Final approval email failed: {e}")
    return {"message": "Step approved successfully"}

@app.post("/api/doa/approve")
@app.post("/api/contracts/doa/{contract_id}/{action}")
async def doa_approve(
    data: dict = Body(None), 
    contract_id: Optional[str] = None, 
    action: Optional[str] = None
):
    # Handle both Body data and Path params for legacy support
    if not data:
        data = {}
    
    c_id = contract_id or data.get("contract_id")
    role = data.get("role")  # initiator, endorser, reviewer, approver
    act = action or data.get("action")  # approve, reject
    comments = data.get("comments", "")
    approved_by = data.get("approved_by", "Admin")
    
    if not c_id or not ObjectId.is_valid(c_id):
        raise HTTPException(400, "Invalid contract ID")
        
    # 1. Find CAS for this contract
    cas = await db.cas.find_one({"contractId": c_id})
    if not cas:
        # If no CAS, fallback to legacy direct approval for basic flows
        if act.lower() == "approve":
            await db.contracts.update_one({"_id": ObjectId(c_id)}, {"$set": {"status": "Approved", "stage": "Approved"}})
            await log_audit(c_id, "Approved", approved_by, role="Admin", details="Direct approval (No CAS found)")
            return {"message": "Direct approval successful"}
        else:
            await db.contracts.update_one({"_id": ObjectId(c_id)}, {"$set": {"status": "Rejected", "stage": "Rejected", "department": "Sales", "rejection_reason": comments or "Rejected during DOA", "rejected_by": approved_by}})
            await log_audit(c_id, "Rejected", approved_by, role="Admin", details="Direct rejection (No CAS found)")
            return {"message": "Direct rejection successful"}

    chain = cas.get("approvalChain", [])
    
    # 2. Find the current pending step
    current_step_index = -1
    for i, step in enumerate(chain):
        if step.get("status") == "Pending":
            current_step_index = i
            break
            
    if current_step_index == -1:
        raise HTTPException(400, "DOA workflow already completed or not initialized")
        
    current_step = chain[current_step_index]
    
    # 3. Check if role matches current step (Strict Enforcement)
    if role and current_step.get("role").lower() != role.lower():
        expected_role = current_step.get("role")
        raise HTTPException(400, f"Waiting for {expected_role} approval. Cannot approve as {role}.")
        
    # 4. Handle Action
    if act.lower() == "approve":
        timestamp = datetime.utcnow().isoformat()
        await db.cas.update_one(
            {"_id": cas["_id"]},
            {"$set": {
                f"approvalChain.{current_step_index}.status": "Approved",
                f"approvalChain.{current_step_index}.approvedBy": approved_by,
                f"approvalChain.{current_step_index}.timestamp": timestamp,
                f"approvalChain.{current_step_index}.comments": comments
            }}
        )
        
        # Advance doa_stage to next or clear if last
        next_role = None
        if current_step_index < len(chain) - 1:
            next_role = chain[current_step_index + 1].get("role")
            
        await db.cas.update_one({"_id": cas["_id"]}, {"$set": {"doa_stage": next_role, "doa_status": "Pending" if next_role else "Approved"}})
        await db.contracts.update_one({"_id": ObjectId(c_id)}, {"$set": {"doa_stage": next_role, "doa_status": "Pending" if next_role else "Approved"}})
        
        await log_audit(
            c_id,
            "DOA Approved",
            approved_by,
            role=current_step.get("role"),
            department=current_step.get("role"),
            details=f"DOA Step approved. Comments: {comments}",
            notes=comments,
            event_type=AUDIT_EVENT_TYPES["DOA_STEP_APPROVED"],
            message=f"{current_step.get('role')} DOA step approved.",
            metadata={"step": current_step.get("role"), "approvedBy": approved_by, "comments": comments},
        )
        
        # 5. Final Approval if last step
        if current_step_index == len(chain) - 1:
            await db.cas.update_one({"_id": cas["_id"]}, {"$set": {"status": "Completed"}})
            await db.contracts.update_one(
                {"_id": ObjectId(c_id)},
                {"$set": {"status": "Completed", "stage": "Completed"}}
            )
            await log_audit(
                c_id,
                "DOA Completed",
                "System",
                role="System",
                details="All DOA approval steps completed.",
                event_type=AUDIT_EVENT_TYPES["DOA_COMPLETED"],
                message="All DOA approval steps completed.",
                metadata={"steps": [step.get("role") for step in chain]},
            )
            await log_audit(
                c_id,
                "Completed",
                "System",
                role="System",
                details="Final DOA approval completed. Contract is now fully completed.",
                event_type=AUDIT_EVENT_TYPES["CONTRACT_COMPLETED"],
                message="Contract completed.",
            )
    else:
        # Rejection
        await db.cas.update_one(
            {"_id": cas["_id"]},
            {"$set": {
                f"approvalChain.{current_step_index}.status": "Rejected",
                f"approvalChain.{current_step_index}.approvedBy": approved_by,
                f"approvalChain.{current_step_index}.timestamp": datetime.utcnow().isoformat(),
                f"approvalChain.{current_step_index}.comments": comments,
                "status": "Rejected",
                "doa_status": "Rejected"
            }}
        )
        await db.contracts.update_one(
            {"_id": ObjectId(c_id)},
            {"$set": {"status": "Rejected", "stage": "Rejected", "department": "Sales", "doa_status": "Rejected", "rejection_reason": comments or "Rejected during DOA", "rejected_by": approved_by}}
        )
        await log_audit(c_id, "Rejected", approved_by, role=current_step.get("role"), department=current_step.get("role"), details=f"DOA Step rejected. Comments: {comments}")

    return {"message": f"Step {act}d successfully"}

# ─── Notification Endpoints ───

@app.get("/api/notifications")
async def get_notifications(role: Optional[str] = Query(None)):
    query = {"read": False}
    if role:
        query["for_role"] = role
    cursor = db.notifications.find(query).sort("createdAt", -1)
    notifs = await cursor.to_list(length=100)
    for n in notifs:
        n["id"] = str(n["_id"])
        del n["_id"]
    return notifs

@app.post("/api/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: str, role: Optional[str] = Query(None)):
    if not ObjectId.is_valid(notif_id):
        raise HTTPException(status_code=400, detail="Invalid notification ID")
    await db.notifications.update_one(
        {"_id": ObjectId(notif_id)},
        {"$set": {"read": True}}
    )
    # Return updated unread list for the role
    query = {"read": False}
    if role:
        query["for_role"] = role
    cursor = db.notifications.find(query).sort("createdAt", -1)
    notifs = await cursor.to_list(length=100)
    for n in notifs:
        n["id"] = str(n["_id"])
        del n["_id"]
    return notifs

@app.get("/api/audit/{contract_id}")
async def get_audit_logs(contract_id: str):
    cursor = db.audit_logs.find({"contractId": contract_id}).sort("timestamp", -1)
    logs = await cursor.to_list(length=100)
    return _curate_audit_logs(logs)

# ─── Admin Endpoints ───

@app.delete("/api/admin/clear-all")
async def clear_all():
    await db.contracts.delete_many({})
    await db.cas.delete_many({})
    await db.notifications.delete_many({})
    return {"message": "All data cleared"}

@app.delete("/api/admin/clear-notifications")
async def clear_notifs():
    await db.notifications.delete_many({})
    return {"message": "Notifications cleared"}

@app.delete("/api/admin/clear-cas")
async def clear_cas():
    await db.cas.delete_many({})
    return {"message": "CAS records cleared"}


@app.post("/api/doa/rules")
async def create_doa_rule(data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    rule = DOARule(
        name=data.get("name", "").strip() or "Untitled Rule",
        conditions=data.get("conditions", {}),
        approvers=[str(item).strip() for item in data.get("approvers", []) if str(item).strip()],
        priority=int(data.get("priority", 100)),
        is_active=bool(data.get("is_active", True)),
    )
    result = await db.doa_rules.insert_one(rule.model_dump(by_alias=True, exclude_none=True))
    created = await db.doa_rules.find_one({"_id": result.inserted_id})
    await db.doa_rule_history.insert_one(
        DOARuleHistory(
            rule_id=str(result.inserted_id),
            updated_by=user_name,
            previous_rule={},
            new_rule={**rule.model_dump(by_alias=True, exclude_none=True), "_id": str(result.inserted_id)},
            action="created",
        ).model_dump(by_alias=True, exclude_none=True)
    )
    created["id"] = str(created["_id"])
    del created["_id"]
    return created


@app.get("/api/doa/rules")
async def list_doa_rules(include_inactive: bool = False):
    query = {} if include_inactive else {"is_active": True}
    cursor = db.doa_rules.find(query).sort([("priority", 1), ("updated_at", -1)])
    rules = await cursor.to_list(length=500)
    for rule in rules:
        rule["id"] = str(rule["_id"])
        del rule["_id"]
    return rules


@app.put("/api/doa/rules/{rule_id}")
async def update_doa_rule(rule_id: str, data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    if not ObjectId.is_valid(rule_id):
        raise HTTPException(400, "Invalid rule ID")
    existing = await db.doa_rules.find_one({"_id": ObjectId(rule_id)})
    if not existing:
        raise HTTPException(404, "DOA rule not found")

    update_fields = {
        "updated_at": datetime.utcnow().isoformat(),
    }
    if "name" in data:
        update_fields["name"] = data["name"].strip() or existing.get("name", "Untitled Rule")
    if "conditions" in data:
        update_fields["conditions"] = data["conditions"]
    if "approvers" in data:
        update_fields["approvers"] = [str(item).strip() for item in data.get("approvers", []) if str(item).strip()]
    if "priority" in data:
        update_fields["priority"] = int(data["priority"])
    if "is_active" in data:
        update_fields["is_active"] = bool(data["is_active"])

    await db.doa_rules.update_one({"_id": ObjectId(rule_id)}, {"$set": update_fields})
    updated = await db.doa_rules.find_one({"_id": ObjectId(rule_id)})
    await db.doa_rule_history.insert_one(
        DOARuleHistory(
            rule_id=rule_id,
            updated_by=user_name,
            previous_rule={**existing, "_id": str(existing["_id"])},
            new_rule={**updated, "_id": str(updated["_id"])},
            action="updated",
        ).model_dump(by_alias=True, exclude_none=True)
    )
    updated["id"] = str(updated["_id"])
    del updated["_id"]
    return updated


@app.delete("/api/doa/rules/{rule_id}")
async def delete_doa_rule(rule_id: str, authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    if not ObjectId.is_valid(rule_id):
        raise HTTPException(400, "Invalid rule ID")
    existing = await db.doa_rules.find_one({"_id": ObjectId(rule_id)})
    if not existing:
        raise HTTPException(404, "DOA rule not found")
    await db.doa_rules.update_one(
        {"_id": ObjectId(rule_id)},
        {"$set": {"is_active": False, "updated_at": datetime.utcnow().isoformat()}}
    )
    updated = await db.doa_rules.find_one({"_id": ObjectId(rule_id)})
    await db.doa_rule_history.insert_one(
        DOARuleHistory(
            rule_id=rule_id,
            updated_by=user_name,
            previous_rule={**existing, "_id": str(existing["_id"])},
            new_rule={**updated, "_id": str(updated["_id"])},
            action="deleted",
        ).model_dump(by_alias=True, exclude_none=True)
    )
    return {"message": "DOA rule deactivated successfully"}


@app.get("/api/doa/rules/history")
async def list_doa_rule_history():
    cursor = db.doa_rule_history.find({}).sort("updated_at", -1)
    history = await cursor.to_list(length=1000)
    for item in history:
        item["id"] = str(item["_id"])
        del item["_id"]
    return history

@app.get("/api/admin/doa-thresholds")
async def get_doa_thresholds():
    settings_doc = await db.settings.find_one({"key": "doa_thresholds"})
    if settings_doc:
        return settings_doc.get("value")
    return {"Manager": {"min": 0, "max": 10000}, "CEO": {"min": 10001, "max": 0}}

@app.put("/api/admin/doa-thresholds")
async def update_doa(thresholds: dict = Body(...)):
    await db.settings.update_one(
        {"key": "doa_thresholds"},
        {"$set": {"value": thresholds}},
        upsert=True
    )
    return {"message": "DOA thresholds updated"}

@app.post("/api/admin/run-lifecycle-check")
async def manual_lifecycle_check():
    await check_expiring_contracts()
    return {"message": "Lifecycle check completed"}

@app.get("/api/admin/users")
async def get_users():
    users = await db.users.find({}).to_list(100)
    for u in users:
        u["id"] = str(u["_id"])
        del u["_id"]
        u.pop("password", None)  # Never send password
    return users

@app.post("/api/admin/users")
async def add_user(user: dict = Body(...)):
    existing = await db.users.find_one({"email": user.get("email")})
    if existing:
        raise HTTPException(400, "User already exists")
    user["status"] = "Active"
    await db.users.insert_one(user)
    return {"message": "User added successfully"}

@app.put("/api/admin/users/{user_id}/role")
async def update_role(user_id: str, data: dict = Body(...)):
    await db.users.update_one(
        {"_id": ObjectId(user_id)},
        {"$set": {"role": data.get("role")}}
    )
    return {"message": "Role updated"}

@app.put("/api/admin/users/{user_id}/status")
async def update_status(user_id: str, data: dict = Body(...)):
    await db.users.update_one(
        {"_id": ObjectId(user_id)},
        {"$set": {"status": data.get("status")}}
    )
    return {"message": "Status updated"}

@app.delete("/api/admin/users/{user_id}")
async def delete_user(user_id: str):
    if not ObjectId.is_valid(user_id):
        raise HTTPException(400, "Invalid user ID")
    result = await db.users.delete_one({"_id": ObjectId(user_id)})
    if result.deleted_count == 0:
        raise HTTPException(404, "User not found")
    return {"message": "User deleted successfully"}

@app.get("/api/admin/users-with-auth")
async def get_users_auth():
    # Only for login validation - includes password
    users = await db.users.find({}).to_list(100)
    for u in users:
        u["id"] = str(u["_id"])
        del u["_id"]
    return users

@app.post("/api/ai/analyze-contract")
async def analyze_contract(data: dict = Body(...)):
    raw_text = data.get("document_text", "")
    structured_data = {}
    try:
        structured_data = json.loads(raw_text) if isinstance(raw_text, str) else {}
    except Exception:
        structured_data = {}

    result = await ai_service.analyze_risk(
        text=data.get("document_text", ""),
        contract_type=data.get("contract_type", "General"),
        jurisdiction=data.get("jurisdiction", "International"),
        value=data.get("value", "Unknown"),
        duration=data.get("duration", "Unknown")
    )

    if not isinstance(result, dict):
        result = {}

    extracted_clauses = structured_data.get("clauses") or []
    ai_extracted_clauses = result.get("extractedClauses") or {}
    key_financial_terms = result.get("keyFinancialTerms") or {}

    def pick_clause(*keywords):
        for clause in extracted_clauses:
            haystack = f"{clause.get('title', '')} {clause.get('type', '')} {clause.get('content', '')}".lower()
            if any(keyword in haystack for keyword in keywords):
                return clause.get("content") or clause.get("text") or clause.get("title")
        return ""

    def normalize_clause_excerpt(value):
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, dict):
            for key in ["content", "text", "value", "summary", "excerpt", "clause"]:
                nested_value = value.get(key)
                if isinstance(nested_value, str) and nested_value.strip():
                    return nested_value.strip()
        if isinstance(value, list):
            parts = [normalize_clause_excerpt(item) for item in value]
            parts = [part for part in parts if part]
            return " ".join(parts[:2]).strip()
        return str(value).strip()

    def first_available(*values):
        for value in values:
            normalized_value = normalize_clause_excerpt(value)
            if normalized_value:
                return normalized_value
        return ""

    contract_type = (
        result.get("contractType")
        or structured_data.get("category")
        or data.get("contract_type")
        or "Not Available"
    )

    try:
        value_num = _parse_contract_value(data.get("value"))
    except Exception:
        value_num = 0.0

    suggested_department = result.get("suggestedDepartment")
    if not suggested_department:
        try:
            workflow = await calculate_review_stages(
                structured_data.get("title") or "Untitled",
                value_num,
                contract_type if contract_type != "Not Available" else "General",
                result.get("overallRiskScore") or "Medium",
                extracted_clauses
            )
            suggested_department = workflow[0] if workflow else None
        except Exception:
            suggested_department = None

    key_clauses = {
        "payment_terms": first_available(
            ai_extracted_clauses.get("payment_terms"),
            ai_extracted_clauses.get("paymentTerms"),
            ai_extracted_clauses.get("payment"),
            key_financial_terms.get("paymentTerms"),
            key_financial_terms.get("payment_terms"),
            pick_clause("payment", "invoice", "fee", "pricing"),
        ),
        "termination": first_available(
            ai_extracted_clauses.get("termination"),
            ai_extracted_clauses.get("termination_clause"),
            ai_extracted_clauses.get("terminationClause"),
            pick_clause("termination", "terminate", "expiry", "expiration"),
        ),
        "confidentiality": first_available(
            ai_extracted_clauses.get("confidentiality"),
            ai_extracted_clauses.get("confidentiality_clause"),
            ai_extracted_clauses.get("confidentialityClause"),
            ai_extracted_clauses.get("nda"),
            pick_clause("confidentiality", "confidential", "non-disclosure", "nda"),
        ),
        "liability": first_available(
            ai_extracted_clauses.get("liability"),
            ai_extracted_clauses.get("liability_clause"),
            ai_extracted_clauses.get("liabilityClause"),
            ai_extracted_clauses.get("indemnification"),
            ai_extracted_clauses.get("indemnity"),
            pick_clause("liability", "indemn", "damages"),
        ),
    }

    result["contractType"] = contract_type
    result["suggestedDepartment"] = suggested_department or "Not Available"
    result["riskScore"] = result.get("riskPercentage", 0)
    result["missingClauses"] = result.get("missingClauses") or []
    result["criticalIssues"] = result.get("criticalIssues") or []
    result["extractedClauses"] = {
        **(ai_extracted_clauses if isinstance(ai_extracted_clauses, dict) else {}),
        **key_clauses,
    }
    result["keyClauses"] = key_clauses
    result["clauseHighlights"] = [
        {"label": "Payment Terms", "value": key_clauses["payment_terms"] or "Not Available"},
        {"label": "Liability", "value": key_clauses["liability"] or "Not Available"},
        {"label": "Termination", "value": key_clauses["termination"] or "Not Available"},
    ]
    return result

@app.post("/api/ai/extract-text")
async def extract_text_endpoint(data: dict = Body(...)):
    result = await ai_service.extract_text(
        text=data.get("document_text", "")
    )
    return result

@app.post("/api/ai/check-compliance")
async def check_compliance_endpoint(data: dict = Body(...)):
    result = await ai_service.check_compliance(
        text=data.get("document_text", ""),
        regulations=data.get("regulations", "GDPR, HIPAA, SOX, CCPA"),
        industry=data.get("industry", "Technology"),
        jurisdiction=data.get("jurisdiction", "International")
    )
    return result

@app.post("/api/ai/route-review")
async def route_review_endpoint(data: dict = Body(...)):
    result = await ai_service.route_review(
        title=data.get("title", "Untitled"),
        value=data.get("value", 0),
        contract_type=data.get("contract_type", "General"),
        jurisdiction=data.get("jurisdiction", "International"),
        risk=data.get("risk", "Medium")
    )
    return result

@app.post("/api/ai/redline-clause")
async def redline_clause_endpoint(data: dict = Body(...)):
    result = await ai_service.redline_clause(
        clause=data.get("clause", ""),
        section=data.get("section", ""),
        issue=data.get("issue", ""),
        company=data.get("company", "Our Company"),
        role=data.get("role", "Vendor"),
        risk_tolerance=data.get("risk_tolerance", "Medium"),
        industry=data.get("industry", "Technology")
    )
    return result

<<<<<<< Updated upstream
@app.post("/api/ai/extract-email")
async def extract_email(data: dict = Body(...)):
    # Handle both frontend 'email_text' and potential backend 'body'
    body_content = data.get("body") or data.get("email_text") or ""
    
    result = await ai_service.extract_email(
        email_from=data.get("email_from", ""),
        subject=data.get("subject", ""),
        date=data.get("date", ""),
        body=body_content
    )
    
    # Flatten the response for frontend compatibility
    # The AIService returns { contractInfo: {...}, dates: {...}, ... }
    # The frontend expects { counterpartyName, contractValue, subject, dates, ... }
    flattened = {
        **(result.get("contractInfo", {})),
        "dates": result.get("dates", {}).get("proposedStartDate", "") + " - " + result.get("dates", {}).get("proposedEndDate", ""),
        "keyRequirements": result.get("keyRequirements", []),
        "sentiment": result.get("sentiment"),
        "extractionConfidence": result.get("extractionConfidence", 0)
    }
    
    # Ensure critical fields expected by 35: if (!data.counterpartyName && !data.contractValue)
    # exist at the root
    if "counterpartyName" not in flattened:
        flattened["counterpartyName"] = result.get("contractInfo", {}).get("counterpartyName", "")
    if "contractValue" not in flattened:
        flattened["contractValue"] = result.get("contractInfo", {}).get("contractValue", "")
        
    return flattened

@app.post("/api/ai/generate-notification")
async def generate_notification_endpoint(data: dict = Body(...)):
    result = await ai_service.generate_notification(
        event=data.get("event", ""),
        contract_title=data.get("contract_title", ""),
        context=data.get("context", "")
    )
    return result

@app.post("/api/ai/track-lifecycle")
async def track_lifecycle_endpoint(data: dict = Body(...)):
    result = await ai_service.track_lifecycle(
        title=data.get("title", ""),
        start_date=data.get("start_date", ""),
        end_date=data.get("end_date", ""),
        renewal_clause=data.get("renewal_clause", ""),
        notice_days=data.get("notice_days", 0),
        value=data.get("value", 0),
        status=data.get("status", "Active")
    )
    return result

=======
>>>>>>> Stashed changes
@app.post("/api/ai/generate-cas-notes")
async def generate_cas_notes(data: dict = Body(...)):
    contract_id = data.get("contract_id", "")
    contract = await db.contracts.find_one(
        {"_id": ObjectId(contract_id)}
    ) if ObjectId.is_valid(contract_id) else None
    
    if not contract:
        reviews = {}
        title = company = business_unit = department = risk_summary = key_issues = initiator_notes = "Unknown"
        value = 0
    else:
        reviews = contract.get("reviews", {})
        title = contract.get("title", "Untitled")
        company = contract.get("company", "Unknown")
        value = contract.get("value", 0)
        contract_type = contract.get("category", "General")
        business_unit = contract.get("business_unit") or contract.get("businessUnit") or ""
        department = contract.get("department", "Legal")
        risk_summary = contract.get("risk_classification", "Medium Risk")
        key_issues = "Standard terms reviewed."
        initiator_notes = contract.get("submittedBy", "Admin")

    key_notes = await ai_service.generate_cas_notes(
        title=title,
        company=company,
        value=value,
        contract_type=data.get("contract_type", contract_type if contract else "General"),
        business_unit=business_unit,
        department=department,
        risk_summary=risk_summary,
        reviews=reviews,
        key_issues=key_issues,
        initiator_notes=initiator_notes
    )
    
    if ObjectId.is_valid(contract_id):
        await db.cas.update_one(
            {"contractId": contract_id},
            {"$set": {"keyNotes": key_notes}}
        )
    
    return {"key_notes": key_notes}

@app.post("/api/ai/search-contracts")
async def search_contracts_endpoint(data: dict = Body(...)):
    contracts = data.get("contracts")
    if not contracts:
        # Fetch from DB if not provided
        cursor = db.contracts.find({})
        db_contracts = await cursor.to_list(length=100)
        contracts = "\n".join([f"ID: {str(c['_id'])} | Title: {c.get('title')} | Company: {c.get('company')} | Value: ${c.get('value', 0)}" for c in db_contracts])
    
    result = await ai_service.search_contracts(
        query=data.get("query", ""),
        search_type=data.get("search_type", "Semantic"),
        contracts_list=contracts
    )
    return result

@app.post("/api/contracts/{contract_id}/documents")
async def upload_document(
    contract_id: str,
    file: UploadFile = File(...),
    uploadedBy: str = Body("Admin"),
    category: str = Body("General"),
    tags: str = Body(""),
):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    # Check file size
    contents = await file.read()
    size_mb = len(contents) / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise HTTPException(status_code=413, detail=f"File too large. Max {MAX_FILE_SIZE_MB}MB allowed.")

    # Determine version
    existing_count = await db.documents.count_documents({"contractId": contract_id})
    version = existing_count + 1

    # Save file to disk
    contract_dir = UPLOAD_DIR / contract_id
    contract_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"v{version}_{file.filename}"
    file_path = contract_dir / safe_name
    with open(file_path, "wb") as f:
        f.write(contents)

    # Save metadata to MongoDB
    tag_list = [t.strip() for t in tags.split(",") if t.strip()] if tags else []
    doc = Document(
        contractId=contract_id,
        fileName=file.filename,
        fileType=file.content_type or "application/octet-stream",
        fileSize=len(contents),
        storagePath=str(file_path),
        version=version,
        uploadedBy=uploadedBy,
        category=category,
        tags=tag_list,
    )
    result = await db.documents.insert_one(
        doc.model_dump(by_alias=True, exclude_none=True)
    )
    return {
        "message": "File uploaded successfully",
        "documentId": str(result.inserted_id),
        "version": version,
        "fileName": file.filename,
        "fileSize": len(contents),
    }

@app.get("/api/contracts/{contract_id}/documents")
async def list_documents(contract_id: str):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
    cursor = db.documents.find({"contractId": contract_id}).sort("uploadedAt", -1)
    docs = await cursor.to_list(length=100)
    for d in docs:
        d["id"] = str(d["_id"])
        del d["_id"]
    return docs

@app.get("/api/documents/{document_id}/download")
async def download_document(document_id: str):
    if not ObjectId.is_valid(document_id):
        raise HTTPException(status_code=400, detail="Invalid document ID")
    doc = await db.documents.find_one({"_id": ObjectId(document_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    file_path = pathlib.Path(doc["storagePath"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    return FileResponse(
        path=str(file_path),
        filename=doc["fileName"],
        media_type=doc["fileType"],
    )

<<<<<<< Updated upstream
=======
@app.get("/api/documents/{document_id}/view")
async def get_document_view_data(document_id: str):
    if not ObjectId.is_valid(document_id):
        raise HTTPException(status_code=400, detail="Invalid document ID")
    doc = await db.documents.find_one({"_id": ObjectId(document_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    file_path = pathlib.Path(doc["storagePath"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")
    
    file_extension = pathlib.Path(doc["fileName"]).suffix.lower()
    
    if file_extension == ".pdf":
        return {"type": "pdf", "url": f"/api/documents/{document_id}/download"}
    
    elif file_extension == ".docx":
        if doc.get("renderedHtml"):
            return {
                "type": "docx",
                "render_mode": "html",
                "content": doc.get("renderedHtml")
            }
        return {
            "type": "docx",
            "render_mode": "docx-preview",
            "url": f"/api/documents/{document_id}/download"
        }
    
    else:
        # Fallback for text files or unknown
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return {"type": "text", "content": content}
        except Exception:
            return {"type": "unsupported", "message": "This file type cannot be viewed in-app."}

@app.get("/api/contracts/{contract_id}/comments")
async def get_contract_comments(contract_id: str, department: Optional[str] = None):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
    
    query = {"contractId": contract_id}
    if department:
        query["department"] = department
        
    cursor = db.comments.find(query).sort("createdAt", 1)
    comments = await cursor.to_list(length=100)
    for c in comments:
        c["id"] = str(c["_id"])
        del c["_id"]
    return comments

@app.post("/api/contracts/{contract_id}/comments")
async def add_contract_comment(contract_id: str, comment_data: ReviewComment):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")
    
    comment_dict = comment_data.model_dump(by_alias=True, exclude_none=True)
    comment_dict["contractId"] = contract_id
    
    result = await db.comments.insert_one(comment_dict)
    return {
        "message": "Comment added successfully",
        "commentId": str(result.inserted_id)
    }

>>>>>>> Stashed changes
@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: str):
    if not ObjectId.is_valid(document_id):
        raise HTTPException(status_code=400, detail="Invalid document ID")
    doc = await db.documents.find_one({"_id": ObjectId(document_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    # Delete file from disk
    file_path = pathlib.Path(doc["storagePath"])
    if file_path.exists():
        file_path.unlink()
    await db.documents.delete_one({"_id": ObjectId(document_id)})
    return {"message": "Document deleted successfully"}

<<<<<<< Updated upstream
@app.post("/api/ai/generate-draft")
async def generate_draft(data: dict = Body(...)):
    try:
        prompt = f"""Generate a professional contract draft for:
Title: {data.get('title')}
Counterparty: {data.get('counterpartyName')}
Value: {data.get('contractValue')}
Duration: {data.get('duration')}
Business Unit: {data.get('businessUnit')}
Category: {data.get('category')}
Risk Level: {data.get('riskLevel')}

Write a professional contract with sections:
1. Parties
2. Scope of Work
3. Consideration and Payment Terms
4. Duration
5. Risk Assessment
6. Liability and Indemnity
7. Termination
8. Confidentiality

Use markdown # for title and ## for sections.
Keep it professional and concise."""

        response = deepseek_client.chat.completions.create(
            model=DEEPSEEK_MODEL,
            messages=[{
                "role": "system",
                "content": "You are a professional contract writer. Write clear, professional contract drafts in markdown format."
            }, {
                "role": "user",
                "content": prompt
            }],
            max_tokens=1500
        )
        draft = response.choices[0].message.content.strip()
        return {"draft": draft}
    except Exception as e:
        return {
            "draft": f"# {data.get('title', 'Contract')}\n\n## Standard Terms\nStandard terms applied."
        }

@app.post("/api/contracts/save-draft")
async def save_draft(data: dict = Body(...), authorization: Optional[str] = Header(None)):
    user_name = get_current_user(authorization)
    try:
        normalized_input = _normalize_contract_input(data)
        contract = Contract(
            title=normalized_input["title"] or "Draft Contract",
            contract_name=normalized_input["title"] or "Draft Contract",
            company=normalized_input["company"] or "Unknown",
            value=normalized_input["value"],
            department=data.get("department") or normalized_input["business_unit"] or "Legal",
            stage="Draft",
            status="Draft",
            submittedBy=normalized_input["submittedBy"],
            created_by=normalized_input["created_by"],
            category=normalized_input["category"],
            risk_classification=normalized_input["risk_classification"],
            duration=normalized_input["duration"],
            business_unit=normalized_input["business_unit"],
            contract_owner=normalized_input["contract_owner"],
            sales_opportunity_id=normalized_input["sales_opportunity_id"] or None,
            opportunity_id=normalized_input["sales_opportunity_id"] or None,
            workflow=[],
            current_stage_index=0,
            startDate=data.get("startDate", ""),
            endDate=data.get("endDate", ""),
            expiryDate=data.get("expiryDate", "")
        )
        result = await db.contracts.insert_one(
            contract.model_dump(by_alias=True, exclude_none=True)
        )
        draft_text = normalized_input["draftText"]
        if draft_text:
            await db.contracts.update_one(
                {"_id": result.inserted_id},
                {"$set": {"draftText": draft_text}}
            )
        await log_audit(str(result.inserted_id), "Created", user_name, role="System", details=f"Draft contract '{contract.title}' saved by {user_name}", department="Draft")
        return {
            "message": "Draft saved successfully",
            "id": str(result.inserted_id)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/ai/extract-file")
async def extract_file(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        text = ""
        
        if file.filename.endswith('.pdf'):
            import PyPDF2
            from io import BytesIO
            pdf_reader = PyPDF2.PdfReader(BytesIO(contents))
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif file.filename.endswith('.docx'):
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(contents))
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            text = contents.decode('utf-8', errors='ignore')

        if not text.strip():
            return {
                "counterparty": "",
                "contractValue": "",
                "effectiveDate": "",
                "expiryDate": "",
                "duration": "",
                "keyDates": "",
                "clauses": [],
                "error": "No readable text was found in this file. If this is a scanned PDF, OCR is required before AI extraction can work.",
                "warnings": ["The uploaded document does not contain selectable text."],
                "extractionStatus": "empty_text"
            }

        result = await ai_service.extract_text(text)
        
        # Defensively handle if AI wraps the JSON object in a list
        if isinstance(result, list):
            result = result[0] if len(result) > 0 else {}
        elif not isinstance(result, dict):
            result = {}
            
        # Safely extract nested properties in case AI returns nulls instead of dicts
        parties = result.get("parties") or {}
        fin_terms = result.get("financialTerms") or {}
        dates = result.get("contractDates") or {}

        extracted_clauses = result.get("clauses", [])
        import uuid
        for idx, clause in enumerate(extracted_clauses):
            if not clause.get("id"):
                clause["id"] = f"cl-{idx}-{str(uuid.uuid4())[:8]}"
            clause["text"] = clause.get("content", "")  # required by frontend preview
            if not clause.get("title"):
                clause["title"] = clause.get("type", f"Clause {idx+1}")

        # Map to legacy format expected by FileUploadTab.jsx
        currency = fin_terms.get("currency") or ""
        total_value = fin_terms.get("totalValue")
        contract_value = ""
        if total_value not in [None, ""]:
            contract_value = f"{currency}{total_value}"

        legacy_mapped = {
            "counterparty": parties.get("vendor") or parties.get("client") or "",
            "counterparty_email": parties.get("email") or parties.get("contact") or "",
            "signer_name": parties.get("signer") or parties.get("representative") or "",
            "contractValue": contract_value,
            "duration": f"{dates.get('effectiveDate', '')} to {dates.get('expiryDate', '')}",
            "keyDates": f"Effective: {dates.get('effectiveDate', '')} | Expiry: {dates.get('expiryDate', '')}",
            "effectiveDate": dates.get('effectiveDate', ''),
            "expiryDate": dates.get('expiryDate', ''),
            "clauses": extracted_clauses,
            "full_data": result,
            "warnings": [],
            "extractionStatus": "ai"
        }
        
        # Clean placeholders so frontend validation triggers cleanly if actually empty
        if legacy_mapped["contractValue"] == "$":
            legacy_mapped["contractValue"] = ""

        has_ai_data = any([
            _clean_extracted_value(legacy_mapped["counterparty"]),
            _clean_extracted_value(legacy_mapped["contractValue"]),
            _clean_extracted_value(legacy_mapped["effectiveDate"]),
            _clean_extracted_value(legacy_mapped["expiryDate"]),
            len(legacy_mapped["clauses"]) > 0,
        ])

        if not has_ai_data:
            fallback = _extract_plain_text_fallback(text)
            has_fallback_data = any([
                _clean_extracted_value(fallback["counterparty"]),
                _clean_extracted_value(fallback["contractValue"]),
                _clean_extracted_value(fallback["effectiveDate"]),
                _clean_extracted_value(fallback["expiryDate"]),
                len(fallback["clauses"]) > 0,
            ])
            if has_fallback_data:
                fallback["full_data"] = result
                fallback["warnings"] = ["AI returned incomplete metadata, so plain-text fallback extraction was used."]
                fallback["extractionStatus"] = "fallback"
                return fallback
            legacy_mapped["warnings"] = ["AI could not confidently extract structured metadata from this file."]
            legacy_mapped["error"] = result.get("dataQuality", {}).get("warnings", ["AI extraction failed."])[0] if isinstance(result, dict) else "AI extraction failed."
            legacy_mapped["extractionStatus"] = "empty_ai_result"
            
        return legacy_mapped
    except Exception as e:
        print(f"File extraction error: {e}")
        # Return empty data so frontend safely displays the "AI could not extract data" error
        return {
            "counterparty": "",
            "contractValue": "",
            "error": str(e)
        }

# --- DigiInk Digital Signature Integration ---
=======
@app.post("/api/contracts/{contract_id}/send-for-signature")
async def send_for_signature(contract_id: str, current_user: dict = Depends(get_current_user)):
    print(f"DEBUG: Entering send_for_signature for contract_id: {contract_id}")
    if not ObjectId.is_valid(contract_id):
        print(f"DEBUG: Invalid contract_id: {contract_id}")
        raise HTTPException(status_code=400, detail="Invalid contract ID")
    
    # 1. Fetch Contract
    print(f"DEBUG: Fetching contract {contract_id} from DB...")
    contract = await db.contracts.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        print(f"DEBUG: Contract {contract_id} not found in DB")
        raise HTTPException(status_code=404, detail="Contract not found")
        
    # 2. Fetch Document
    print(f"DEBUG: Fetching latest document for contract {contract_id}")
    doc = await db.documents.find_one({"contractId": contract_id}, sort=[("uploadedAt", -1)])
    if not doc:
        print(f"DEBUG: No document found for contract {contract_id}")
        raise HTTPException(status_code=404, detail="No document found for this contract")
    
    # 2.1 Robust Path Resolution
    # storagePath is the field set by the upload endpoint
    db_path = doc.get("storagePath") or doc.get("file_path", "")
    print(f"DEBUG: DB Document storagePath: {db_path}")
    
    # Extract filename correctly regardless of it being a Windows or Linux path
    filename = PureWindowsPath(db_path).name if db_path else doc.get("fileName", "")
    print(f"DEBUG: Extracted filename: {filename}")
    
    # Search for the file in possible locations (only accept actual files, not directories)
    candidate_paths = [
        # 1. Exact stored path
        db_path,
        # 2. Inside contract-specific folder in uploads/ (Modern convention)
        os.path.join("uploads", contract_id, filename) if filename else None,
        # 3. Directly in uploads/
        os.path.join("uploads", filename) if filename else None,
        # 4. In uploaded_contracts/ (Legacy convention)
        os.path.join("uploaded_contracts", filename) if filename else None,
    ]
    
    file_path = None
    for path in candidate_paths:
        if path and os.path.isfile(path):  # isfile() rejects directories
            file_path = path
            print(f"DEBUG: Successfully resolved file to: {file_path}")
            break
    
    if not file_path:
        # Last resort: walk the contract upload dir and find any file
        contract_upload_dir = os.path.join("uploads", contract_id)
        if os.path.isdir(contract_upload_dir):
            for f in os.listdir(contract_upload_dir):
                candidate = os.path.join(contract_upload_dir, f)
                if os.path.isfile(candidate):
                    file_path = candidate
                    print(f"DEBUG: Found file by directory scan: {file_path}")
                    break
    
    if not file_path:
        print(f"DEBUG: File not found. candidates tried: {[p for p in candidate_paths if p]}")
        raise HTTPException(status_code=404, detail="Document file not found on disk")
        
    # 3. Prepare recipients
    initiator_email = contract.get("submittedBy", ADMIN_EMAIL)
    print(f"DEBUG: Initiator email: {initiator_email}, Admin email: {ADMIN_EMAIL}")
    
    # Simple whitelist check for initiator and sender
    try:
        await check_domain_whitelist(initiator_email)
        await check_domain_whitelist(ADMIN_EMAIL)
    except Exception as e:
        print(f"DEBUG: Whitelist check failed: {str(e)}")
        raise e

    recipients = [
        {"email": initiator_email, "name": "Initiator", "role": "signer"},
        {"email": ADMIN_EMAIL, "name": contract.get("company", "Enterprise Services"), "role": "signer"}
    ]
    
    try:
        print(f"DEBUG: Calling digink_service.create_document...")
        result = await digink_service.create_document(
            sender_email=ADMIN_EMAIL,
            title=contract.get("title", "Contract Signature Request"),
            file_path=file_path,
            recipients=recipients
        )
        
        # Save DigInk tracking info
        digink_id = result.get("document_id") or result.get("id")
        print(f"DEBUG: DigInk success! ID: {digink_id}")
        digink_doc = {
            "contractId": contract_id,
            "diginkDocumentId": digink_id,
            "status": "Sent",
            "createdAt": datetime.utcnow().isoformat(),
            "updatedAt": datetime.utcnow().isoformat()
        }
        await db.digink_documents.insert_one(digink_doc)
        
        # Update contract stage
        await db.contracts.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": {"stage": "Sent for Signature", "status": "Sent for Signature"}}
        )
        
        await log_action("Send for Signature", current_user.get("email", "Admin"), current_user.get("role", "User"), 
                         f"Contract '{contract.get('title')}' sent for signature via DigInk.", contract_id)
        
        print(f"DEBUG: Returning success for {contract_id}")
        return {"message": "Document sent for signature successfully", "digink_id": digink_id}
    except Exception as e:
        error_msg = str(e)
        print(f"DEBUG: DigInk Service Error: {error_msg}")
        await log_action("Signature Failure", current_user.get("email", "Admin"), current_user.get("role", "User"), 
                         f"Failed to send '{contract.get('title')}' for signature: {error_msg}", contract_id)
        raise HTTPException(status_code=500, detail=error_msg)

@app.post("/api/webhook/digink")
async def digink_webhook(payload: dict = Body(...)):
    doc_id = payload.get("document_id")
    status = payload.get("status")
    
    if not doc_id or not status:
        return {"status": "ignored"}
        
    digink_doc = await db.digink_documents.find_one({"diginkDocumentId": doc_id})
    if not digink_doc:
        return {"status": "not_found"}
        
    await db.digink_documents.update_one(
        {"diginkDocumentId": doc_id},
        {"$set": {"status": status, "updatedAt": datetime.utcnow().isoformat()}}
    )
    
    if status == "signed":
        await db.contracts.update_one(
            {"_id": ObjectId(digink_doc["contractId"])},
            {"$set": {"stage": "Signed", "status": "Approved"}}
        )
        
        contract = await db.contracts.find_one({"_id": ObjectId(digink_doc["contractId"])})
        if contract:
            await add_notification(
                for_user=contract.get("submittedBy"),
                message=f"Contract '{contract.get('title')}' has been fully signed via DigInk!",
                type="success",
                contract_id=str(contract["_id"]),
                contract_title=contract.get("title")
            )
            
    return {"status": "success"}

@app.get("/api/contracts/{contract_id}/digink-status")
async def get_digink_status(contract_id: str):
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(400, "Invalid contract ID")
    
    digink_doc = await db.digink_documents.find_one({"contractId": contract_id})
    if not digink_doc:
        return {"status": "Not Sent"}
    
    return {
        "status": digink_doc.get("status"),
        "diginkDocumentId": digink_doc.get("diginkDocumentId"),
        "updatedAt": digink_doc.get("updatedAt")
    }
>>>>>>> Stashed changes
